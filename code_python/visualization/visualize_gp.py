import json
from pathlib import Path
from typing import List, Optional, Any, Dict
import os 
import re

# visualization imports
# import cmcrameri.cm as cmc
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from matplotlib.collections import PolyCollection
from matplotlib.transforms import Bbox

import numpy as np
import pandas as pd
import seaborn as sns

#statistics imports
# import krippendorff
import pingouin as pg
from scipy.stats import kendalltau
from sklearn.metrics import auc
from scipy.stats import wilcoxon

#docs
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# internal imports
from visualization_setting import set_plot_style, save_img_path, ensure_long_path

set_plot_style()

HERE = Path(__file__).resolve().parent
# DATASETS = HERE.parent.parent / "datasets" / "Validation datasets"
RESULTS = HERE.parent.parent / "results"





def plot_average_feature_importances(scores_data: Dict[str, Any], save_loc: Path, file_extension: str, figsize: tuple=(10,7)) -> None:
    all_model_fi: List[Dict[str, float]] = []
    all_shap_fi: List[Dict[str, float]] = []

    # Collect model importances
    if isinstance(scores_data.get("feature_importance_MDI"), list):
        all_model_fi.extend(scores_data["feature_importance_MDI"])
    else:
        for value in scores_data.values():
            if isinstance(value, dict) and isinstance(value.get("feature_importance_MDI"), list):
                all_model_fi.extend(value["feature_importance_MDI"])

    # Collect SHAP importances
    if isinstance(scores_data.get("feature_importance_SHAP"), list):
        all_shap_fi.extend(scores_data["feature_importance_SHAP"])
    else:
        for value in scores_data.values():
            if isinstance(value, dict) and isinstance(value.get("feature_importance_SHAP"), list):
                all_shap_fi.extend(value["feature_importance_SHAP"])

    # Build DataFrames
    df_model = pd.DataFrame(all_model_fi) if all_model_fi else pd.DataFrame()
    df_shap = pd.DataFrame(all_shap_fi) if all_shap_fi else pd.DataFrame()
    if df_model.empty and df_shap.empty:
        raise ValueError("No feature importance data found in scores_data.")

    # Compute means and stds
    mean_model, std_model, mean_shap, std_shap = None, None, None, None
    if not df_model.empty:
        mean_model = df_model.mean()
        std_model = df_model.std()
    if not df_shap.empty:
        mean_shap = df_shap.mean()
        std_shap = df_shap.std()

    # Determine feature order (prefer model if available)
    if mean_model is not None:
        features = mean_model.sort_values(ascending=False).index
    else:
        features = mean_shap.sort_values(ascending=False).index

    # Align both importances
    if mean_model is not None:
        mean_model = mean_model.reindex(features, fill_value=0)
        std_model = std_model.reindex(features, fill_value=0)
    if mean_shap is not None:
        mean_shap = mean_shap.reindex(features, fill_value=0)
        std_shap = std_shap.reindex(features, fill_value=0)

    # ---- Select top 15 most important features ----
    if mean_model is not None and mean_shap is not None:
        combined_mean = (mean_model + mean_shap) / 2
    elif mean_model is not None:
        combined_mean = mean_model
    else:
        combined_mean = mean_shap

    top_features = combined_mean.sort_values(ascending=False).head(15).index

    # Restrict to top 15
    mean_model = mean_model[top_features] if mean_model is not None else None
    std_model = std_model[top_features] if std_model is not None else None
    mean_shap = mean_shap[top_features] if mean_shap is not None else None
    std_shap = std_shap[top_features] if std_shap is not None else None
    features = top_features
    # ------------------------------------------------------------

    # ----- NEW VIOLIN + SWARM PLOT SECTION -----
    # fig, ax = plt.subplots(figsize=figsize)

    # # Build long-format dataframe
    # plot_data = []

    # if not df_model.empty:
    #     for feat in features:
    #         for val in df_model[feat].dropna().values:
    #             plot_data.append([feat, val, "Model Importance"])

    # if not df_shap.empty:
    #     for feat in features:
    #         for val in df_shap[feat].dropna().values:
    #             plot_data.append([feat, val, "SHAP Importance"])

    # plot_df = pd.DataFrame(plot_data, columns=["Feature", "Value", "Type"])

    # Colors: red for model, blue for shap
    # palette = {
    #     "Model Importance": "#BB3A5A",  # red
    #     "SHAP Importance": "#1f77b4",   # blue
    # }

    # # VIOLIN PLOTS (side-by-side)
    # sns.violinplot(
    #     data=plot_df,
    #     x="Feature",
    #     y="Value",
    #     hue="Type",
    #     palette=palette,
    #     inner=None,
    #     linewidth=1.2,
    #     cut=0,
    #     dodge=True,
    #     ax=ax,
    # )

    # # SWARMPLOTS
    # sns.swarmplot(
    #     data=plot_df,
    #     x="Feature",
    #     y="Value",
    #     hue="Type",
    #     dodge=True,
    #     palette=palette, 
    #     size=3,
    #     alpha=0.7,
    #     ax=ax,
    # )

    # # Fix duplicate legends (remove swarm's entries)
    # handles, labels = ax.get_legend_handles_labels()
    # ax.legend(handles[:2], labels[:2], frameon=False)

    # ax.set_xticklabels(features, rotation=75, ha="right")
    # ax.set_ylabel("Feature Importance")
    # ax.set_xlabel("Top 15 Features")
    # ax.grid(axis="y", linestyle="--", alpha=0.4)

    # plt.tight_layout()
    # save_img_path(save_loc / "feature importance", f"feature_importance_top15_{file_extension}.png")
    # # plt.show()
    # plt.close()
    return df_model, df_shap



def get_lengthscale_stat(
        scores: dict,
        expert_rank=None,
        feature_stability=True,
        mean_std=True
    ) -> Dict:

    all_rows = []
    feat_samples = {}   # <-- collect raw samples for mean/std over seeds, folds, draws

    for key, value in scores.items():
        if isinstance(value, dict) and "test_lengthscale" in value:
            fold_list = value["test_lengthscale"]

            for fold_dict in fold_list:
                row_data = {}

                for feat_name, samples in fold_dict.items():
                    if samples is not None:
                        samples = np.asarray(samples)

                        # Row dataframe: mean across draws only for df_ls
                        row_data[feat_name] = np.mean(samples)

                        # For global stats: keep ALL draws across folds + seeds
                        if feat_name not in feat_samples:
                            feat_samples[feat_name] = []
                        feat_samples[feat_name].append(samples.flatten())
                    else:
                        row_data[feat_name] = np.nan

                all_rows.append(row_data)

    # main df like before (mean over draws per fold/seed)
    df_ls = pd.DataFrame(all_rows)

    # ---- NEW: mean + std over ALL seeds × folds × draws ----
    mean_std_stats = {}
    if mean_std or expert_rank is not None:
        for feat, arr_list in feat_samples.items():
            arr = np.concatenate(arr_list)
            mean_std_stats[feat] = {
                "mean": np.mean(arr),
                "std": np.std(arr, ddof=1)
            }
            df_mean_std = pd.DataFrame(mean_std_stats).T

    kendalls_w_result=  kendalls_w(df_ls, tie_corrected=False)["Kendall's W"] if feature_stability else None
        

    kendall_tau_result = None
    if expert_rank is not None:
        expert_series = expert_rank.iloc[0].copy()

        assert set(expert_series.index) == set(df_mean_std.index), \
            "Features in expert_rank and length scale mean must match exactly!"

        ls_rank = df_mean_std["mean"].rank(method="average", ascending=True)
        ls_rank_aligned = ls_rank.loc[expert_series.index]
        print("expert rank\n", expert_series)
        print("model rank (aligned)\n", ls_rank_aligned)
        print("real ls values\n", df_mean_std["mean"].loc[expert_series.index])
        kendall_tau_result = kendalltau(
            expert_series.values,
            ls_rank_aligned.values
        ).statistic

    stat_results = {
        "kendalls_w":kendalls_w_result,
        "kendall_tau": kendall_tau_result,
        "mean_std": mean_std_stats if mean_std else None,
    }
    # kendallTau_results = kendalltau(x=, y=)
    return stat_results


def kendalls_w(df_input, tie_corrected=True):
    """
    Kendall's W for agreement across raters (rows) on items (columns).
    If tie_corrected=True, applies the standard tie correction.
    """
    m = len(df_input)          # raters / folds
    n = len(df_input.columns)  # items / features

    # Rank each row
    ranks = df_input.rank(axis=1, ascending=False, method="average")

    # Sum of ranks for each feature
    R = ranks.sum(axis=0)
    R_bar = m * (n + 1) / 2.0

    # S term
    S = np.sum((R - R_bar) ** 2)

    if not tie_corrected:
        # Classic denominator
        S_max = (m**2 * (n**3 - n)) / 12.0
        W = S / S_max if S_max != 0 else np.nan
        chi_sq = m * (n - 1) * W
        df_chi = n - 1
        return {
            "m (Runs)": m,
            "n (Features)": n,
            "Kendall's W": W,
            "Chi-square": chi_sq,
            "Degrees of Freedom": df_chi,
            "Tie Corrected": False
        }

    # -------- Tie corrected version --------
    T_total = 0
    for _, row in df_input.iterrows():
        counts = row.value_counts()
        ties = counts[counts > 1]
        if len(ties) > 0:
            T_total += np.sum(ties**3 - ties)

    denom = (m**2 * (n**3 - n)) - (m * T_total)
    W = (12 * S) / denom if denom != 0 else np.nan

    chi_sq = m * (n - 1) * W
    df_chi = n - 1

    return {
        "m (Runs)": m,
        "n (Features)": n,
        "Kendall's W": W,
        "Chi-square": chi_sq,
        "Degrees of Freedom": df_chi,
        "Tie Corrected": True,
        "Total Tie Correction Term (Σ(t^3−t))": T_total
    }



def get_scores(data: Dict, metric: List[str]):
    special = {"rmse", "r2", "mae"}

    return {
        score: (
            f"{round(data[f'{score}_avg'], 2)} ± {round(data[f'{score}_stdev'], 2)}"
            if score.lower() in special
            else f"{int(round(data[score]))}"
        )
        for score in metric
    }


def create_word_table_table(rows_data, folder_path, file_name="results_gp_table.docx"):
    document = Document()

    rows = len(rows_data) + 3
    cols = 9   # <-- 9 columns now (Combination + 2 models × 3 metrics)
    table = document.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"

    # ---------------- HEADER LEVEL 1 ----------------
    h = table.rows[0].cells
    h[0].text = "Combination"
    h[3].text = "GPyro HMC"
    h[6].text = "GPyTorch MAP"

    h[0].merge(h[2])   # Combination spans Count/FP/Mixing
    h[3].merge(h[5])   # GPyro HMC spans RMSE / R² / runtime
    h[6].merge(h[8])   # GPyTorch MAP spans RMSE / R² / runtime

    # ---------------- HEADER LEVEL 2 ----------------
    h2 = table.rows[1].cells
    h2[0].text = "Kernel"
    h2[3].text = "RMSE"
    h2[4].text = "R²"
    h2[5].text = "Run time (sec)"
    h2[6].text = "RMSE"
    h2[7].text = "R²"
    h2[8].text = "Run time (sec)"
    h2[0].merge(h2[1])

    # ---------------- HEADER LEVEL 3 ----------------
    h3 = table.rows[2].cells
    h3[0].text = "Count"
    h3[1].text = "FP"
    h3[2].text = "Mixing method"

    # ---- VERTICAL MERGES ----
    table.rows[1].cells[2].merge(table.rows[2].cells[2])
    for i in range(3, 9):
        table.rows[1].cells[i].merge(table.rows[2].cells[i])

    # ---------------- DATA ROWS ----------------
    r = 3
    for row in rows_data:
        cells = table.rows[r].cells
        for c, v in enumerate(row):
            cells[c].text = str(v)
        r += 1

    # ---------------- STYLING ----------------
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)

    os.makedirs(folder_path, exist_ok=True)
    document.save(folder_path / file_name)
    print(f"Saved Word table → {file_name}")



PLS_Ranks = pd.DataFrame([{
                "Xn": 1,
                "Mw (g/mol)": 2,
                "Concentration (mg/ml)": 3,
                "Temperature SANS/SLS/DLS/SEC (K)": 4,
                "polymer dD": 5,
                "solvent dD": 6,
                "polymer dP": 7,
                "solvent dP": 8,
                "polymer dH": 9,
                "solvent dH": 10,
                "fp": 11,
                "PDI": 12
            }], 
            index=["expert_rank"]
            )

BMS_Ranks = pd.DataFrame([{
                "HOMO_D (eV)": 1,
                "LUMO_A (eV)": 2,
                "Eg_D (eV)": 3,
                "Eg_A (eV)": 4,
                "Ehl_D (eV)": 5,
                "Ehl_A (eV)": 6,
                "HOMO_A (eV)": 7,
                "LUMO_D (eV)": 8,
                "fp_Donor": 9,
                "fp_Acceptor": 10,
                "D:A ratio (m/m)": 11,
                "temperature of thermal annealing": 12,
                "solvent additive conc. (% v/v)": 13,
                "HTL energy level (eV)": 14,
                "ETL energy level (eV)": 15
            }],
            index=["expert_rank"]
            )

PAPER = {
    "Robust Learning from Literature Data_Model Generalizability and Uncertainty for Predicting Conjugated Polymer Solution Conformation": {
        "target": ["target_log Rg (nm)"],
        "expert_impt": PLS_Ranks
    },
    "Beyond molecular structure_ critically assessing machine learning for designing organic photovoltaic materials and devices": {
        "target": ["target_calculated PCE (%)"],
        "expert_impt": BMS_Ranks
    },
    "Machine Learning for Polymer Design to Enhance Pervaporation-Based Organic Recovery": {
        "target": ["target_log (Separation factor)", "target_log (Total flux)"],
        "expert_impt": None
    },
    "Machine Learning-Enabled Prediction and High-Throughput Screening of Polymer Membranes for Pervaporation Separation": {
        "target": ["target_log (Separation factor)", "target_log (Total flux)"],
        "expert_impt": None
    },
    "Understanding and Designing a High-Performance Ultrafiltration Membrane Using Machine Learning": {
        "target": [
            "target_flux decline ratio (%)",
            "target_flux recovery ratio (%)",
            "target_irreversible fouling ratio(%)",
            "target_organic compound removal (%)",
            "target_reversible fouling ratio (%)",
            r"target_water permeability (LMH\bar)"
        ],
        "expert_impt": None
    },
    "Miniaturization of Popular Reactions from the Medicinal Chemists Toolbox for Ultrahigh_Throughput Experimentation": {
        "target": ["target_Approx Conv (%)"],
        "expert_impt": None
    },
}



fp_sk_kernels = ["TanimotoMatern32", "TanimotoMatern52", "TanimotoRBF", "Tanimoto"]
fp_bit_kernels = ["Matern32", "Matern52", "RBF"]
count_kernels    = ["Matern32", "Matern52", "RBF"]
mixing_methods  = ["sum", "product", "averageProduct"]
mixing_labels   = {"sum": "Sum", "product": "Product", "averageProduct": "Av(count)×FP"}
TREE_MODELS = {"RF", "XGBR", "NGB"}
GNN_MODELS = {"GNN", "GCN", "GAT", "GIN", "MPNN", "DMPNN"}
MODEL_TYPE_COLORS = {
    "tree": "#7093B9",
    "gp": "#E45756",
    "gnn": "#72B7B2",
}
MODELS= ["RF", "XGBR", "NGB", "GpyroHMC", "GPytorchMAP"]
DEFAULT_SCORE_METRICS = ["rmse", "r2", "mae", "cvpp_ama", "nll", "ece", "Cv", "RUSC"]
BASE_MASTER_RESULT_COLUMNS = [
    "paper",
    "target",
    "model",
    "fp kernel",
    "count kernel",
    "mixing method",
]
TIME_COLUMNS = [
    "Running time (GPU)",
    "Running time (CPU)",
]



# def plot_barplot(model_specs, save_dir: Path, figsize=(6, 4)):
#     data = build_barplot_data(model_specs)

#     labels = [spec["label"] for spec in model_specs]
#     means = [np.mean(data[label]) for label in labels]
#     stds = [np.std(data[label]) for label in labels]

#     fig, ax = plt.subplots(figsize=figsize)

#     x = np.arange(len(labels))
#     bars = ax.bar(
#         x,
#         means,
#         yerr=stds,
#         capsize=5,
#         color=["#AA5A6E", "#AA5A6E","#AA5A6E","#AA5A6E", "#6B9DB4"],
#         edgecolor="white",
#         linewidth=0.8,
#         error_kw={"elinewidth": 1.5, "ecolor": "black", "capthick": 1.5},
#     )

#     for bar, mean, std in zip(bars, means, stds):
#         ax.text(
#             bar.get_x() + bar.get_width() / 2,
#             bar.get_height() + std + 0.005,
#             f"{mean:.2f}",
#             ha="center",
#             va="bottom",
#             fontsize=13,
#         )

#     ax.set_xticks(x)
#     ax.tick_params(axis="both", labelsize=14)
#     ax.set_xticklabels(labels, fontsize=14, rotation=45)
#     ax.set_ylabel("R²", fontsize=16, fontweight="bold")
#     ax.set_xlabel("Model", fontsize=16, fontweight="bold")
#     ax.set_ylim(0, 1.0)
#     ax.set_yticks(np.arange(0, 1.1, 0.1))

#     plt.tight_layout()
#     save_img_path(save_dir, "Tree_based_model_result_across_datasets.png")
#     plt.close(fig)
#     print(f"Saved barplot at {save_dir}")





def _is_tree_model(model: str) -> bool:
    return model in TREE_MODELS


def _score_file_templates(
    model: str,
    fp_k: Optional[str] = None,
    count_k: Optional[str] = None,
    mix_method: Optional[str] = None,
    use_gpu: bool = False,
) -> List[str]:
    if _is_tree_model(model):
        if use_gpu:
            return []
        return [f"(ECFP3_count_512-COUNT)_{model}_hypOFF_Standard_Standard_scores"]

    device_suffix = "_GPU" if use_gpu else ""
    return [
        (
            f"(ECFP3_count_512-COUNT)_"
            f"({model}_{fp_k}-{count_k}_{mix_method})"
            f"{suffix}_hypOFF_Standard_Standard{device_suffix}_scores"
        )
        for suffix in ["", "_mean"]
    ]


def _find_score_path(
    paper_loc: Path,
    model: str,
    fp_k: Optional[str] = None,
    count_k: Optional[str] = None,
    mix_method: Optional[str] = None,
    use_gpu: bool = False,
) -> Optional[Path]:
    for tmpl in _score_file_templates(model, fp_k, count_k, mix_method, use_gpu):
        path = ensure_long_path(paper_loc / f"{tmpl}.json")
        if path.exists():
            return path
    return None


def _read_json(path: Optional[Path]) -> Optional[Dict[str, Any]]:
    if path is None:
        return None

    with open(path, "r") as f:
        return json.load(f)


def _metric_score(
    data: Optional[Dict[str, Any]],
    metric: str,
    suffix: str,
) -> Optional[Any]:
    if data is None:
        return None

    return data.get(f"{metric}_{suffix}")


def _seed_fold_scores(
    data: Optional[Dict[str, Any]],
    metric: str,
) -> Optional[List[Any]]:
    if data is None:
        return None

    scores = []
    seeds = [key for key, value in data.items() if isinstance(value, dict)]

    for seed in sorted(
        seeds,
        key=lambda value: (0, int(value)) if str(value).isdigit() else (1, str(value)),
    ):
        values = data[seed].get(f"test_{metric}")
        if isinstance(values, list):
            scores.extend(values)

    return scores or None


def _load_score_files(
    paper_loc: Path,
    model: str,
    fp_k: Optional[str] = None,
    count_k: Optional[str] = None,
    mix_method: Optional[str] = None,
) -> tuple[Optional[Dict[str, Any]], Optional[Dict[str, Any]]]:
    cpu_data = _read_json(_find_score_path(paper_loc, model, fp_k, count_k, mix_method))
    gpu_data = None

    if not _is_tree_model(model):
        gpu_data = _read_json(
            _find_score_path(
                paper_loc,
                model,
                fp_k,
                count_k,
                mix_method,
                use_gpu=True,
            )
        )

    return cpu_data, gpu_data


def _metric_result_columns(score_metrics: List[str]) -> List[str]:
    columns = []
    for metric in score_metrics:
        columns.extend([
            f"{metric}_avg",
            f"{metric}_stdev",
            f"{metric}_seed_fold_scores",
        ])
    return columns


def _master_result_columns(score_metrics: List[str]) -> List[str]:
    return BASE_MASTER_RESULT_COLUMNS + _metric_result_columns(score_metrics) + TIME_COLUMNS


def _score_row_values(
    cpu_data: Optional[Dict[str, Any]],
    gpu_data: Optional[Dict[str, Any]],
    score_metrics: List[str],
) -> Dict[str, Any]:
    scores = {}
    for metric in score_metrics:
        scores.update({
            f"{metric}_avg": _metric_score(cpu_data, metric, "avg"),
            f"{metric}_stdev": _metric_score(cpu_data, metric, "stdev"),
            f"{metric}_seed_fold_scores": _seed_fold_scores(cpu_data, metric),
        })

    scores.update({
        "Running time (GPU)": None if gpu_data is None else gpu_data.get("run_time_sec"),
        "Running time (CPU)": None if cpu_data is None else cpu_data.get("run_time_sec"),
    })
    return scores


def load_score(
    paper_loc: Path,
    model: str,
    fp_k: Optional[str] = None,
    count_k: Optional[str] = None,
    mix_method: Optional[str] = None,
    score_metrics: Optional[List[str]] = None,
) -> Dict[str, Any]:
    """Load score values for one master-results row."""
    score_metrics = DEFAULT_SCORE_METRICS if score_metrics is None else score_metrics
    cpu_data, gpu_data = _load_score_files(paper_loc, model, fp_k, count_k, mix_method)
    return _score_row_values(cpu_data, gpu_data, score_metrics)


def _save_master_performance_data(df: pd.DataFrame, save_path: Path) -> None:
    save_path = Path(save_path)
    os.makedirs(ensure_long_path(save_path.parent), exist_ok=True)

    df.to_pickle(ensure_long_path(Path(f"{save_path}.pkl")))

    output = df.copy()
    for column in output.columns:
        if column.endswith("_seed_fold_scores"):
            output[column] = output[column].apply(
                lambda value: json.dumps(value) if value is not None else None
            )
    output.to_csv(ensure_long_path(Path(f"{save_path}.csv")), index=False)


def build_master_performance_data(
    save_path: Optional[Path] = RESULTS / "master_performance_data",
    score_metrics: Optional[List[str]] = None,
) -> pd.DataFrame:
    score_metrics = (
        DEFAULT_SCORE_METRICS
        if score_metrics is None
        else list(dict.fromkeys(score_metrics))
    )
    rows = []

    for paper_name, paper_info in PAPER.items():
        for target in paper_info["target"]:
            paper_loc = RESULTS / paper_name / target

            for model in MODELS:
                if _is_tree_model(model):
                    cpu_data, gpu_data = _load_score_files(
                        paper_loc=paper_loc,
                        model=model,
                    )
                    score_info = _score_row_values(cpu_data, gpu_data, score_metrics)
                    rows.append({
                        "paper": paper_name,
                        "target": target,
                        "model": model,
                        "fp kernel": None,
                        "count kernel": None,
                        "mixing method": None,
                        **score_info,
                    })
                    continue

                for fp_k in fp_sk_kernels + fp_bit_kernels:
                    for count_k in count_kernels:
                        for mix in mixing_methods:
                            cpu_data, gpu_data = _load_score_files(
                                paper_loc=paper_loc,
                                model=model,
                                fp_k=fp_k,
                                count_k=count_k,
                                mix_method=mix,
                            )
                            score_info = _score_row_values(cpu_data, gpu_data, score_metrics)
                            rows.append({
                                "paper": paper_name,
                                "target": target,
                                "model": model,
                                "fp kernel": fp_k,
                                "count kernel": count_k,
                                "mixing method": mix,
                                **score_info,
                            })

    df = pd.DataFrame(rows, columns=_master_result_columns(score_metrics), dtype=object)
    if save_path is not None:
        _save_master_performance_data(df, save_path)

    return df



def add_legend_box(
    indices, legend, color="grey", linestyle="dashed", alpha=1.0, linewidth=1.5, pad=5
):
    # Utilitary function to add a box in the legend of a figure.
    fig = legend.figure
    renderer = fig.canvas.get_renderer()

    texts = legend.get_texts()
    handles = legend.legend_handles

    text_boxes = [texts[i].get_window_extent(renderer) for i in indices]
    handle_boxes = [handles[i].get_window_extent(renderer) for i in indices]

    all_boxes = text_boxes + handle_boxes

    if not all_boxes:
        return

    full_box = Bbox.union(all_boxes).padded(pad)

    trans = fig.transFigure.inverted()
    full_box_fig = trans.transform_bbox(full_box)

    rect = patches.Rectangle(
        (full_box_fig.x0, full_box_fig.y0),
        full_box_fig.width,
        full_box_fig.height,
        transform=fig.transFigure,
        facecolor="none",
        edgecolor=color,
        linestyle=linestyle,
        linewidth=linewidth,
        alpha=alpha,
        zorder=5,
    )
    fig.patches.append(rect)


def _coerce_score_list(value: Any) -> Optional[List[Any]]:
    if isinstance(value, list):
        return value
    if isinstance(value, np.ndarray):
        return value.tolist()
    if isinstance(value, str):
        try:
            parsed = json.loads(value)
        except json.JSONDecodeError:
            return None
        return parsed if isinstance(parsed, list) else None
    return None


def _selection_values(values: Any) -> Optional[List[Any]]:
    if values is None:
        return None
    if isinstance(values, str):
        return [values]
    return list(values)


def _filter_selection(
    df: pd.DataFrame,
    column: str,
    values: Any,
    keep_missing: bool = False,
) -> pd.DataFrame:
    selected = _selection_values(values)
    if selected is None:
        return df

    selected_lower = {str(value).lower() for value in selected}
    column_values = df[column].astype(str).str.lower()
    mask = column_values.isin(selected_lower)
    if keep_missing:
        mask = mask | df[column].isna()
    return df[mask].copy()


def _kernel_label(row: pd.Series, include_model: bool = False) -> str:
    model = str(row["model"])
    if pd.isna(row["fp kernel"]) and pd.isna(row["count kernel"]):
        return model

    label = f"{row['fp kernel']}-{row['count kernel']}_{row['mixing method']}"
    return f"{model}: {label}" if include_model else label


def _metric_higher_is_better(metric: str) -> bool:
    return metric in {"r2", "RUSC", "cvpp_ama"}


def _expand_master_scores_for_profile(
    df: pd.DataFrame,
    metric: str,
    model: Any,
    fp_kernels: Any = None,
    count_kernels: Any = None,
    mixing_methods: Any = None,
) -> pd.DataFrame:
    score_col = f"{metric}_seed_fold_scores"
    use_seed_fold_scores = True
    if score_col not in df.columns and metric in df.columns:
        score_col = metric
        use_seed_fold_scores = False
    if score_col not in df.columns:
        raise ValueError(f"Missing required column: {score_col}")

    rows = []
    selected_models = _selection_values(model)
    model_df = _filter_selection(df, "model", selected_models)
    model_df = _filter_selection(model_df, "fp kernel", fp_kernels, keep_missing=True)
    model_df = _filter_selection(model_df, "count kernel", count_kernels, keep_missing=True)
    model_df = _filter_selection(model_df, "mixing method", mixing_methods, keep_missing=True)
    include_model = selected_models is None or len(selected_models) > 1

    for _, row in model_df.iterrows():
        if use_seed_fold_scores:
            scores = _coerce_score_list(row[score_col])
            if scores is None:
                continue
        else:
            scores = [row[score_col]]

        dataset = row["paper"]
        target = row["target"]
        kernel = _kernel_label(row, include_model=include_model)
        for repeat_idx, score in enumerate(scores):
            try:
                score = float(score)
            except (TypeError, ValueError):
                continue
            if np.isnan(score):
                continue
            rows.append({
                "dataset": dataset,
                "target": target,
                "repeat": repeat_idx,
                "kernel": kernel,
                "model": row["model"],
                "fp kernel": row["fp kernel"],
                "count kernel": row["count kernel"],
                "mixing method": row["mixing method"],
                metric: score,
            })

    return pd.DataFrame(rows)



def _model_config_label(row: pd.Series, include_kernel_config: bool = True) -> str:
    model = str(row["model"])
    if not include_kernel_config:
        return model
    if pd.isna(row["fp kernel"]) and pd.isna(row["count kernel"]):
        return model

    # mix = mixing_labels.get(str(row["mixing method"]))
    suffix = "SK" if "tanimoto" in row["fp kernel"].lower() else "Bitwise"
    return f"{model} ({suffix})"


def _model_type_color(model: Any) -> str:
    model_name = str(model)
    model_upper = model_name.upper()
    if model_name in TREE_MODELS:
        return MODEL_TYPE_COLORS["tree"]
    if model_name in GNN_MODELS or any(token in model_upper for token in GNN_MODELS):
        return MODEL_TYPE_COLORS["gnn"]
    if "GP" in model_upper:
        return MODEL_TYPE_COLORS["gp"]
    return "#808080"


def _kernel_triple_values(kernel_triples: Any) -> Optional[List[tuple]]:
    if kernel_triples is None:
        return None
    if (
        isinstance(kernel_triples, tuple)
        and len(kernel_triples) == 3
        and not any(isinstance(value, (list, tuple)) for value in kernel_triples)
    ):
        return [kernel_triples]

    triples = list(kernel_triples)
    for triple in triples:
        if not isinstance(triple, (list, tuple)) or len(triple) != 3:
            raise ValueError(
                "kernel_triples must be a triple or a list of "
                "(fp_kernel, count_kernel, mixing_method) triples."
            )
    return triples


def _filter_kernel_triples(
    df: pd.DataFrame,
    kernel_triples: Any,
    keep_missing: bool = True,
) -> pd.DataFrame:
    triples = _kernel_triple_values(kernel_triples)
    if triples is None:
        return df

    selected = {
        (str(fp).lower(), str(count).lower(), str(mix).lower())
        for fp, count, mix in triples
    }
    row_triples = zip(
        df["fp kernel"].astype(str).str.lower(),
        df["count kernel"].astype(str).str.lower(),
        df["mixing method"].astype(str).str.lower(),
    )
    mask = pd.Series(
        [row_triple in selected for row_triple in row_triples],
        index=df.index,
    )
    if keep_missing:
        mask = mask | (
            df["fp kernel"].isna()
            & df["count kernel"].isna()
            & df["mixing method"].isna()
        )
    return df[mask].copy()


def plot_model_comparison(
    df: pd.DataFrame,
    metric: str = "r2",
    model: Any = "GPytorchMAP",
    kernel_triples: Any = None,
    include_kernel_config: bool = True,
    dataset_as_experiment_points: bool = False,
    figsize: tuple = (9, 5),
    fontsize: int = 12,
    title: Optional[str] = None,
    x_label: str = "Model",
    y_label: Optional[str] = None,
    x_tick_rotation: int = 35,
    y_lim: Optional[tuple] = None,
    log_y: bool = False,
    show: bool = True,
    high_quality: bool = True,
    save_dir: Optional[Path] = HERE / "result_analysis",
    file_name: Optional[str] = None,
) -> pd.DataFrame:
    """
    Make a box plot of a metric score for selected models/kernels.

    Score metrics use ``<metric>_seed_fold_scores``. Scalar columns such as
    ``Running time (CPU)`` are plotted directly across dataset/model rows.
    For GP models, pass ``kernel_triples`` as one triple or a list of
    ``(fp_kernel, count_kernel, mixing_method)`` triples.
    Set ``log_y=True`` to show the y-axis on a log scale.
    """
    df = _filter_kernel_triples(df, kernel_triples, keep_missing=True)
    is_scalar_metric = metric in df.columns and f"{metric}_seed_fold_scores" not in df.columns
    plot_df = _expand_master_scores_for_profile(
        df,
        metric=metric,
        model=model,
    )
    if plot_df.empty:
        raise ValueError(
            f"No {metric} values found for model={model} "
            f"with kernel_triples={kernel_triples}."
        )

    x_col = "model configuration"
    plot_df[x_col] = plot_df.apply(
        lambda row: _model_config_label(row, include_kernel_config),
        axis=1,
    )

    if dataset_as_experiment_points:
        group_cols = [
            "dataset",
            "target",
            x_col,
            "model",
            "fp kernel",
            "count kernel",
            "mixing method",
        ]
        plot_df = (
            plot_df.groupby(group_cols, dropna=False, as_index=False)[metric]
            .mean()
            .copy()
        )

    order = plot_df[x_col].drop_duplicates().tolist()
    violin_palette = {
        label: _model_type_color(plot_df.loc[plot_df[x_col] == label, "model"].iloc[0])
        for label in order
    }
    fig, ax = plt.subplots(figsize=figsize)
    sns.violinplot(
        data=plot_df,
        x=x_col,
        y=metric,
        hue=x_col,
        order=order,
        hue_order=order,
        ax=ax,
        width=0.65,
        palette=violin_palette,
        inner_kws=dict(box_width=10, whis_width=2, color=".65",solid_capstyle="round"),
        cut=0,
        linewidth=1,
        dodge=False,
        legend=False,
    )
    positions = list(range(len(order)))
    for collection in ax.collections:
        if not isinstance(collection, PolyCollection):
            continue
        for path in collection.get_paths():
            vertices = path.vertices
            if len(vertices) == 0:
                continue
            center = min(positions, key=lambda pos: abs(pos - np.median(vertices[:, 0])))
            vertices[:, 0] = np.maximum(vertices[:, 0], center)
    for line in ax.lines:
        x_data = np.asarray(line.get_xdata(), dtype=float)
        if len(x_data) == 0:
            continue
        center = min(positions, key=lambda pos: abs(pos - np.mean(x_data)))
        line.set_xdata(np.maximum(x_data, center))

    # if show_points:
    #     sns.stripplot(
    #         data=plot_df,
    #         x=x_col,
    #         y=metric,
    #         order=order,
    #         ax=ax,
    #         color="black",
    #         size=3,
    #         alpha=point_alpha,
    #         jitter=0.22,
    #     )
    #     for collection in ax.collections:
    #         if isinstance(collection, PolyCollection) or not hasattr(collection, "get_offsets"):
    #             continue
    #         offsets = collection.get_offsets()
    #         if len(offsets) == 0:
    #             continue
    #         for offset in offsets:
    #             center = min(positions, key=lambda pos: abs(pos - offset[0]))
    #             offset[0] = max(offset[0], center)
    #         collection.set_offsets(offsets)

    ax.set_xlabel(x_label, fontsize=fontsize, fontweight="bold")
    ax.set_ylabel(y_label or metric.capitalize(), fontsize=fontsize, fontweight="bold")
    if title is not None:
        ax.set_title(title, fontsize=fontsize + 2)
    ax.tick_params(axis="both", labelsize=fontsize-2)
    ax.set_xticks(range(len(order)))
    ax.set_xticklabels(
        order,
        rotation=x_tick_rotation,
        ha="right" if x_tick_rotation else "center",
    )
    if log_y:
        positive_values = pd.to_numeric(plot_df[metric], errors="coerce")
        if (positive_values.dropna() <= 0).any():
            raise ValueError("log_y=True requires all plotted values to be positive.")
        ax.set_yscale("log")

    if y_lim is not None:
        ax.set_ylim(*y_lim)
    elif log_y:
        ax.set_ylim(bottom=pd.to_numeric(plot_df[metric], errors="coerce").min() * 0.8)
    elif is_scalar_metric:
        ax.set_ylim(bottom=0)
    else:
        ax.set_ylim(0, 1.05)
    plt.tight_layout()

    if save_dir is not None:
        save_dir = ensure_long_path(Path(save_dir))
        os.makedirs(save_dir, exist_ok=True)
        if file_name is None:
            model_name = "_".join(str(value) for value in (_selection_values(model) or ["all"]))
            file_name = f"{model_name}_{metric}_boxplot.png"
        fig.savefig(
            ensure_long_path(save_dir / file_name),
            bbox_inches="tight",
            format="png",
            dpi=900 if high_quality else 100,
        )

    if show:
        plt.show()
    else:
        plt.close(fig)


def plot_model_profile_comparison(
    prof_results: pd.DataFrame,
    model: Any = None,
    kernel_triples: Any = None,
    metric: Optional[str] = None,
    auc_column: str = "auc",
    include_kernel_config: bool = True,
    figsize: tuple = (7, 5),
    fontsize: int = 12,
    title: Optional[str] = None,
    x_label: str = "Model",
    y_label: str = "Performance profile AUC",
    x_tick_rotation: int = 35,
    y_lim: tuple = (0, 1.05),
    show_values: bool = True,
    show: bool = True,
    high_quality: bool = True,
    save_dir: Optional[Path] = HERE / "result_analysis",
    file_name: Optional[str] = None,
) -> pd.DataFrame:
    """
    Draw a barplot of profile AUC values from ``performance_plot_with_ranks``.

    Pass ``kernel_triples`` as one triple or a list of
    ``(fp_kernel, count_kernel, mixing_method)`` triples to select exact GP
    kernel combinations. Tree models are kept when their kernel columns are
    empty.
    """
    required_columns = {
        "model",
        "fp kernel",
        "count kernel",
        "mixing method",
        auc_column,
    }
    missing_columns = required_columns.difference(prof_results.columns)
    if missing_columns:
        raise ValueError(
            "prof_results is missing required columns: "
            + ", ".join(sorted(missing_columns))
        )

    plot_df = prof_results.copy()
    plot_df = _filter_selection(plot_df, "model", model)
    plot_df = _filter_kernel_triples(plot_df, kernel_triples, keep_missing=True)
    if metric is not None and "metric" in plot_df.columns:
        plot_df = _filter_selection(plot_df, "metric", metric)

    plot_df[auc_column] = pd.to_numeric(plot_df[auc_column], errors="coerce")
    plot_df = plot_df.dropna(subset=[auc_column])
    if plot_df.empty:
        raise ValueError(
            f"No profile AUC rows found for model={model} "
            f"with kernel_triples={kernel_triples}."
        )

    x_col = "model configuration"
    plot_df[x_col] = plot_df.apply(
        lambda row: _model_config_label(row, include_kernel_config),
        axis=1,
    )
    selected_models = _selection_values(model)
    selected_triples = _kernel_triple_values(kernel_triples)
    if selected_models is None:
        order = plot_df[x_col].drop_duplicates().tolist()
    else:
        order = []
        model_values = plot_df["model"].astype(str).str.lower()
        fp_values = plot_df["fp kernel"].astype(str).str.lower()
        count_values = plot_df["count kernel"].astype(str).str.lower()
        mix_values = plot_df["mixing method"].astype(str).str.lower()
        missing_kernel_mask = (
            plot_df["fp kernel"].isna()
            & plot_df["count kernel"].isna()
            & plot_df["mixing method"].isna()
        )

        for model_name in selected_models:
            model_mask = model_values == str(model_name).lower()
            for label in plot_df.loc[model_mask & missing_kernel_mask, x_col].drop_duplicates():
                if label not in order:
                    order.append(label)

            if selected_triples is None:
                labels = plot_df.loc[model_mask & ~missing_kernel_mask, x_col].drop_duplicates()
                for label in labels:
                    if label not in order:
                        order.append(label)
                continue

            for fp_kernel, count_kernel, mix_method in selected_triples:
                triple_mask = (
                    (fp_values == str(fp_kernel).lower())
                    & (count_values == str(count_kernel).lower())
                    & (mix_values == str(mix_method).lower())
                )
                for label in plot_df.loc[model_mask & triple_mask, x_col].drop_duplicates():
                    if label not in order:
                        order.append(label)

        for label in plot_df[x_col].drop_duplicates():
            if label not in order:
                order.append(label)

    order_lookup = {label: idx for idx, label in enumerate(order)}
    plot_df["_plot_order"] = plot_df[x_col].map(order_lookup)
    plot_df = (
        plot_df.sort_values("_plot_order", kind="stable")
        .drop(columns="_plot_order")
        .copy()
    )
    bar_palette = {
        label: _model_type_color(plot_df.loc[plot_df[x_col] == label, "model"].iloc[0])
        for label in order
    }

    fig, ax = plt.subplots(figsize=figsize)
    sns.barplot(
        data=plot_df,
        x=x_col,
        y=auc_column,
        hue=x_col,
        order=order,
        hue_order=order,
        palette=bar_palette,
        width=0.67,
        dodge=False,
        legend=False,
        errorbar=None,
        ax=ax,
    )

    if show_values:
        for patch in ax.patches:
            height = patch.get_height()
            if np.isnan(height):
                continue
            ax.text(
                patch.get_x() + patch.get_width() / 2,
                height + 0.01,
                f"{height:.2f}",
                ha="center",
                va="bottom",
                fontsize=fontsize - 2,
            )

    ax.set_xlabel(x_label, fontsize=fontsize, fontweight="bold")
    ax.set_ylabel(y_label, fontsize=fontsize, fontweight="bold")
    if title is not None:
        ax.set_title(title, fontsize=fontsize + 2)
    ax.tick_params(axis="both", labelsize=fontsize - 2)
    ax.set_xticks(range(len(order)))
    ax.set_xticklabels(
        order,
        rotation=x_tick_rotation,
        ha="right" if x_tick_rotation else "center",
    )
    ax.set_ylim(*y_lim)
    plt.tight_layout()

    if save_dir is not None:
        save_dir = ensure_long_path(Path(save_dir))
        os.makedirs(save_dir, exist_ok=True)
        if file_name is None:
            model_name = "_".join(str(value) for value in (_selection_values(model) or ["all"]))
            file_name = f"{model_name}_profile_auc_barplot.png"
        fig.savefig(
            ensure_long_path(save_dir / file_name),
            bbox_inches="tight",
            format="png",
            dpi=900 if high_quality else 100,
        )

    if show:
        plt.show()
    else:
        plt.close(fig)

    return plot_df


def _rank_percentile_matrix(
    score_matrix: pd.DataFrame,
    higher_is_better: bool,
) -> pd.DataFrame:
    rank_matrix = score_matrix.rank(
        axis=1,
        ascending=not higher_is_better,
        method="min",
    )
    valid_counts = score_matrix.notna().sum(axis=1).replace(0, np.nan)
    return rank_matrix.div(valid_counts, axis=0)


def _profile_task_columns(profile_df: pd.DataFrame) -> List[str]:
    task_cols = [column for column in ["dataset", "target"] if column in profile_df.columns]
    if not task_cols:
        raise ValueError("profile_df must contain at least a dataset or target column.")
    return task_cols


def _profile_index_columns(
    profile_df: pd.DataFrame,
    include_repeat: bool,
) -> List[str]:
    index_cols = _profile_task_columns(profile_df)
    if include_repeat and "repeat" in profile_df.columns:
        index_cols = index_cols + ["repeat"]
    return index_cols


def _aggregate_profile_repeats(
    profile_df: pd.DataFrame,
    metric: str,
) -> pd.DataFrame:
    group_cols = [
        column
        for column in [
            "dataset",
            "target",
            "kernel",
            "model",
            "fp kernel",
            "count kernel",
            "mixing method",
        ]
        if column in profile_df.columns
    ]
    return (
        profile_df.groupby(group_cols, dropna=False, as_index=False)[metric]
        .mean()
        .copy()
    )


def _cat_gp_percentile_rank_matrix(
    profile_df: pd.DataFrame,
    metric: str,
    higher_is_better: bool,
    index_cols: List[str],
) -> pd.DataFrame:
    ranked_df = profile_df.copy()
    task_cols = _profile_task_columns(ranked_df)
    ranked_df["rank_repeat"] = ranked_df.groupby(
        task_cols,
        dropna=False,
    )[metric].rank(
        ascending=not higher_is_better,
        method="min",
    )
    max_rank = ranked_df.groupby(
        task_cols,
        dropna=False,
    )["rank_repeat"].transform("max")
    ranked_df["percentile_rank_repeat"] = ranked_df["rank_repeat"].div(
        max_rank.replace(0, np.nan)
    )

    return ranked_df.pivot_table(
        index=index_cols,
        columns="kernel",
        values="percentile_rank_repeat",
        aggfunc="first",
    )


def _profile_score_matrix(
    profile_df: pd.DataFrame,
    metric: str,
    index_cols: List[str],
) -> pd.DataFrame:
    return profile_df.pivot_table(
        index=index_cols,
        columns="kernel",
        values=metric,
        aggfunc="first",
    )


def _profile_curve_from_percentiles(
    percentiles: pd.Series,
    tau_grid: np.ndarray,
) -> Optional[np.ndarray]:
    values = percentiles.dropna().to_numpy(dtype=float)
    if len(values) == 0:
        return None

    return np.array([(values <= tau).mean() for tau in tau_grid], dtype=float)


def _profile_curves(
    percentile_matrix: pd.DataFrame,
    tau_grid: np.ndarray,
) -> Dict[str, np.ndarray]:
    curves = {}
    for kernel_name in percentile_matrix.columns:
        curve = _profile_curve_from_percentiles(
            percentile_matrix[kernel_name],
            tau_grid,
        )
        if curve is not None:
            curves[kernel_name] = curve

    return curves


def _repeat_profile_error_bands(
    profile_df: pd.DataFrame,
    metric: str,
    higher_is_better: bool,
    tau_grid: np.ndarray,
    error_band_stat: str,
) -> Dict[str, np.ndarray]:
    error_band_stat = error_band_stat.lower()
    if error_band_stat not in {"std", "sem"}:
        raise ValueError("error_band_stat must be either 'std' or 'sem'.")

    curves_by_kernel = {}
    if "repeat" not in profile_df.columns:
        return curves_by_kernel

    index_cols = _profile_index_columns(profile_df, include_repeat=False)
    for _, repeat_df in profile_df.groupby("repeat", dropna=False):
        repeat_percentile_matrix = _cat_gp_percentile_rank_matrix(
            repeat_df,
            metric,
            higher_is_better,
            index_cols,
        )
        repeat_curves = _profile_curves(repeat_percentile_matrix, tau_grid)
        for kernel_name, curve in repeat_curves.items():
            curves_by_kernel.setdefault(kernel_name, []).append(curve)

    error_bands = {}
    for kernel_name, curves in curves_by_kernel.items():
        curve_matrix = np.vstack(curves)
        if curve_matrix.shape[0] <= 1:
            error_bands[kernel_name] = np.zeros(curve_matrix.shape[1])
            continue

        error = np.nanstd(curve_matrix, axis=0, ddof=1)
        if error_band_stat == "sem":
            error = error / np.sqrt(curve_matrix.shape[0])
        error_bands[kernel_name] = error

    return error_bands


def _profile_auc_dataframe(
    profile_df: pd.DataFrame,
    aucs: Dict[str, float],
    sorted_names: List[str],
    metric: str,
) -> pd.DataFrame:
    metadata_cols = ["kernel", "model", "fp kernel", "count kernel", "mixing method"]
    metadata = (
        profile_df[metadata_cols]
        .drop_duplicates(subset=["kernel"])
        .set_index("kernel")
    )
    rows = []

    for rank, kernel_name in enumerate(sorted_names, start=1):
        meta = metadata.loc[kernel_name]
        fp_kernel = meta["fp kernel"]
        count_kernel = meta["count kernel"]
        kernel_combination = (
            None
            if pd.isna(fp_kernel) or pd.isna(count_kernel)
            else f"{fp_kernel}-{count_kernel}"
        )

        rows.append({
            "rank": rank,
            "metric": metric,
            "model": meta["model"],
            "fp kernel": fp_kernel,
            "count kernel": count_kernel,
            "mixing method": meta["mixing method"],
            "auc": aucs[kernel_name],
        })

    return pd.DataFrame(rows)


def performance_plot_with_ranks(
    df: pd.DataFrame,
    metric: str = "r2",
    model: Any = "GPytorchMAP",
    fp_kernels: Any = None,
    count_kernels: Any = None,
    mixing_methods: Any = None,
    kernel_triples: Any = None,
    p_threshold: float = 0.05,
    title: Optional[str] = None,
    show: bool = True,
    high_quality: bool = True,
    dataset_as_experiment_points: bool = False,
    show_error_band: bool = False,
    error_band_stat: str = "std",
    error_band_alpha: float = 0.18,
    fontsize: float = 12,
    figsize: tuple = (7, 5),
    legend_ncols: Optional[int] = None,
    legend_nrows: Optional[int] = None,
    save_dir: Optional[Path] = HERE / "result_analysis",
    file_name: Optional[str] = "performance_profile.png",
):
    """
    Plot model/kernel performance profiles from the master result dataset.

    Following the cat_gp profile code, each row is one
    ``dataset * target * seed/fold`` repeat for a model/kernel. Percentile
    ranks are computed within each ``dataset * target`` task using all selected
    model/kernel repeats, then the plotted profile is the CDF of those ranks.
    Returns a DataFrame with one AUC row per model/kernel combination.

    For exact GP kernel choices, pass ``kernel_triples`` as one triple or a
    list of ``(fp_kernel, count_kernel, mixing_method)`` triples. The older
    separate ``fp_kernels``, ``count_kernels``, and ``mixing_methods`` filters
    are still supported.

    If ``dataset_as_experiment_points`` is True, the main profile first averages
    seed/fold scores within each dataset/target so each target contributes one
    profile point. If ``show_error_band`` is True, this dataset-level mode is
    used and the shaded band is the std/sem across seed/fold repeat profiles.
    ``legend_ncols`` and ``legend_nrows`` control the legend layout above the
    plot. If both are omitted, up to three columns are used.
    """
    df = _filter_kernel_triples(df, kernel_triples, keep_missing=True)
    profile_df = _expand_master_scores_for_profile(
        df,
        metric=metric,
        model=model,
        fp_kernels=fp_kernels,
        count_kernels=count_kernels,
        mixing_methods=mixing_methods,
    )
    if profile_df.empty:
        raise ValueError(
            f"No {metric} seed/fold scores found for model={model} "
            f"with fp_kernels={fp_kernels}, count_kernels={count_kernels}, "
            f"mixing_methods={mixing_methods}, and kernel_triples={kernel_triples}."
        )

    higher_is_better = _metric_higher_is_better(metric)

    if show_error_band:
        dataset_as_experiment_points = True

    profile_rank_df = profile_df
    if dataset_as_experiment_points:
        profile_rank_df = _aggregate_profile_repeats(profile_df, metric)

    include_repeat = not dataset_as_experiment_points
    profile_index_cols = _profile_index_columns(
        profile_rank_df,
        include_repeat=include_repeat,
    )
    profile_score_matrix = _profile_score_matrix(
        profile_rank_df,
        metric,
        profile_index_cols,
    )
    percentile_matrix = _cat_gp_percentile_rank_matrix(
        profile_rank_df,
        metric,
        higher_is_better,
        profile_index_cols,
    )

    linn = np.linspace(0, 1, max(len(percentile_matrix), 2))
    all_percentages = _profile_curves(percentile_matrix, linn)
    error_bands = (
        _repeat_profile_error_bands(
            profile_df,
            metric,
            higher_is_better,
            linn,
            error_band_stat,
        )
        if show_error_band
        else {}
    )

    aucs = {
        kernel_name: auc(linn, percentages)
        for kernel_name, percentages in all_percentages.items()
    }
    sorted_names = [
        kernel_name
        for kernel_name, _ in sorted(aucs.items(), key=lambda x: x[1], reverse=True)
    ]
    auc_df = _profile_auc_dataframe(profile_df, aucs, sorted_names, metric)


    groups_of_same_ranks = []
    curr_l = [0]
    alternative = "greater" if higher_is_better else "less"
    for i in range(len(sorted_names)):
        if i == len(sorted_names) - 1:
            if len(curr_l) >= 2:
                groups_of_same_ranks.append(curr_l)
            continue

        paired_scores = profile_score_matrix[
            [sorted_names[i], sorted_names[i + 1]]
        ].dropna()
        if paired_scores.empty:
            continue

        try:
            _, p = wilcoxon(
                paired_scores[sorted_names[i]],
                paired_scores[sorted_names[i + 1]],
                alternative=alternative,
            )
        except ValueError:
            p = 1.0

        if p < p_threshold:
            if len(curr_l) >= 2:
                groups_of_same_ranks.append(curr_l)
            curr_l = [i + 1]
        else:
            curr_l.append(i + 1)

    fig, ax = plt.subplots(figsize=figsize)
    colors = sns.color_palette("tab20", n_colors=max(len(sorted_names), 1))

    for color, kernel_name in zip(colors, sorted_names):
        curve = np.asarray(all_percentages[kernel_name], dtype=float)
        ax.plot(
            linn,
            curve,
            color=color,
            label=f"{kernel_name}, {aucs[kernel_name]:.2f}",
        )
        if kernel_name in error_bands:
            error = np.asarray(error_bands[kernel_name], dtype=float)
            ax.fill_between(
                linn,
                np.clip(curve - error, 0, 1),
                np.clip(curve + error, 0, 1),
                color=color,
                alpha=error_band_alpha,
                linewidth=0,
            )

    if legend_ncols is not None and legend_ncols < 1:
        raise ValueError("legend_ncols must be at least 1.")
    if legend_nrows is not None and legend_nrows < 1:
        raise ValueError("legend_nrows must be at least 1.")

    if legend_ncols is None:
        if legend_nrows is None:
            legend_ncols = max(1, min(3, len(sorted_names)))
        else:
            legend_ncols = int(np.ceil(len(sorted_names) / legend_nrows))

    legend = ax.legend(
        loc="lower center",
        bbox_to_anchor=(0.5, 1.02),
        ncol=legend_ncols,
        borderaxespad=0,
        frameon=False,
        fontsize=fontsize,
    )

    fig.canvas.draw()

    ax.set_xlabel(r"$\tau$", fontsize=fontsize, fontweight="bold")
    ax.set_ylabel(r"$p_i(\tau)$", fontsize=fontsize, fontweight="bold")
    if title is not None:
        ax.set_title(
            title or f"{model} {metric} performance profile",
            fontsize=fontsize+2,
            fontweight="bold"
        )
    ax.yaxis.set_tick_params(labelsize=fontsize-2)
    ax.xaxis.set_tick_params(labelsize=fontsize-2)

    fig.canvas.draw()
    for group_rank in groups_of_same_ranks:
        add_legend_box(group_rank, legend, linewidth=1.2, alpha=1.0, pad=3)

    if save_dir is not None:
        save_dir = ensure_long_path(Path(save_dir))
        os.makedirs(save_dir, exist_ok=True)
        if file_name is None:
            file_name = f"{model}_{metric}_performance_profile.png"
        fig.savefig(
            ensure_long_path(save_dir / file_name),
            bbox_inches="tight",
            format="png",
            dpi=900 if high_quality else 100,
        )

    if show:
        plt.show()
    else:
        plt.close(fig)

    return auc_df


def plot_profile_auc_heatmap(
    auc_df: pd.DataFrame,
    model: Any = None,
    fp_kernels: Any = None,
    count_kernels: Any = None,
    mixing_methods: Any = None,
    auc_column: str = "auc",
    figsize: tuple = (12, 4),
    title: Optional[str] = None,
    vmin: float = 0,
    vmax: float = 1,
    fontsize: int = 12,
    show: bool = True,
    high_quality: bool = True,
    plot_heatmap: bool = True,
    save_dir: Optional[Path] = HERE / "result_analysis",
    file_name: Optional[str] = "profile_auc_heatmap.png",
) -> pd.DataFrame:
    """
    Plot profile-performance AUC values as a heatmap.

    The input should be the DataFrame returned by ``performance_plot_with_ranks``.
    Rows are hybridization methods and columns are FP/count kernel combinations.
    """
    required_columns = {
        "model",
        "fp kernel",
        "count kernel",
        "mixing method",
        auc_column,
    }
    missing_columns = required_columns.difference(auc_df.columns)
    if missing_columns:
        raise ValueError(
            "auc_df is missing required columns: "
            + ", ".join(sorted(missing_columns))
        )

    plot_df = auc_df.copy()
    plot_df = _filter_selection(plot_df, "model", model)
    plot_df = _filter_selection(plot_df, "fp kernel", fp_kernels)
    plot_df = _filter_selection(plot_df, "count kernel", count_kernels)
    plot_df = _filter_selection(plot_df, "mixing method", mixing_methods)
    plot_df = plot_df.dropna(
        subset=["fp kernel", "count kernel", "mixing method", auc_column]
    ).copy()

    if plot_df.empty:
        raise ValueError(
            "No profile-AUC rows available for the requested model/kernel filters."
        )

    plot_df[auc_column] = pd.to_numeric(plot_df[auc_column], errors="coerce")
    plot_df = plot_df.dropna(subset=[auc_column])
    if plot_df.empty:
        raise ValueError("No numeric profile-AUC values available to plot.")

    plot_df["hybridization method"] = plot_df["mixing method"].map(
        lambda value: mixing_labels.get(str(value), str(value))
    )
    plot_df["kernel combination"] = (
        plot_df["fp kernel"].astype(str)
        + "-"
        + plot_df["count kernel"].astype(str)
    )

    heatmap_matrix = plot_df.pivot_table(
        index="hybridization method",
        columns="kernel combination",
        values=auc_column,
        aggfunc="mean",
    )

    selected_mixes = _selection_values(mixing_methods) or globals()["mixing_methods"]
    row_order = [
        mixing_labels.get(str(method), str(method))
        for method in selected_mixes
        if mixing_labels.get(str(method), str(method)) in heatmap_matrix.index
    ]
    row_order.extend(row for row in heatmap_matrix.index if row not in row_order)
    heatmap_matrix = heatmap_matrix.reindex(row_order)

    selected_fp = _selection_values(fp_kernels) or (fp_sk_kernels + fp_bit_kernels)
    selected_count = _selection_values(count_kernels) or globals()["count_kernels"]
    column_order = [
        f"{fp_kernel}-{count_kernel}"
        for fp_kernel in selected_fp
        for count_kernel in selected_count
        if f"{fp_kernel}-{count_kernel}" in heatmap_matrix.columns
    ]
    column_order.extend(
        column for column in heatmap_matrix.columns if column not in column_order
    )
    heatmap_matrix = heatmap_matrix.reindex(columns=column_order)

    annotations = heatmap_matrix.copy().astype(object)
    for row_label in heatmap_matrix.index:
        for column_label in heatmap_matrix.columns:
            value = heatmap_matrix.loc[row_label, column_label]
            annotations.loc[row_label, column_label] = (
                "" if pd.isna(value) else f"{value:.2f}"
            )

    if not plot_heatmap:
        return heatmap_matrix

    x_tick_labels = [
        str(column).replace("-", ":\n", 1)
        for column in heatmap_matrix.columns
    ]

    fig, ax = plt.subplots(figsize=figsize)
    custom_cmap = sns.color_palette("viridis", as_cmap=True)
    custom_cmap.set_bad(color="lightgray")
    hmap = sns.heatmap(
        heatmap_matrix,
        annot=annotations,
        fmt="",
        cmap=custom_cmap,
        cbar=True,
        vmin=vmin,
        vmax=vmax,
        ax=ax,
        mask=heatmap_matrix.isnull(),
        linewidths=0.5,
        linecolor="white",
        xticklabels=x_tick_labels,
        yticklabels=heatmap_matrix.index.tolist(),
        annot_kws={"fontsize": fontsize},
    )

    cbar = hmap.collections[0].colorbar
    if cbar is not None:
        cbar.set_label(
            "Performance profile AUC",
            fontsize=fontsize,
            fontweight="bold",
        )
        cbar.ax.tick_params(labelsize=fontsize)

    ax.set_xlabel("FP Kernel : Count Kernel", fontsize=fontsize + 2, fontweight="bold")
    ax.set_ylabel("Hybridization method", fontsize=fontsize + 2, fontweight="bold")
    if title is not None:
        ax.set_title(title, fontsize=fontsize + 3, fontweight="bold")
    ax.set_xticklabels(
        x_tick_labels,
        rotation=45,
        ha="right",
        fontsize=fontsize,
    )
    ax.set_yticklabels(ax.get_yticklabels(), rotation=0, fontsize=fontsize)

    plt.tight_layout()

    if save_dir is not None:
        save_dir = ensure_long_path(Path(save_dir))
        os.makedirs(save_dir, exist_ok=True)
        if file_name is None:
            file_name = "profile_auc_heatmap.png"
        fig.savefig(
            ensure_long_path(save_dir / file_name),
            bbox_inches="tight",
            format="png",
            dpi=900 if high_quality else 100,
        )

    if show:
        plt.show()
    else:
        plt.close(fig)

    return heatmap_matrix



if __name__ == "__main__":

    build_master_performance_data(save_path=RESULTS/"master_performance_data"/"Tree_and_GP",
                                   score_metrics=DEFAULT_SCORE_METRICS)

    # COMBINED_RESULTS: pd.DataFrame = RESULTS / "master_performance_data"/ "Tree_and_GP.pkl"
    # result_df = pd.read_pickle(ensure_long_path(COMBINED_RESULTS))
    # prof_results = performance_plot_with_ranks(
    #     df=result_df,
    #     metric="r2",
    #     model=["GPytorchMAP"],
    #     fp_kernels=["TanimotoRBF", "TanimotoMatern32", "TanimotoMatern52", "Tanimoto"],
    #     count_kernels=["RBF", "Matern32", "Matern52"],
    #     # title="GPyTorch MAP R² Performance Profile",
    #     dataset_as_experiment_points=False,
    #     # show_error_band=True,
    #     # show=True,
    #     save_dir= HERE / "result_analysis",
    #     file_name="GpytorchMAP_SK_r2_performance_profile.png",
    # )
    # for metric in ["nll", "cvpp_ama", "ece"]:
    #     prof_results = performance_plot_with_ranks(
    #         df=result_df,
    #         metric=metric,
    #         model=["RF", "XGBR", "NGB", "GPytorchMAP"],
    #         kernel_triples=[
    #             ("Matern32", "Matern32", "averageProduct"),
    #             ("TanimotoMatern32", "Matern32", "averageProduct"),
    #         ],
    #         legend_ncols=2,
    #         fontsize=17,
    #         figsize=(7, 5),
    #         show=True,
    #         show_error_band=False,
    #         save_dir= HERE / "result_analysis",
    #         file_name=f"{metric}_model_performance_profile_curve_comparison.png",
    #     )
    #     metric_name = "AMA" if metric == "cvpp_ama" else metric.upper()
    #     plot_model_profile_comparison(
    #         prof_results=prof_results,
    #         model=["RF", "XGBR", "NGB", "GPytorchMAP"],
    #         kernel_triples=[
    #             ("Matern32", "Matern32", "averageProduct"),
    #             ("TanimotoMatern32", "Matern32", "averageProduct"),
    #         ],
    #         metric=metric,
    #         y_label=f"Profile AUC of {metric_name}",
    #         fontsize=17,
    #         figsize=(5, 5),
    #         save_dir=HERE / "result_analysis",
    #         file_name=f"{metric}_model_profile_comparison.png",
    #     )

    # plot_profile_auc_heatmap(
    #     auc_df=prof_results,
    #     fontsize=13.4,
    #     figsize=(15, 4.1),
    #     model=["GPytorchMAP"],
    #     fp_kernels=["RBF", "Matern32", "Matern52"],
    #     count_kernels=["RBF", "Matern32", "Matern52"],
    #     # title="GPyTorch MAP R² Profile AUC Heatmap",
    #     save_dir=HERE / "result_analysis",
    #     file_name="GPytorchMAP_bitwise_r2_profile_auc_heatmap.png",
    # )


    # plot_model_comparison(
    #     df=result_df,
    #     metric="r2",
    #     model=["RF", "XGBR","NGB","GPytorchMAP", "GpyroHMC"],
    #     kernel_triples=[
    #         ("Matern32", "Matern32", "averageProduct"),
    #         ("TanimotoMatern32", "Matern32", "averageProduct"),
    #         ],
    #     y_label="R²",
    #     fontsize=17,
    #     show=True,
    #     # y_lim=(0,.2),
    #     figsize=(5, 5),
    #     # log_y=True,
    #     save_dir=HERE / "result_analysis",
    #     file_name="r2_model_comparison.png",
    # )
