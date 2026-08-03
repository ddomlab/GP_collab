import json
from pathlib import Path
from typing import List, Optional, Any, Dict
import os 
import re

# visualization imports
# import cmcrameri.cm as cmc
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from matplotlib.collections import PolyCollection
from matplotlib.transforms import Bbox

import numpy as np
import pandas as pd
import seaborn as sns

#statistics imports
# import krippendorff
import pingouin as pg
from scipy.stats import kendalltau
from sklearn.metrics import auc
from scipy.stats import wilcoxon
from sklearn.metrics._scorer import r2_scorer

#docs
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# internal imports
from visualization_setting import set_plot_style, save_img_path, ensure_long_path

set_plot_style()

HERE = Path(__file__).resolve().parent
# DATASETS = HERE.parent.parent / "datasets" / "Validation datasets"
RESULTS = HERE.parent.parent / "results"





def _numeric_feature_value(value: Any) -> float:
    if value is None:
        return np.nan

    try:
        values = np.asarray(value, dtype=float)
    except (TypeError, ValueError):
        return np.nan

    if values.size == 0:
        return np.nan

    finite_values = values[~np.isnan(values)]
    return float(np.mean(finite_values)) if finite_values.size else np.nan


def _feature_records_dataframe(
    feature_records: Any,
    fill_missing_value: Optional[float] = None,
) -> pd.DataFrame:
    if feature_records is None:
        df_feature = pd.DataFrame()
    elif isinstance(feature_records, pd.DataFrame):
        df_feature = feature_records.copy()
    elif isinstance(feature_records, dict):
        df_feature = pd.DataFrame([feature_records])
    else:
        df_feature = pd.DataFrame(
            [record for record in feature_records if isinstance(record, dict)]
        )

    for column in df_feature.columns:
        df_feature[column] = df_feature[column].map(_numeric_feature_value)

    df_feature = df_feature.dropna(axis=1, how="all")
    if fill_missing_value is not None:
        df_feature = df_feature.fillna(fill_missing_value)

    return df_feature


def kendalls_w(df_input, tie_corrected=True, fill_missing_value: Optional[float] = None):
    """
    Kendall's W for agreement across raters (rows) on items (columns).
    If tie_corrected=True, applies the standard tie correction.
    ``df_input`` can be a DataFrame or a list of fold-level feature dictionaries.
    Missing feature values are left as NaN by default; incomplete feature
    columns are excluded before computing W because Kendall's W needs every
    rater to rank the same items.
    """
    df_input = _feature_records_dataframe(df_input, fill_missing_value=fill_missing_value)
    if fill_missing_value is None:
        df_input = df_input.dropna(axis=1, how="any")

    m = len(df_input)          # raters / folds
    n = len(df_input.columns)  # items / features

    if m < 2 or n < 2:
        return {
            "m (Runs)": m,
            "n (Features)": n,
            "Kendall's W": np.nan,
            "Chi-square": np.nan,
            "Degrees of Freedom": max(n - 1, 0),
            "Tie Corrected": tie_corrected
        }

    # Rank each row
    ranks = df_input.rank(axis=1, ascending=False, method="average")

    # Sum of ranks for each feature
    R = ranks.sum(axis=0)
    R_bar = m * (n + 1) / 2.0

    # S term
    S = np.sum((R - R_bar) ** 2)

    if not tie_corrected:
        # Classic denominator
        S_max = (m**2 * (n**3 - n)) / 12.0
        W = S / S_max if S_max != 0 else np.nan
        chi_sq = m * (n - 1) * W
        df_chi = n - 1
        return {
            "m (Runs)": m,
            "n (Features)": n,
            "Kendall's W": W,
            "Chi-square": chi_sq,
            "Degrees of Freedom": df_chi,
            "Tie Corrected": False
        }

    # -------- Tie corrected version --------
    T_total = 0
    for _, row in df_input.iterrows():
        counts = row.value_counts()
        ties = counts[counts > 1]
        if len(ties) > 0:
            T_total += np.sum(ties**3 - ties)

    denom = (m**2 * (n**3 - n)) - (m * T_total)
    W = (12 * S) / denom if denom != 0 else np.nan

    chi_sq = m * (n - 1) * W
    df_chi = n - 1

    return {
        "m (Runs)": m,
        "n (Features)": n,
        "Kendall's W": W,
        "Chi-square": chi_sq,
        "Degrees of Freedom": df_chi,
        "Tie Corrected": True,
        "Total Tie Correction Term (Σ(t^3−t))": T_total
    }



def get_scores(data: Dict, metric: List[str]):
    special = {"rmse", "r2", "mae"}

    return {
        score: (
            f"{round(data[f'{score}_avg'], 2)} ± {round(data[f'{score}_stdev'], 2)}"
            if score.lower() in special
            else f"{int(round(data[score]))}"
        )
        for score in metric
    }


def create_word_table_table(rows_data, folder_path, file_name="results_gp_table.docx"):
    document = Document()

    rows = len(rows_data) + 3
    cols = 9   # <-- 9 columns now (Combination + 2 models × 3 metrics)
    table = document.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"

    # ---------------- HEADER LEVEL 1 ----------------
    h = table.rows[0].cells
    h[0].text = "Combination"
    h[3].text = "GPyro HMC"
    h[6].text = "GPyTorch MAP"

    h[0].merge(h[2])   # Combination spans Count/FP/Mixing
    h[3].merge(h[5])   # GPyro HMC spans RMSE / R² / runtime
    h[6].merge(h[8])   # GPyTorch MAP spans RMSE / R² / runtime

    # ---------------- HEADER LEVEL 2 ----------------
    h2 = table.rows[1].cells
    h2[0].text = "Kernel"
    h2[3].text = "RMSE"
    h2[4].text = "R²"
    h2[5].text = "Run time (sec)"
    h2[6].text = "RMSE"
    h2[7].text = "R²"
    h2[8].text = "Run time (sec)"
    h2[0].merge(h2[1])

    # ---------------- HEADER LEVEL 3 ----------------
    h3 = table.rows[2].cells
    h3[0].text = "Count"
    h3[1].text = "FP"
    h3[2].text = "Mixing method"

    # ---- VERTICAL MERGES ----
    table.rows[1].cells[2].merge(table.rows[2].cells[2])
    for i in range(3, 9):
        table.rows[1].cells[i].merge(table.rows[2].cells[i])

    # ---------------- DATA ROWS ----------------
    r = 3
    for row in rows_data:
        cells = table.rows[r].cells
        for c, v in enumerate(row):
            cells[c].text = str(v)
        r += 1

    # ---------------- STYLING ----------------
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)

    os.makedirs(folder_path, exist_ok=True)
    document.save(folder_path / file_name)
    print(f"Saved Word table → {file_name}")



PLS_Ranks = pd.DataFrame([{
                "Xn": 1,
                "Mw (g/mol)": 2,
                "Concentration (mg/ml)": 3,
                "Temperature SANS/SLS/DLS/SEC (K)": 4,
                "polymer dD": 5,
                "solvent dD": 6,
                "polymer dP": 7,
                "solvent dP": 8,
                "polymer dH": 9,
                "solvent dH": 10,
                "fp": 11,
                "PDI": 12
            }], 
            index=["expert_rank"]
            )

BMS_Ranks = pd.DataFrame([{
                "HOMO_D (eV)": 1,
                "LUMO_A (eV)": 2,
                "Eg_D (eV)": 3,
                "Eg_A (eV)": 4,
                "Ehl_D (eV)": 5,
                "Ehl_A (eV)": 6,
                "HOMO_A (eV)": 7,
                "LUMO_D (eV)": 8,
                "fp_Donor": 9,
                "fp_Acceptor": 10,
                "D:A ratio (m/m)": 11,
                "temperature of thermal annealing": 12,
                "solvent additive conc. (% v/v)": 13,
                "HTL energy level (eV)": 14,
                "ETL energy level (eV)": 15
            }],
            index=["expert_rank"]
            )

PAPER = {
    "Robust Learning from Literature Data_Model Generalizability and Uncertainty for Predicting Conjugated Polymer Solution Conformation": {
        "target": ["target_log Rg (nm)"],
        "expert_impt": PLS_Ranks,
        "n_datapoints":[256]
    },
    "Beyond molecular structure_ critically assessing machine learning for designing organic photovoltaic materials and devices": {
        "target": ["target_calculated PCE (%)"],
        "expert_impt": BMS_Ranks,
        "n_datapoints":[558]
    },
    "Machine Learning for Polymer Design to Enhance Pervaporation-Based Organic Recovery": {
        "target": ["target_log (Separation factor)", "target_log (Total flux)"],
        "expert_impt": None,
        "n_datapoints": [2311,2283]
    },
    "Machine Learning-Enabled Prediction and High-Throughput Screening of Polymer Membranes for Pervaporation Separation": {
        "target": ["target_log (Separation factor)", "target_log (Total flux)"],
        "expert_impt": None,
        "n_datapoints": [681,681]
    },
    "Understanding and Designing a High-Performance Ultrafiltration Membrane Using Machine Learning": {
        "target": [
            "target_flux decline ratio (%)",
            "target_flux recovery ratio (%)",
            "target_irreversible fouling ratio(%)",
            "target_organic compound removal (%)",
            "target_reversible fouling ratio (%)",
            r"target_water permeability (LMH\bar)"
        ],
        "expert_impt": None,
        "n_datapoints": [318,318,318,318,318,318]
    },
    "Miniaturization of Popular Reactions from the Medicinal Chemists Toolbox for Ultrahigh_Throughput Experimentation": {
        "target": ["target_Approx Conv (%)"],
        "expert_impt": None,
        "n_datapoints": [768]
    },
}



fp_sk_kernels = ["TanimotoMatern32", "TanimotoMatern52", "TanimotoRBF", "Tanimoto"]
fp_bit_kernels = ["Matern32", "Matern52", "RBF"]
count_kernels    = ["Matern32", "Matern52", "RBF"]
mixing_methods  = ["sum", "product", "averageProduct"]
HYBRIDIZATION_METHOD_COLORS = {
    "sum": "#2A9C9D",
    "product": "#C77B7B",
    "averageProduct": "#EDC525",
}
mixing_labels   = {"sum": "Sum", "product": "Product", "averageProduct": "Av(count)×FP"}
TREE_MODELS = {"RF", "XGBR", "NGB"}
GNN_MODELS = {"GNN", "GCN", "GAT", "GIN", "MPNN", "DMPNN"}
MODEL_TYPE_COLORS = {
    "tree": "#7093B9",
    "gp": "#E45756",
    "mgk": "#9C1E1E",
    "gnn": "#72B7B2",
}

INDIVIDUAL_MODEL_COLORS = {
    "RF": "#174C85",
    "XGBR": "#5482B3",
    "NGB": "#94AFCC",
    "GPytorchMAP (Bitwise)": "#D88F8F",
    "GPytorchMAP (SK)": "#E03B3B",
    "GpyroHMC (SK)": "#7E1919",
    "MGK": "#CB6D01",
}

MODELS= ["RF", "XGBR", "NGB", "GpyroHMC", "GPytorchMAP", "MGK"]
GPU_SCORE_MODELS = {"MGK"}
DEFAULT_SCORE_METRICS = ["rmse", "r2", "mae", "cvpp_ama", "nll", "ece", "Cv", "RUSC"]
BASE_MASTER_RESULT_COLUMNS = [
    "paper",
    "target",
    "model",
    "fp kernel",
    "count kernel",
    "mixing method",
]
TIME_COLUMNS = [
    "Running time (GPU)",
    "Running time (CPU)",
]
OOF_SCORE_COLUMNS = [
    "OOF_R2",
]
FEATURE_DATA_COLUMNS = [
    "lengthscale",
    "feature_importance_MDI",
    "feature_importance_SHAP",
]
FEATURE_STABILITY_COLUMNS = [
    "lengthscale_kendalls_w",
    "feature_importance_MDI_kendalls_w",
    "feature_importance_SHAP_kendalls_w",
]

label_conversion_source = {
    "r2": "R²",
    "OOF_R2": "R²",
    "rmse": "RMSE",
    "mae": "MAE",
    "nll": "NLL",
    "cvpp_ama": "AMA",
    "ece": "ECE"
}

# def plot_barplot(model_specs, save_dir: Path, figsize=(6, 4)):
#     data = build_barplot_data(model_specs)

#     labels = [spec["label"] for spec in model_specs]
#     means = [np.mean(data[label]) for label in labels]
#     stds = [np.std(data[label]) for label in labels]

#     fig, ax = plt.subplots(figsize=figsize)

#     x = np.arange(len(labels))
#     bars = ax.bar(
#         x,
#         means,
#         yerr=stds,
#         capsize=5,
#         color=["#AA5A6E", "#AA5A6E","#AA5A6E","#AA5A6E", "#6B9DB4"],
#         edgecolor="white",
#         linewidth=0.8,
#         error_kw={"elinewidth": 1.5, "ecolor": "black", "capthick": 1.5},
#     )

#     for bar, mean, std in zip(bars, means, stds):
#         ax.text(
#             bar.get_x() + bar.get_width() / 2,
#             bar.get_height() + std + 0.005,
#             f"{mean:.2f}",
#             ha="center",
#             va="bottom",
#             fontsize=13,
#         )

#     ax.set_xticks(x)
#     ax.tick_params(axis="both", labelsize=14)
#     ax.set_xticklabels(labels, fontsize=14, rotation=45)
#     ax.set_ylabel("R²", fontsize=16, fontweight="bold")
#     ax.set_xlabel("Model", fontsize=16, fontweight="bold")
#     ax.set_ylim(0, 1.0)
#     ax.set_yticks(np.arange(0, 1.1, 0.1))

#     plt.tight_layout()
#     save_img_path(save_dir, "Tree_based_model_result_across_datasets.png")
#     plt.close(fig)
#     print(f"Saved barplot at {save_dir}")





def _is_tree_model(model: str) -> bool:
    return model in TREE_MODELS


def _uses_gpu_scores(model: str) -> bool:
    return model in GPU_SCORE_MODELS


def _kernel_configs_for_model(model: str) -> List[tuple[str, str, str]]:
    return [
        (fp_k, count_k, mix_method)
        for fp_k in fp_sk_kernels + fp_bit_kernels
        for count_k in count_kernels
        for mix_method in mixing_methods
    ]


def _score_file_templates(
    model: str,
    fp_k: Optional[str] = None,
    count_k: Optional[str] = None,
    mix_method: Optional[str] = None,
    use_gpu: bool = False,
) -> List[str]:
    if _is_tree_model(model):
        if use_gpu:
            return []
        return [f"(ECFP3_count_512-COUNT)_{model}_hypOFF_Standard_Standard_scores"]

    if model == "MGK":
        if (
            not use_gpu
            or count_k not in count_kernels
            or mix_method not in mixing_methods
        ):
            return []
        return [
            (
                f"(MG-COUNT)_(MGK_Graph-{count_k}_{mix_method})"
                f"{suffix}_hypOFF_Standard_Standard_GPU_scores"
            )
            for suffix in ["", "_mean"]
        ]

    device_suffix = "_GPU" if use_gpu else ""
    return [
        (
            f"(ECFP3_count_512-COUNT)_"
            f"({model}_{fp_k}-{count_k}_{mix_method})"
            f"{suffix}_hypOFF_Standard_Standard{device_suffix}_scores"
        )
        for suffix in ["", "_mean"]
    ]


def _find_score_path(
    paper_loc: Path,
    model: str,
    fp_k: Optional[str] = None,
    count_k: Optional[str] = None,
    mix_method: Optional[str] = None,
    use_gpu: bool = False,
) -> Optional[Path]:
    for tmpl in _score_file_templates(model, fp_k, count_k, mix_method, use_gpu):
        path = ensure_long_path(paper_loc / f"{tmpl}.json")
        if path.exists():
            return path
    return None


def _read_json(path: Optional[Path]) -> Optional[Dict[str, Any]]:
    if path is None:
        return None

    with open(path, "r") as f:
        return json.load(f)


def _prediction_path_from_score_path(score_path: Optional[Path]) -> Optional[Path]:
    if score_path is None:
        return None

    score_name = score_path.name
    if not score_name.endswith("_scores.json"):
        return None

    prediction_stem = score_name.removesuffix("_scores.json") + "_predictions"
    for suffix in [".csv", ".json"]:
        prediction_path = ensure_long_path(score_path.with_name(f"{prediction_stem}{suffix}"))
        if prediction_path.exists():
            return prediction_path

    return None


def _prediction_target_column(
    columns: List[str],
    target: Optional[str] = None,
) -> Optional[str]:
    target_names = []
    if target is not None:
        target_names.extend([target, target.removeprefix("target_")])

    for target_name in target_names:
        if target_name in columns:
            return target_name

    non_prediction_columns = [
        column
        for column in columns
        if not re.match(r"^seed_.+_y_(pred|std)$", str(column))
    ]
    return non_prediction_columns[0] if len(non_prediction_columns) == 1 else None


def _pooled_oof_r2_score(
    prediction_path: Optional[Path],
    target: Optional[str] = None,
) -> Optional[float]:
    if prediction_path is None:
        return None

    if prediction_path.suffix.lower() != ".csv":
        return None

    prediction_df = pd.read_csv(prediction_path)
    target_col = _prediction_target_column(prediction_df.columns.tolist(), target)
    if target_col is None:
        return None

    pred_cols = [
        column
        for column in prediction_df.columns
        if re.match(r"^seed_.+_y_pred$", str(column))
    ]
    if not pred_cols:
        return None

    y_true = pd.to_numeric(prediction_df[target_col], errors="coerce").to_numpy()
    pooled_true = []
    pooled_pred = []
    for pred_col in pred_cols:
        y_pred = pd.to_numeric(prediction_df[pred_col], errors="coerce").to_numpy()
        valid_mask = np.isfinite(y_true) & np.isfinite(y_pred)
        pooled_true.extend(y_true[valid_mask])
        pooled_pred.extend(y_pred[valid_mask])

    if len(pooled_true) < 2:
        return None

    return float(
        r2_scorer._score_func(
            np.asarray(pooled_true),
            np.asarray(pooled_pred),
            **r2_scorer._kwargs,
        )
    )


def _metric_score(
    data: Optional[Dict[str, Any]],
    metric: str,
    suffix: str,
) -> Optional[Any]:
    if data is None:
        return None

    return data.get(f"{metric}_{suffix}")


def _seed_sort_key(value: Any) -> tuple:
    return (0, int(value)) if str(value).isdigit() else (1, str(value))


def _seed_items(data: Optional[Dict[str, Any]]) -> List[tuple[str, Dict[str, Any]]]:
    if data is None:
        return []

    return [
        (key, data[key])
        for key in sorted(
            [key for key, value in data.items() if isinstance(value, dict)],
            key=_seed_sort_key,
        )
    ]


def _seed_fold_scores(
    data: Optional[Dict[str, Any]],
    metric: str,
) -> Optional[List[Any]]:
    scores = []

    for _, seed_data in _seed_items(data):
        values = seed_data.get(f"test_{metric}")
        if isinstance(values, list):
            scores.extend(values)

    return scores or None


def _load_score_files(
    paper_loc: Path,
    model: str,
    fp_k: Optional[str] = None,
    count_k: Optional[str] = None,
    mix_method: Optional[str] = None,
) -> tuple[
    Optional[Dict[str, Any]],
    Optional[Dict[str, Any]],
    Optional[Path],
    Optional[Path],
]:
    cpu_path = _find_score_path(paper_loc, model, fp_k, count_k, mix_method)
    cpu_data = _read_json(cpu_path)
    gpu_path = None
    gpu_data = None

    if not _is_tree_model(model):
        gpu_path = _find_score_path(
            paper_loc,
            model,
            fp_k,
            count_k,
            mix_method,
            use_gpu=True,
        )
        gpu_data = _read_json(
            gpu_path
        )

    return cpu_data, gpu_data, cpu_path, gpu_path


def _metric_result_columns(score_metrics: List[str]) -> List[str]:
    columns = []
    for metric in score_metrics:
        columns.extend([
            f"{metric}_avg",
            f"{metric}_stdev",
            f"{metric}_seed_fold_scores",
        ])
    return columns


def _master_result_columns(score_metrics: List[str]) -> List[str]:
    return (
        BASE_MASTER_RESULT_COLUMNS
        + FEATURE_DATA_COLUMNS
        + FEATURE_STABILITY_COLUMNS
        + _metric_result_columns(score_metrics)
        + OOF_SCORE_COLUMNS
        + TIME_COLUMNS
    )


def _seed_fold_feature_records(
    data: Optional[Dict[str, Any]],
    key: str,
) -> Optional[List[Dict[str, Any]]]:
    records = []

    if data is None:
        return None

    top_level_values = data.get(key)
    if isinstance(top_level_values, list):
        records.extend(
            value for value in top_level_values if isinstance(value, dict)
        )
    elif isinstance(top_level_values, dict):
        records.append(top_level_values)

    for _, seed_data in _seed_items(data):
        values = seed_data.get(key)
        if isinstance(values, list):
            records.extend(value for value in values if isinstance(value, dict))
        elif isinstance(values, dict):
            records.append(values)

    return records or None


def _feature_kendalls_w(records: Optional[List[Dict[str, Any]]]) -> Optional[float]:
    if not records:
        return None

    value = kendalls_w(records)["Kendall's W"]
    return None if pd.isna(value) else value


def _feature_row_values(
    cpu_data: Optional[Dict[str, Any]],
    gpu_data: Optional[Dict[str, Any]],
) -> Dict[str, Any]:
    source_data = cpu_data if cpu_data is not None else gpu_data
    lengthscale = _seed_fold_feature_records(source_data, "test_lengthscale")
    mdi = _seed_fold_feature_records(source_data, "test_feature_importance_MDI")
    shap = _seed_fold_feature_records(source_data, "test_feature_importance_SHAP")

    return {
        "lengthscale": lengthscale,
        "feature_importance_MDI": mdi,
        "feature_importance_SHAP": shap,
        "lengthscale_kendalls_w": _feature_kendalls_w(lengthscale),
        "feature_importance_MDI_kendalls_w": _feature_kendalls_w(mdi),
        "feature_importance_SHAP_kendalls_w": _feature_kendalls_w(shap),
    }


def _score_row_values(
    cpu_data: Optional[Dict[str, Any]],
    gpu_data: Optional[Dict[str, Any]],
    score_metrics: List[str],
    prefer_gpu_scores: bool = False,
    target: Optional[str] = None,
    cpu_score_path: Optional[Path] = None,
    gpu_score_path: Optional[Path] = None,
) -> Dict[str, Any]:
    scores = {}
    score_data = gpu_data if prefer_gpu_scores else cpu_data
    score_path = gpu_score_path if prefer_gpu_scores else cpu_score_path
    for metric in score_metrics:
        scores.update({
            f"{metric}_avg": _metric_score(score_data, metric, "avg"),
            f"{metric}_stdev": _metric_score(score_data, metric, "stdev"),
            f"{metric}_seed_fold_scores": _seed_fold_scores(score_data, metric),
        })

    scores.update({
        "OOF_R2": _pooled_oof_r2_score(
            _prediction_path_from_score_path(score_path),
            target,
        ),
        "Running time (GPU)": None if gpu_data is None else gpu_data.get("run_time_sec"),
        "Running time (CPU)": None if cpu_data is None else cpu_data.get("run_time_sec"),
    })
    return scores


def load_score(
    paper_loc: Path,
    model: str,
    fp_k: Optional[str] = None,
    count_k: Optional[str] = None,
    mix_method: Optional[str] = None,
    score_metrics: Optional[List[str]] = None,
) -> Dict[str, Any]:
    """Load score values for one master-results row."""
    score_metrics = DEFAULT_SCORE_METRICS if score_metrics is None else score_metrics
    cpu_data, gpu_data, cpu_path, gpu_path = _load_score_files(
        paper_loc,
        model,
        fp_k,
        count_k,
        mix_method,
    )
    return _score_row_values(
        cpu_data,
        gpu_data,
        score_metrics,
        prefer_gpu_scores=_uses_gpu_scores(model),
        cpu_score_path=cpu_path,
        gpu_score_path=gpu_path,
    )


def _save_master_performance_data(df: pd.DataFrame, save_path: Path) -> None:
    save_path = Path(save_path)
    os.makedirs(ensure_long_path(save_path.parent), exist_ok=True)

    df.to_pickle(ensure_long_path(Path(f"{save_path}.pkl")))

    output = df.copy()
    for column in output.columns:
        if column.endswith("_seed_fold_scores") or column in FEATURE_DATA_COLUMNS:
            output[column] = output[column].apply(
                lambda value: json.dumps(value) if value is not None else None
            )
    output.to_csv(ensure_long_path(Path(f"{save_path}.csv")), index=False)


def build_master_performance_data(
    save_path: Optional[Path] = RESULTS / "master_performance_data",
    score_metrics: Optional[List[str]] = None,
) -> pd.DataFrame:
    score_metrics = (
        DEFAULT_SCORE_METRICS
        if score_metrics is None
        else list(dict.fromkeys(score_metrics))
    )
    rows = []

    for paper_name, paper_info in PAPER.items():
        for target in paper_info["target"]:
            paper_loc = RESULTS / paper_name / target

            for model in MODELS:
                if _is_tree_model(model):
                    cpu_data, gpu_data, cpu_path, gpu_path = _load_score_files(
                        paper_loc=paper_loc,
                        model=model,
                    )
                    score_info = _score_row_values(
                        cpu_data,
                        gpu_data,
                        score_metrics,
                        target=target,
                        cpu_score_path=cpu_path,
                        gpu_score_path=gpu_path,
                    )
                    feature_info = _feature_row_values(cpu_data, gpu_data)
                    rows.append({
                        "paper": paper_name,
                        "target": target,
                        "model": model,
                        "fp kernel": None,
                        "count kernel": None,
                        "mixing method": None,
                        **feature_info,
                        **score_info,
                    })
                    continue

                if model == "MGK":
                    for count_k in count_kernels:
                        for mix in mixing_methods:
                            cpu_data, gpu_data, cpu_path, gpu_path = _load_score_files(
                                paper_loc=paper_loc,
                                model=model,
                                count_k=count_k,
                                mix_method=mix,
                            )
                            score_info = _score_row_values(
                                cpu_data,
                                gpu_data,
                                score_metrics,
                                prefer_gpu_scores=True,
                                target=target,
                                cpu_score_path=cpu_path,
                                gpu_score_path=gpu_path,
                            )
                            feature_info = _feature_row_values(cpu_data, gpu_data)
                            rows.append({
                                "paper": paper_name,
                                "target": target,
                                "model": model,
                                "fp kernel": "Graph",
                                "count kernel": count_k,
                                "mixing method": mix,
                                **feature_info,
                                **score_info,
                            })
                    continue

                for fp_k, count_k, mix in _kernel_configs_for_model(model):
                    cpu_data, gpu_data, cpu_path, gpu_path = _load_score_files(
                        paper_loc=paper_loc,
                        model=model,
                        fp_k=fp_k,
                        count_k=count_k,
                        mix_method=mix,
                    )
                    score_info = _score_row_values(
                        cpu_data,
                        gpu_data,
                        score_metrics,
                        prefer_gpu_scores=_uses_gpu_scores(model),
                        target=target,
                        cpu_score_path=cpu_path,
                        gpu_score_path=gpu_path,
                    )
                    feature_info = _feature_row_values(cpu_data, gpu_data)
                    rows.append({
                        "paper": paper_name,
                        "target": target,
                        "model": model,
                        "fp kernel": fp_k,
                        "count kernel": count_k,
                        "mixing method": mix,
                        **feature_info,
                        **score_info,
                    })

    df = pd.DataFrame(rows, columns=_master_result_columns(score_metrics), dtype=object)
    if save_path is not None:
        _save_master_performance_data(df, save_path)

    return df



def add_legend_box(
    indices, legend, color="grey", linestyle="dashed", alpha=1.0, linewidth=1.5, pad=5
):
    # Utilitary function to add a box in the legend of a figure.
    fig = legend.figure
    renderer = fig.canvas.get_renderer()

    texts = legend.get_texts()
    handles = legend.legend_handles

    text_boxes = [texts[i].get_window_extent(renderer) for i in indices]
    handle_boxes = [handles[i].get_window_extent(renderer) for i in indices]

    all_boxes = text_boxes + handle_boxes

    if not all_boxes:
        return

    full_box = Bbox.union(all_boxes).padded(pad)

    trans = fig.transFigure.inverted()
    full_box_fig = trans.transform_bbox(full_box)

    rect = patches.Rectangle(
        (full_box_fig.x0, full_box_fig.y0),
        full_box_fig.width,
        full_box_fig.height,
        transform=fig.transFigure,
        facecolor="none",
        edgecolor=color,
        linestyle=linestyle,
        linewidth=linewidth,
        alpha=alpha,
        zorder=5,
    )
    fig.patches.append(rect)


def _coerce_score_list(value: Any) -> Optional[List[Any]]:
    if isinstance(value, list):
        return value
    if isinstance(value, np.ndarray):
        return value.tolist()
    if isinstance(value, str):
        try:
            parsed = json.loads(value)
        except json.JSONDecodeError:
            return None
        return parsed if isinstance(parsed, list) else None
    return None


def _selection_values(values: Any) -> Optional[List[Any]]:
    if values is None:
        return None
    if isinstance(values, str):
        return [values]
    return list(values)


def _filter_selection(
    df: pd.DataFrame,
    column: str,
    values: Any,
    keep_missing: bool = False,
) -> pd.DataFrame:
    selected = _selection_values(values)
    if selected is None:
        return df

    selected_lower = {str(value).lower() for value in selected}
    column_values = df[column].astype(str).str.lower()
    mask = column_values.isin(selected_lower)
    if keep_missing:
        mask = mask | df[column].isna()
    return df[mask].copy()


def _kernel_label(row: pd.Series, include_model: bool = False) -> str:
    model = str(row["model"])
    if pd.isna(row["fp kernel"]) and pd.isna(row["count kernel"]):
        return model

    label = f"{row['fp kernel']}-{row['count kernel']}_{row['mixing method']}"
    return f"{model}: {label}" if include_model else label


def _metric_higher_is_better(metric: str) -> bool:
    return (
        str(metric).strip().lower() in {"r2", "oof_r2", "rusc", "cvpp_ama"}
        or metric in FEATURE_STABILITY_COLUMNS
        or _is_feature_stability_metric(metric)
    )


def run_topsis(
    df: pd.DataFrame,
    criteria_weights: Any,
    criteria_types: Any,
) -> pd.Series:
    criteria_weights = np.asarray(criteria_weights, dtype=float)
    criteria_types = np.asarray(criteria_types, dtype=int)

    norm_df = df / (np.sqrt((df**2).sum(axis=0)) + 1e-12)
    weighted_df = norm_df * criteria_weights

    ideal_pos = np.where(
        criteria_types == 1,
        weighted_df.max(axis=0),
        weighted_df.min(axis=0),
    )
    ideal_neg = np.where(
        criteria_types == 1,
        weighted_df.min(axis=0),
        weighted_df.max(axis=0),
    )

    dist_pos = np.sqrt(((weighted_df - ideal_pos) ** 2).sum(axis=1))
    dist_neg = np.sqrt(((weighted_df - ideal_neg) ** 2).sum(axis=1))

    closeness = dist_neg / (dist_pos + dist_neg + 1e-12)
    return pd.Series(closeness, index=df.index, name="TOPSIS_Score")


def _mixing_method_label(value: Any) -> str:
    clean_labels = {"averageProduct": "Av(count)xFP"}
    value = str(value)
    return clean_labels.get(value, mixing_labels.get(value, value))


def _expand_master_scores_for_profile(
    df: pd.DataFrame,
    metric: str,
    model: Any,
    fp_kernels: Any = None,
    count_kernels: Any = None,
    mixing_methods: Any = None,
    tree_feature_importance: str = "MDI",
) -> pd.DataFrame:
    rows = []
    selected_models = _selection_values(model)
    model_df = _filter_selection(df, "model", selected_models)
    model_df = _filter_selection(model_df, "fp kernel", fp_kernels, keep_missing=True)
    model_df = _filter_selection(model_df, "count kernel", count_kernels, keep_missing=True)
    model_df = _filter_selection(model_df, "mixing method", mixing_methods, keep_missing=True)
    include_model = selected_models is None or len(selected_models) > 1
    feature_stability_metric = _is_feature_stability_metric(metric)

    if feature_stability_metric:
        tree_stability_col = _tree_feature_stability_column(tree_feature_importance)
        required_columns = {"lengthscale_kendalls_w"}
        if model_df["model"].apply(_is_tree_model).any():
            required_columns.add(tree_stability_col)
        missing_columns = required_columns.difference(df.columns)
        if missing_columns:
            raise ValueError(
                "df is missing required columns: "
                + ", ".join(sorted(missing_columns))
            )
        use_seed_fold_scores = False
        score_col = None
    else:
        score_col = f"{metric}_seed_fold_scores"
        use_seed_fold_scores = True
        if score_col not in df.columns and metric in df.columns:
            score_col = metric
            use_seed_fold_scores = False
        if score_col not in df.columns:
            raise ValueError(f"Missing required column: {score_col}")

    for _, row in model_df.iterrows():
        feature_stability_source = None
        tree_feature_source = None
        if feature_stability_metric:
            score_col = _feature_stability_column_for_model(
                row["model"],
                tree_feature_importance,
            )
            scores = [row[score_col]]
            feature_stability_source = _feature_stability_source_label(score_col)
            tree_feature_source = _feature_stability_source_label(
                _tree_feature_stability_column(tree_feature_importance)
            )
        elif use_seed_fold_scores:
            scores = _coerce_score_list(row[score_col])
            if scores is None:
                continue
        else:
            scores = [row[score_col]]

        dataset = row["paper"]
        target = row["target"]
        kernel = _kernel_label(row, include_model=include_model)
        for repeat_idx, score in enumerate(scores):
            try:
                score = float(score)
            except (TypeError, ValueError):
                continue
            if np.isnan(score):
                continue
            row_data = {
                "dataset": dataset,
                "target": target,
                "repeat": repeat_idx,
                "kernel": kernel,
                "model": row["model"],
                "fp kernel": row["fp kernel"],
                "count kernel": row["count kernel"],
                "mixing method": row["mixing method"],
                metric: score,
            }
            if feature_stability_metric:
                row_data["feature stability source"] = feature_stability_source
                row_data["tree feature importance"] = tree_feature_source
            rows.append(row_data)

    return pd.DataFrame(rows)



def _model_config_label(row: pd.Series, include_kernel_config: bool = True) -> str:
    model = str(row["model"])
    if not include_kernel_config:
        return model
    if model == "MGK":
        return model
    if pd.isna(row["fp kernel"]) and pd.isna(row["count kernel"]):
        return model

    # mix = mixing_labels.get(str(row["mixing method"]))
    suffix = "SK" if "tanimoto" in row["fp kernel"].lower() else "Bitwise"
    return f"{model} ({suffix})"


def _model_config_sort_key(row: pd.Series, model_order: Optional[Dict[str, int]] = None) -> tuple:
    model = str(row["model"])
    model_rank = len(model_order) if model_order is not None else 0
    if model_order is not None:
        model_rank = model_order.get(model.lower(), model_rank)

    fp_kernel = row["fp kernel"]
    if _is_tree_model(model):
        family_rank = 0
    elif model == "MGK":
        family_rank = 3
    elif pd.notna(fp_kernel) and "tanimoto" not in str(fp_kernel).lower():
        family_rank = 1
    elif pd.notna(fp_kernel) and "tanimoto" in str(fp_kernel).lower():
        family_rank = 2
    else:
        family_rank = 4

    return (family_rank, model_rank, model)


def _is_tree_model(model: Any) -> bool:
    return str(model).upper() in TREE_MODELS


def _model_group_sort_rank(model: Any) -> int:
    model_name = str(model)
    model_upper = model_name.upper()
    if _is_tree_model(model_name):
        return 0
    if "GP" in model_upper and model_name != "MGK":
        return 1
    if model_name == "MGK":
        return 2
    if model_name in GNN_MODELS or any(token in model_upper for token in GNN_MODELS):
        return 3
    return 4


def _model_order_sort_key(model: Any, model_order: Optional[Dict[str, int]] = None) -> tuple:
    model_name = str(model)
    model_rank = len(model_order) if model_order is not None else 0
    if model_order is not None:
        model_rank = model_order.get(model_name.lower(), model_rank)
    return (_model_group_sort_rank(model_name), model_rank, model_name)


def _model_type_color(model: Any) -> str:
    model_name = str(model)
    model_upper = model_name.upper()
    if _is_tree_model(model_name):
        return MODEL_TYPE_COLORS["tree"]
    if model_name == "MGK":
        return MODEL_TYPE_COLORS["mgk"]
    if model_name in GNN_MODELS or any(token in model_upper for token in GNN_MODELS):
        return MODEL_TYPE_COLORS["gnn"]
    if "GP" in model_upper:
        return MODEL_TYPE_COLORS["gp"]
    return "#808080"


def _tree_feature_stability_column(tree_feature_importance: str) -> str:
    aliases = {
        "mdi": "feature_importance_MDI_kendalls_w",
        "model": "feature_importance_MDI_kendalls_w",
        "model_importance": "feature_importance_MDI_kendalls_w",
        "feature_importance_mdi": "feature_importance_MDI_kendalls_w",
        "feature_importance_mdi_kendalls_w": "feature_importance_MDI_kendalls_w",
        "shap": "feature_importance_SHAP_kendalls_w",
        "feature_importance_shap": "feature_importance_SHAP_kendalls_w",
        "feature_importance_shap_kendalls_w": "feature_importance_SHAP_kendalls_w",
    }
    key = str(tree_feature_importance).strip().lower()
    if key not in aliases:
        raise ValueError("tree_feature_importance must be either 'MDI' or 'SHAP'.")
    return aliases[key]


def _feature_stability_source_label(column: str) -> str:
    labels = {
        "lengthscale_kendalls_w": "Lengthscale",
        "feature_importance_MDI_kendalls_w": "MDI",
        "feature_importance_SHAP_kendalls_w": "SHAP",
    }
    return labels.get(column, column)


def _is_feature_stability_metric(metric: Any) -> bool:
    metric_key = re.sub(r"[\s\-]+", "_", str(metric).strip().lower())
    return metric_key in {
        "feature_stability",
        "feature_importance_stability",
        "feature_importance_kendalls_w",
        "feature_importance_stability_kendalls_w",
        "tree_feature_importance_stability",
    }


def _feature_stability_column_for_model(
    model: Any,
    tree_feature_importance: str,
) -> str:
    if _is_tree_model(model):
        return _tree_feature_stability_column(tree_feature_importance)
    return "lengthscale_kendalls_w"


def _filter_metric_selection(
    df: pd.DataFrame,
    metric: Any,
    column: str = "metric",
) -> pd.DataFrame:
    selected = _selection_values(metric)
    if selected is None or column not in df.columns:
        return df

    selected_lower = {str(value).strip().lower() for value in selected}
    column_values = df[column].astype(str).str.strip().str.lower()
    mask = column_values.isin(selected_lower)
    if any(_is_feature_stability_metric(value) for value in selected):
        mask = mask | df[column].map(_is_feature_stability_metric)
    return df[mask].copy()


def _kernel_triple_values(kernel_triples: Any) -> Optional[List[tuple]]:
    if kernel_triples is None:
        return None
    if (
        isinstance(kernel_triples, tuple)
        and len(kernel_triples) == 3
        and not any(isinstance(value, (list, tuple)) for value in kernel_triples)
    ):
        return [kernel_triples]

    triples = list(kernel_triples)
    for triple in triples:
        if not isinstance(triple, (list, tuple)) or len(triple) != 3:
            raise ValueError(
                "kernel_triples must be a triple or a list of "
                "(fp_kernel, count_kernel, mixing_method) triples."
            )
    return triples


def _filter_kernel_triples(
    df: pd.DataFrame,
    kernel_triples: Any,
    keep_missing: bool = True,
) -> pd.DataFrame:
    triples = _kernel_triple_values(kernel_triples)
    if triples is None:
        return df

    selected = {
        (str(fp).lower(), str(count).lower(), str(mix).lower())
        for fp, count, mix in triples
    }
    row_triples = zip(
        df["fp kernel"].astype(str).str.lower(),
        df["count kernel"].astype(str).str.lower(),
        df["mixing method"].astype(str).str.lower(),
    )
    mask = pd.Series(
        [row_triple in selected for row_triple in row_triples],
        index=df.index,
    )
    if keep_missing:
        mask = mask | (
            df["fp kernel"].isna()
            & df["count kernel"].isna()
            & df["mixing method"].isna()
        )
    return df[mask].copy()


def plot_model_comparison(
    df: pd.DataFrame,
    metric: str = "r2",
    model: Any = "GPytorchMAP",
    kernel_triples: Any = None,
    include_kernel_config: bool = True,
    dataset_as_experiment_points: bool = False,
    figsize: tuple = (9, 5),
    fontsize: int = 12,
    title: Optional[str] = None,
    x_label: str = "Model",
    y_label: Optional[str] = None,
    x_tick_rotation: int = 35,
    y_lim: Optional[tuple] = None,
    log_y: bool = False,
    show: bool = True,
    high_quality: bool = True,
    save_dir: Optional[Path] = HERE / "result_analysis",
    file_name: Optional[str] = None,
) -> pd.DataFrame:
    """
    Make a box plot of a metric score for selected models/kernels.

    Score metrics use ``<metric>_seed_fold_scores``. Scalar columns such as
    ``Running time (CPU)`` are plotted directly across dataset/model rows.
    For GP models, pass ``kernel_triples`` as one triple or a list of
    ``(fp_kernel, count_kernel, mixing_method)`` triples.
    Set ``log_y=True`` to show the y-axis on a log scale.
    """
    df = _filter_kernel_triples(df, kernel_triples, keep_missing=True)
    is_scalar_metric = metric in df.columns and f"{metric}_seed_fold_scores" not in df.columns
    plot_df = _expand_master_scores_for_profile(
        df,
        metric=metric,
        model=model,
    )
    if plot_df.empty:
        raise ValueError(
            f"No {metric} values found for model={model} "
            f"with kernel_triples={kernel_triples}."
        )

    x_col = "model configuration"
    plot_df[x_col] = plot_df.apply(
        lambda row: _model_config_label(row, include_kernel_config),
        axis=1,
    )

    if dataset_as_experiment_points:
        group_cols = [
            "dataset",
            "target",
            x_col,
            "model",
            "fp kernel",
            "count kernel",
            "mixing method",
        ]
        plot_df = (
            plot_df.groupby(group_cols, dropna=False, as_index=False)[metric]
            .mean()
            .copy()
        )

    selected_models = _selection_values(model) or MODELS
    model_order = {str(model_name).lower(): idx for idx, model_name in enumerate(selected_models)}
    config_order = plot_df[
        [x_col, "model", "fp kernel", "count kernel", "mixing method"]
    ].drop_duplicates(subset=[x_col])
    config_order["sort_key"] = config_order.apply(
        lambda row: _model_config_sort_key(row, model_order),
        axis=1,
    )
    order = config_order.sort_values("sort_key", kind="mergesort")[x_col].tolist()
    violin_palette = {
        label: _model_type_color(plot_df.loc[plot_df[x_col] == label, "model"].iloc[0])
        for label in order
    }
    fig, ax = plt.subplots(figsize=figsize)
    sns.violinplot(
        data=plot_df,
        x=x_col,
        y=metric,
        hue=x_col,
        order=order,
        hue_order=order,
        ax=ax,
        width=0.65,
        palette=violin_palette,
        inner_kws=dict(box_width=10, whis_width=2, color=".65",solid_capstyle="round"),
        cut=0,
        linewidth=1,
        dodge=False,
        legend=False,
    )
    positions = list(range(len(order)))
    for collection in ax.collections:
        if not isinstance(collection, PolyCollection):
            continue
        for path in collection.get_paths():
            vertices = path.vertices
            if len(vertices) == 0:
                continue
            center = min(positions, key=lambda pos: abs(pos - np.median(vertices[:, 0])))
            vertices[:, 0] = np.maximum(vertices[:, 0], center)
    for line in ax.lines:
        x_data = np.asarray(line.get_xdata(), dtype=float)
        if len(x_data) == 0:
            continue
        center = min(positions, key=lambda pos: abs(pos - np.mean(x_data)))
        line.set_xdata(np.maximum(x_data, center))

    # if show_points:
    #     sns.stripplot(
    #         data=plot_df,
    #         x=x_col,
    #         y=metric,
    #         order=order,
    #         ax=ax,
    #         color="black",
    #         size=3,
    #         alpha=point_alpha,
    #         jitter=0.22,
    #     )
    #     for collection in ax.collections:
    #         if isinstance(collection, PolyCollection) or not hasattr(collection, "get_offsets"):
    #             continue
    #         offsets = collection.get_offsets()
    #         if len(offsets) == 0:
    #             continue
    #         for offset in offsets:
    #             center = min(positions, key=lambda pos: abs(pos - offset[0]))
    #             offset[0] = max(offset[0], center)
    #         collection.set_offsets(offsets)

    ax.set_xlabel(x_label, fontsize=fontsize, fontweight="bold")
    ax.set_ylabel(y_label or metric.capitalize(), fontsize=fontsize, fontweight="bold")
    if title is not None:
        ax.set_title(title, fontsize=fontsize + 2)
    ax.tick_params(axis="both", labelsize=fontsize-2)
    ax.set_xticks(range(len(order)))
    ax.set_xticklabels(
        order,
        rotation=x_tick_rotation,
        ha="right" if x_tick_rotation else "center",
    )
    if log_y:
        positive_values = pd.to_numeric(plot_df[metric], errors="coerce")
        if (positive_values.dropna() <= 0).any():
            raise ValueError("log_y=True requires all plotted values to be positive.")
        ax.set_yscale("log")

    if y_lim is not None:
        ax.set_ylim(*y_lim)
    elif log_y:
        ax.set_ylim(bottom=pd.to_numeric(plot_df[metric], errors="coerce").min() * 0.8)
    elif is_scalar_metric:
        ax.set_ylim(bottom=0)
    else:
        ax.set_ylim(0, 1.05)
    plt.tight_layout()

    if save_dir is not None:
        save_dir = ensure_long_path(Path(save_dir))
        os.makedirs(save_dir, exist_ok=True)
        if file_name is None:
            model_name = "_".join(str(value) for value in (_selection_values(model) or ["all"]))
            file_name = f"{model_name}_{metric}_boxplot.png"
        fig.savefig(
            ensure_long_path(save_dir / file_name),
            bbox_inches="tight",
            format="png",
            dpi=900 if high_quality else 100,
        )

    if show:
        plt.show()
    else:
        plt.close(fig)


def plot_model_feature_importance_stability_comparison(
    df: pd.DataFrame,
    model: Any = None,
    kernel_triples: Any = None,
    tree_feature_importance: str = "MDI",
    include_kernel_config: bool = True,
    dataset_as_experiment_points: bool = False,
    figsize: tuple = (9, 5),
    fontsize: int = 12,
    title: Optional[str] = None,
    x_label: str = "Model",
    y_label: str = "Kendall's W (feature stability)",
    x_tick_rotation: int = 35,
    y_lim: Optional[tuple] = (0, 1.05),
    show: bool = True,
    high_quality: bool = True,
    save_dir: Optional[Path] = HERE / "result_analysis",
    file_name: Optional[str] = None,
) -> pd.DataFrame:
    """
    Plot feature-importance stability across selected model configurations.

    GP-style models use ``lengthscale_kendalls_w``. Tree models use either
    ``feature_importance_MDI_kendalls_w`` or
    ``feature_importance_SHAP_kendalls_w``, selected with
    ``tree_feature_importance="MDI"`` or ``"SHAP"``.
    """
    tree_stability_col = _tree_feature_stability_column(tree_feature_importance)
    required_columns = {
        "paper",
        "target",
        "model",
        "fp kernel",
        "count kernel",
        "mixing method",
        "lengthscale_kendalls_w",
    }
    missing_columns = required_columns.difference(df.columns)
    if missing_columns:
        raise ValueError(
            "df is missing required columns: "
            + ", ".join(sorted(missing_columns))
        )

    filtered_df = _filter_kernel_triples(df, kernel_triples, keep_missing=True)
    filtered_df = _filter_selection(filtered_df, "model", model)
    if filtered_df["model"].apply(_is_tree_model).any() and tree_stability_col not in df.columns:
        raise ValueError(f"df is missing required column: {tree_stability_col}")

    rows = []
    for _, row in filtered_df.iterrows():
        stability_col = (
            tree_stability_col
            if _is_tree_model(row["model"])
            else "lengthscale_kendalls_w"
        )
        try:
            stability_value = float(row[stability_col])
        except (TypeError, ValueError):
            continue
        if np.isnan(stability_value):
            continue

        rows.append({
            "dataset": row["paper"],
            "target": row["target"],
            "model": row["model"],
            "fp kernel": row["fp kernel"],
            "count kernel": row["count kernel"],
            "mixing method": row["mixing method"],
            "feature stability source": _feature_stability_source_label(stability_col),
            "feature stability": stability_value,
        })

    plot_df = pd.DataFrame(rows)
    if plot_df.empty:
        raise ValueError(
            "No feature-importance stability values found for "
            f"model={model}, kernel_triples={kernel_triples}, and "
            f"tree_feature_importance={tree_feature_importance}."
        )

    x_col = "model configuration"
    plot_df[x_col] = plot_df.apply(
        lambda row: _model_config_label(row, include_kernel_config),
        axis=1,
    )

    if dataset_as_experiment_points:
        group_cols = [
            "dataset",
            "target",
            x_col,
            "model",
            "fp kernel",
            "count kernel",
            "mixing method",
            "feature stability source",
        ]
        plot_df = (
            plot_df.groupby(group_cols, dropna=False, as_index=False)["feature stability"]
            .mean()
            .copy()
        )

    selected_models = _selection_values(model) or MODELS
    model_order = {str(model_name).lower(): idx for idx, model_name in enumerate(selected_models)}
    config_order = plot_df[
        [x_col, "model", "fp kernel", "count kernel", "mixing method"]
    ].drop_duplicates(subset=[x_col])
    config_order["sort_key"] = config_order.apply(
        lambda row: _model_config_sort_key(row, model_order),
        axis=1,
    )
    order = config_order.sort_values("sort_key", kind="mergesort")[x_col].tolist()
    violin_palette = {
        label: _model_type_color(plot_df.loc[plot_df[x_col] == label, "model"].iloc[0])
        for label in order
    }

    fig, ax = plt.subplots(figsize=figsize)
    sns.violinplot(
        data=plot_df,
        x=x_col,
        y="feature stability",
        hue=x_col,
        order=order,
        hue_order=order,
        ax=ax,
        width=0.65,
        palette=violin_palette,
        inner_kws=dict(box_width=10, whis_width=2, color=".65", solid_capstyle="round"),
        cut=0,
        linewidth=1,
        dodge=False,
        legend=False,
    )
    positions = list(range(len(order)))
    for collection in ax.collections:
        if not isinstance(collection, PolyCollection):
            continue
        for path in collection.get_paths():
            vertices = path.vertices
            if len(vertices) == 0:
                continue
            center = min(positions, key=lambda pos: abs(pos - np.median(vertices[:, 0])))
            vertices[:, 0] = np.maximum(vertices[:, 0], center)
    for line in ax.lines:
        x_data = np.asarray(line.get_xdata(), dtype=float)
        if len(x_data) == 0:
            continue
        center = min(positions, key=lambda pos: abs(pos - np.mean(x_data)))
        line.set_xdata(np.maximum(x_data, center))

    ax.set_xlabel(x_label, fontsize=fontsize, fontweight="bold")
    ax.set_ylabel(y_label, fontsize=fontsize, fontweight="bold")
    if title is not None:
        ax.set_title(title, fontsize=fontsize + 2)
    ax.tick_params(axis="both", labelsize=fontsize - 2)
    ax.set_xticks(range(len(order)))
    ax.set_xticklabels(
        order,
        rotation=x_tick_rotation,
        ha="right" if x_tick_rotation else "center",
    )
    if y_lim is not None:
        ax.set_ylim(*y_lim)
    plt.tight_layout()

    if save_dir is not None:
        save_dir = ensure_long_path(Path(save_dir))
        os.makedirs(save_dir, exist_ok=True)
        if file_name is None:
            model_name = "_".join(str(value) for value in (_selection_values(model) or ["all"]))
            tree_source = _feature_stability_source_label(tree_stability_col).lower()
            file_name = f"{model_name}_{tree_source}_feature_stability_comparison.png"
        fig.savefig(
            ensure_long_path(save_dir / file_name),
            bbox_inches="tight",
            format="png",
            dpi=900 if high_quality else 100,
        )

    if show:
        plt.show()
    else:
        plt.close(fig)

    return plot_df


def plot_hybridization_method_comparison(
    df: pd.DataFrame,
    metric: str = "r2",
    model: Any = "GPytorchMAP",
    fp_kernels: Any = None,
    count_kernels: Any = None,
    mixing_methods: Any = None,
    kernel_triples: Any = None,
    dataset_as_experiment_points: bool = False,
    figsize: tuple = (6, 5),
    fontsize: int = 12,
    title: Optional[str] = None,
    x_label: str = "Hybridization method",
    y_label: Optional[str] = None,
    x_tick_rotation: int = 0,
    y_lim: Optional[tuple] = None,
    log_y: bool = False,
    show: bool = True,
    high_quality: bool = True,
    save_dir: Optional[Path] = HERE / "result_analysis",
    file_name: Optional[str] = None,
) -> pd.DataFrame:
    """
    Make a violin plot comparing GP kernel hybridization methods.

    Hybridization is stored in ``result_df`` as ``mixing method``. Score
    metrics use ``<metric>_seed_fold_scores`` when available; scalar columns
    such as ``Running time (CPU)`` are plotted directly. Use ``fp_kernels`` and
    ``count_kernels`` to compare methods within a kernel family, or
    ``kernel_triples`` to select exact ``(fp, count, mixing)`` combinations.
    """
    df = _filter_kernel_triples(df, kernel_triples, keep_missing=False)
    is_scalar_metric = metric in df.columns and f"{metric}_seed_fold_scores" not in df.columns
    plot_df = _expand_master_scores_for_profile(
        df,
        metric=metric,
        model=model,
        fp_kernels=fp_kernels,
        count_kernels=count_kernels,
        mixing_methods=mixing_methods,
    )
    plot_df = plot_df.dropna(subset=["mixing method"]).copy()
    if plot_df.empty:
        raise ValueError(
            f"No {metric} values found for model={model} with "
            f"fp_kernels={fp_kernels}, count_kernels={count_kernels}, "
            f"mixing_methods={mixing_methods}, and kernel_triples={kernel_triples}."
        )

    plot_df["hybridization method"] = plot_df["mixing method"].map(_mixing_method_label)
    selected_models = _selection_values(model)
    include_model = selected_models is None or len(selected_models) > 1
    x_col = "model hybridization method" if include_model else "hybridization method"
    if include_model:
        plot_df[x_col] = (
            plot_df["model"].astype(str)
            + "\n"
            + plot_df["hybridization method"].astype(str)
        )

    if dataset_as_experiment_points:
        group_cols = [
            "dataset",
            "target",
            x_col,
            "hybridization method",
            "model",
            "mixing method",
        ]
        group_cols = list(dict.fromkeys(group_cols))
        plot_df = (
            plot_df.groupby(group_cols, dropna=False, as_index=False)[metric]
            .mean()
            .copy()
        )

    selected_mixes = _selection_values(mixing_methods) or globals()["mixing_methods"]
    selected_models = selected_models or plot_df["model"].drop_duplicates().tolist()
    method_order = {str(method).lower(): idx for idx, method in enumerate(selected_mixes)}
    model_order = {str(model_name).lower(): idx for idx, model_name in enumerate(selected_models)}

    config_order = plot_df[[x_col, "model", "mixing method"]].drop_duplicates(subset=[x_col])
    config_order["sort_key"] = config_order.apply(
        lambda row: (
            *_model_order_sort_key(row["model"], model_order),
            method_order.get(str(row["mixing method"]).lower(), len(method_order)),
            str(row["mixing method"]),
        ),
        axis=1,
    )
    config_order = config_order.sort_values("sort_key", kind="mergesort")
    order = config_order[x_col].tolist()
    order.extend(label for label in plot_df[x_col].drop_duplicates() if label not in order)

    fallback_colors = sns.color_palette("Set2", n_colors=max(len(order), 1))
    violin_palette = {}
    for idx, (_, row) in enumerate(config_order.iterrows()):
        method = str(row["mixing method"])
        violin_palette[row[x_col]] = HYBRIDIZATION_METHOD_COLORS.get(
            method,
            fallback_colors[idx % len(fallback_colors)],
        )

    fig, ax = plt.subplots(figsize=figsize)
    sns.violinplot(
        data=plot_df,
        x=x_col,
        y=metric,
        hue=x_col,
        order=order,
        hue_order=order,
        ax=ax,
        width=0.65,
        palette=violin_palette,
        inner_kws=dict(box_width=10, whis_width=2, color=".65", solid_capstyle="round"),
        cut=0,
        linewidth=1,
        dodge=False,
        legend=False,
    )
    positions = list(range(len(order)))
    for collection in ax.collections:
        if not isinstance(collection, PolyCollection):
            continue
        for path in collection.get_paths():
            vertices = path.vertices
            if len(vertices) == 0:
                continue
            center = min(positions, key=lambda pos: abs(pos - np.median(vertices[:, 0])))
            vertices[:, 0] = np.maximum(vertices[:, 0], center)
    for line in ax.lines:
        x_data = np.asarray(line.get_xdata(), dtype=float)
        if len(x_data) == 0:
            continue
        center = min(positions, key=lambda pos: abs(pos - np.mean(x_data)))
        line.set_xdata(np.maximum(x_data, center))

    ax.set_xlabel(x_label, fontsize=fontsize, fontweight="bold")
    ax.set_ylabel(y_label or metric.capitalize(), fontsize=fontsize, fontweight="bold")
    if title is not None:
        ax.set_title(title, fontsize=fontsize + 2)
    ax.tick_params(axis="both", labelsize=fontsize - 2)
    ax.set_xticks(range(len(order)))
    ax.set_xticklabels(
        order,
        rotation=x_tick_rotation,
        ha="right" if x_tick_rotation else "center",
    )
    if log_y:
        positive_values = pd.to_numeric(plot_df[metric], errors="coerce")
        if (positive_values.dropna() <= 0).any():
            raise ValueError("log_y=True requires all plotted values to be positive.")
        ax.set_yscale("log")

    if y_lim is not None:
        ax.set_ylim(*y_lim)
    elif log_y:
        ax.set_ylim(bottom=pd.to_numeric(plot_df[metric], errors="coerce").min() * 0.8)
    elif metric == "r2":
        ax.set_ylim(0, 1.05)
    elif is_scalar_metric or not _metric_higher_is_better(metric):
        ax.set_ylim(bottom=0)
    plt.tight_layout()

    if save_dir is not None:
        save_dir = ensure_long_path(Path(save_dir))
        os.makedirs(save_dir, exist_ok=True)
        if file_name is None:
            model_name = "_".join(str(value) for value in (_selection_values(model) or ["all"]))
            file_name = f"{model_name}_{metric}_hybridization_method_comparison.png"
        fig.savefig(
            ensure_long_path(save_dir / file_name),
            bbox_inches="tight",
            format="png",
            dpi=900 if high_quality else 100,
        )

    if show:
        plt.show()
    else:
        plt.close(fig)

    return plot_df


def plot_model_profile_comparison(
    prof_results: pd.DataFrame,
    model: Any = None,
    kernel_triples: Any = None,
    metric: Optional[str] = None,
    auc_column: str = "auc",
    tree_feature_importance: Optional[str] = None,
    include_kernel_config: bool = True,
    figsize: tuple = (7, 5),
    fontsize: int = 12,
    title: Optional[str] = None,
    x_label: str = "Model",
    y_label: str = "Performance profile AUC",
    x_tick_rotation: int = 35,
    y_lim: tuple = (0, 1.05),
    show_values: bool = True,
    show: bool = True,
    high_quality: bool = True,
    save_dir: Optional[Path] = HERE / "result_analysis",
    file_name: Optional[str] = None,
) -> pd.DataFrame:
    """
    Draw a barplot of profile AUC values from ``performance_plot_with_ranks``.

    Pass ``kernel_triples`` as one triple or a list of
    ``(fp_kernel, count_kernel, mixing_method)`` triples to select exact GP
    kernel combinations. Tree models are kept when their kernel columns are
    empty.

    If the input came from ``performance_plot_with_ranks`` with
    ``metric="feature_stability"``, pass ``tree_feature_importance`` to select
    MDI- or SHAP-based tree stability AUC rows when both are present.
    """
    required_columns = {
        "model",
        "fp kernel",
        "count kernel",
        "mixing method",
        auc_column,
    }
    missing_columns = required_columns.difference(prof_results.columns)
    if missing_columns:
        raise ValueError(
            "prof_results is missing required columns: "
            + ", ".join(sorted(missing_columns))
        )

    plot_df = prof_results.copy()
    plot_df = _filter_selection(plot_df, "model", model)
    plot_df = _filter_kernel_triples(plot_df, kernel_triples, keep_missing=True)
    plot_df = _filter_metric_selection(plot_df, metric)
    if tree_feature_importance is not None:
        tree_feature_source = _feature_stability_source_label(
            _tree_feature_stability_column(tree_feature_importance)
        )
        if "tree feature importance" in plot_df.columns:
            plot_df = _filter_selection(
                plot_df,
                "tree feature importance",
                tree_feature_source,
            )

    plot_df[auc_column] = pd.to_numeric(plot_df[auc_column], errors="coerce")
    plot_df = plot_df.dropna(subset=[auc_column])
    if plot_df.empty:
        raise ValueError(
            f"No profile AUC rows found for model={model} "
            f"with kernel_triples={kernel_triples}."
        )

    x_col = "model configuration"
    plot_df[x_col] = plot_df.apply(
        lambda row: _model_config_label(row, include_kernel_config),
        axis=1,
    )
    selected_models = _selection_values(model) or plot_df["model"].drop_duplicates().tolist()
    model_order = {str(model_name).lower(): idx for idx, model_name in enumerate(selected_models)}
    config_order = plot_df[
        [x_col, "model", "fp kernel", "count kernel", "mixing method"]
    ].drop_duplicates(subset=[x_col])
    config_order["sort_key"] = config_order.apply(
        lambda row: _model_config_sort_key(row, model_order),
        axis=1,
    )
    order = config_order.sort_values("sort_key", kind="mergesort")[x_col].tolist()
    order.extend(label for label in plot_df[x_col].drop_duplicates() if label not in order)

    order_lookup = {label: idx for idx, label in enumerate(order)}
    plot_df["_plot_order"] = plot_df[x_col].map(order_lookup)
    plot_df = (
        plot_df.sort_values("_plot_order", kind="stable")
        .drop(columns="_plot_order")
        .copy()
    )
    bar_palette = {
        label: _model_type_color(plot_df.loc[plot_df[x_col] == label, "model"].iloc[0])
        for label in order
    }

    fig, ax = plt.subplots(figsize=figsize)
    sns.barplot(
        data=plot_df,
        x=x_col,
        y=auc_column,
        hue=x_col,
        order=order,
        hue_order=order,
        palette=bar_palette,
        width=0.67,
        dodge=False,
        legend=False,
        errorbar=None,
        ax=ax,
    )

    if show_values:
        for patch in ax.patches:
            height = patch.get_height()
            if np.isnan(height):
                continue
            ax.text(
                patch.get_x() + patch.get_width() / 2,
                height + 0.01,
                f"{height:.2f}",
                ha="center",
                va="bottom",
                fontsize=fontsize - 2,
            )

    ax.set_xlabel(x_label, fontsize=fontsize, fontweight="bold")
    ax.set_ylabel(y_label, fontsize=fontsize-2, fontweight="bold")
    if title is not None:
        ax.set_title(title, fontsize=fontsize + 2)
    ax.tick_params(axis="both", labelsize=fontsize - 2)
    ax.set_xticks(range(len(order)))
    ax.set_xticklabels(
        order,
        rotation=x_tick_rotation,
        ha="right" if x_tick_rotation else "center",
    )
    ax.set_ylim(*y_lim)
    plt.tight_layout()

    if save_dir is not None:
        save_dir = ensure_long_path(Path(save_dir))
        os.makedirs(save_dir, exist_ok=True)
        if file_name is None:
            model_name = "_".join(str(value) for value in (_selection_values(model) or ["all"]))
            file_name = f"{model_name}_profile_auc_barplot.png"
        fig.savefig(
            ensure_long_path(save_dir / file_name),
            bbox_inches="tight",
            format="png",
            dpi=900 if high_quality else 100,
        )

    if show:
        plt.show()
    else:
        plt.close(fig)

    return plot_df


def plot_hybridization_profile_comparison(
    prof_results: pd.DataFrame,
    model: Any = "GPytorchMAP",
    fp_kernels: Any = None,
    count_kernels: Any = None,
    mixing_methods: Any = None,
    kernel_triples: Any = None,
    metric: Optional[str] = None,
    auc_column: str = "auc",
    figsize: tuple = (5, 5),
    fontsize: int = 12,
    title: Optional[str] = None,
    x_label: str = "Hybridization method",
    y_label: str = "Performance profile AUC",
    x_tick_rotation: int = 0,
    y_lim: tuple = (0, 1.05),
    show_values: bool = True,
    show: bool = True,
    high_quality: bool = True,
    save_dir: Optional[Path] = HERE / "result_analysis",
    file_name: Optional[str] = None,
) -> pd.DataFrame:
    """
    Draw a barplot comparing hybridization methods from profile AUC results.

    The input should be the DataFrame returned by
    ``performance_plot_with_ranks`` or
    ``performance_plot_hybridization_with_ranks``. Hybridization is stored as
    ``mixing method``. Use ``fp_kernels`` and ``count_kernels`` to compare
    methods within a kernel family, or ``kernel_triples`` to select exact
    ``(fp_kernel, count_kernel, mixing_method)`` combinations. When the input
    has one AUC per kernel combination, bars summarize mean +/- std across
    kernels. When the input has one direct AUC per hybridization method, bars
    show that direct AUC.
    """
    required_columns = {
        "model",
        "fp kernel",
        "count kernel",
        "mixing method",
        auc_column,
    }
    missing_columns = required_columns.difference(prof_results.columns)
    if missing_columns:
        raise ValueError(
            "prof_results is missing required columns: "
            + ", ".join(sorted(missing_columns))
        )

    plot_df = prof_results.copy()
    plot_df = _filter_selection(plot_df, "model", model)
    plot_df = _filter_selection(plot_df, "fp kernel", fp_kernels, keep_missing=True)
    plot_df = _filter_selection(plot_df, "count kernel", count_kernels, keep_missing=True)
    plot_df = _filter_selection(plot_df, "mixing method", mixing_methods)
    plot_df = _filter_kernel_triples(plot_df, kernel_triples, keep_missing=True)
    if metric is not None and "metric" in plot_df.columns:
        plot_df = _filter_selection(plot_df, "metric", metric)

    plot_df[auc_column] = pd.to_numeric(plot_df[auc_column], errors="coerce")
    plot_df = plot_df.dropna(subset=["mixing method", auc_column]).copy()
    if plot_df.empty:
        raise ValueError(
            f"No profile AUC rows found for model={model} with "
            f"fp_kernels={fp_kernels}, count_kernels={count_kernels}, "
            f"mixing_methods={mixing_methods}, and kernel_triples={kernel_triples}."
        )

    plot_df["hybridization method"] = plot_df["mixing method"].map(_mixing_method_label)
    selected_models = _selection_values(model)
    include_model = selected_models is None or len(selected_models) > 1
    x_col = "model hybridization method" if include_model else "hybridization method"
    if include_model:
        plot_df[x_col] = (
            plot_df["model"].astype(str)
            + "\n"
            + plot_df["hybridization method"].astype(str)
        )

    selected_mixes = _selection_values(mixing_methods) or globals()["mixing_methods"]
    selected_models = selected_models or plot_df["model"].drop_duplicates().tolist()
    method_order = {str(method).lower(): idx for idx, method in enumerate(selected_mixes)}
    model_order = {str(model_name).lower(): idx for idx, model_name in enumerate(selected_models)}

    config_order = plot_df[[x_col, "model", "mixing method"]].drop_duplicates(subset=[x_col])
    config_order["sort_key"] = config_order.apply(
        lambda row: (
            *_model_order_sort_key(row["model"], model_order),
            method_order.get(str(row["mixing method"]).lower(), len(method_order)),
            str(row["mixing method"]),
        ),
        axis=1,
    )
    config_order = config_order.sort_values("sort_key", kind="mergesort")
    order = config_order[x_col].tolist()
    order.extend(label for label in plot_df[x_col].drop_duplicates() if label not in order)

    order_lookup = {label: idx for idx, label in enumerate(order)}
    plot_df["_plot_order"] = plot_df[x_col].map(order_lookup)
    plot_df = (
        plot_df.sort_values("_plot_order", kind="stable")
        .drop(columns="_plot_order")
        .copy()
    )

    direct_profile_auc = (
        (
            "profile grouping" in plot_df.columns
            and plot_df["profile grouping"].eq("hybridization method").all()
        )
        or plot_df.groupby(x_col, dropna=False)[auc_column].size().max() == 1
    )
    if direct_profile_auc:
        summary_df = plot_df.drop_duplicates(subset=[x_col]).copy()
        summary_df[f"{auc_column}_mean"] = summary_df[auc_column]
        if f"{auc_column}_std" not in summary_df.columns:
            summary_df[f"{auc_column}_std"] = 0.0
        if f"{auc_column}_count" not in summary_df.columns:
            summary_df[f"{auc_column}_count"] = 1
    else:
        summary_group_cols = list(dict.fromkeys([
            x_col,
            "hybridization method",
            "model",
            "mixing method",
        ]))
        summary_df = (
            plot_df.groupby(summary_group_cols, dropna=False)[auc_column]
            .agg(["mean", "std", "count"])
            .reset_index()
            .rename(columns={
                "mean": f"{auc_column}_mean",
                "std": f"{auc_column}_std",
                "count": f"{auc_column}_count",
            })
        )
    summary_df[f"{auc_column}_std"] = summary_df[f"{auc_column}_std"].fillna(0.0)
    summary_df["_plot_order"] = summary_df[x_col].map(order_lookup)
    summary_df = (
        summary_df.sort_values("_plot_order", kind="stable")
        .drop(columns="_plot_order")
        .copy()
    )
    summary_lookup = summary_df.set_index(x_col)[
        [f"{auc_column}_mean", f"{auc_column}_std", f"{auc_column}_count"]
    ]
    plot_df = plot_df.join(summary_lookup, on=x_col)

    fallback_colors = sns.color_palette("Set2", n_colors=max(len(order), 1))
    bar_palette = {}
    for idx, (_, row) in enumerate(config_order.iterrows()):
        method = str(row["mixing method"])
        bar_palette[row[x_col]] = HYBRIDIZATION_METHOD_COLORS.get(
            method,
            fallback_colors[idx % len(fallback_colors)],
        )

    fig, ax = plt.subplots(figsize=figsize)
    sns.barplot(
        data=summary_df,
        x=x_col,
        y=f"{auc_column}_mean",
        hue=x_col,
        order=order,
        hue_order=order,
        palette=bar_palette,
        width=0.67,
        dodge=False,
        legend=False,
        errorbar=None,
        ax=ax,
    )

    summary_by_label = summary_df.set_index(x_col)
    max_label_y = None
    for patch, label in zip(ax.patches, order):
        if label not in summary_by_label.index:
            continue
        mean_value = float(summary_by_label.loc[label, f"{auc_column}_mean"])
        std_value = float(summary_by_label.loc[label, f"{auc_column}_std"])
        x_position = patch.get_x() + patch.get_width() / 2
        if not direct_profile_auc:
            ax.errorbar(
                x_position,
                mean_value,
                yerr=std_value,
                fmt="none",
                ecolor="black",
                elinewidth=1.4,
                capsize=4,
                capthick=1.4,
                zorder=4,
            )
        label_y = mean_value + (0 if direct_profile_auc else std_value) + 0.015
        max_label_y = label_y if max_label_y is None else max(max_label_y, label_y)

    if show_values:
        for patch, label in zip(ax.patches, order):
            if label not in summary_by_label.index:
                continue
            mean_value = float(summary_by_label.loc[label, f"{auc_column}_mean"])
            std_value = float(summary_by_label.loc[label, f"{auc_column}_std"])
            label_text = (
                f"{mean_value:.2f}"
                if direct_profile_auc
                else f"{mean_value:.2f} ± {std_value:.2f}"
            )
            ax.text(
                patch.get_x() + patch.get_width() / 2,
                mean_value + (0 if direct_profile_auc else std_value) + 0.015,
                label_text,
                ha="center",
                va="bottom",
                fontsize=fontsize - 2,
            )

    ax.set_xlabel(x_label, fontsize=fontsize, fontweight="bold")
    ax.set_ylabel(y_label, fontsize=fontsize, fontweight="bold")
    if title is not None:
        ax.set_title(title, fontsize=fontsize + 2)
    ax.tick_params(axis="both", labelsize=fontsize - 2)
    ax.set_xticks(range(len(order)))
    ax.set_xticklabels(
        order,
        rotation=x_tick_rotation,
        ha="right" if x_tick_rotation else "center",
    )
    ax.set_ylim(*y_lim)
    if max_label_y is not None:
        bottom, top = ax.get_ylim()
        ax.set_ylim(bottom, max(top, max_label_y + 0.08))
    plt.tight_layout()

    if save_dir is not None:
        save_dir = ensure_long_path(Path(save_dir))
        os.makedirs(save_dir, exist_ok=True)
        if file_name is None:
            model_name = "_".join(str(value) for value in (_selection_values(model) or ["all"]))
            file_name = f"{model_name}_hybridization_profile_auc_barplot.png"
        fig.savefig(
            ensure_long_path(save_dir / file_name),
            bbox_inches="tight",
            format="png",
            dpi=900 if high_quality else 100,
        )

    if show:
        plt.show()
    else:
        plt.close(fig)

    return plot_df


def plot_model_performance_TOPSIS(
    df: pd.DataFrame,
    metrics: List[str],
    criteria_weights: Optional[List[float]] = None,
    criteria_types: Optional[List[int]] = None,
    model: Any = None,
    fp_kernels: Any = None,
    count_kernels: Any = None,
    mixing_methods: Any = None,
    kernel_triples: Any = None,
    tree_feature_importance: str = "MDI",
    include_kernel_config: bool = True,
    dataset_as_experiment_points: bool = True,
    figsize: tuple = (7, 5),
    fontsize: int = 12,
    title: Optional[str] = None,
    x_label: str = "Model",
    y_label: str = "TOPSIS score",
    x_tick_rotation: int = 35,
    y_lim: tuple = (0, 1.05),
    show_values: bool = True,
    show: bool = True,
    high_quality: bool = True,
    save_dir: Optional[Path] = HERE / "result_analysis",
    file_name: Optional[str] = None,
) -> pd.DataFrame:
    """
    Compare model configurations with TOPSIS using one or more metrics.

    ``df`` should be the master ``result_df`` from
    ``build_master_performance_data``. Each metric is averaged before TOPSIS
    by exact model/FP/count/mixing configuration. Tree models, which have no
    kernel columns, remain one row per model. TOPSIS is then run on that mean
    table, with model configurations as alternatives and ``metrics`` as
    criteria. ``include_kernel_config`` is kept only for backward-compatible
    calls and does not collapse selected configurations.
    """
    if not metrics:
        raise ValueError("metrics must contain at least one metric name.")

    metrics = list(dict.fromkeys(metrics))
    if criteria_weights is None:
        criteria_weights = [1.0] * len(metrics)
    if len(criteria_weights) != len(metrics):
        raise ValueError("criteria_weights must have the same length as metrics.")

    if criteria_types is None:
        criteria_types = [
            1 if _metric_higher_is_better(metric) else -1
            for metric in metrics
        ]
    if len(criteria_types) != len(metrics):
        raise ValueError("criteria_types must have the same length as metrics.")
    if any(criteria_type not in {-1, 1} for criteria_type in criteria_types):
        raise ValueError("criteria_types values must be 1 or -1.")

    filtered_df = _filter_kernel_triples(df, kernel_triples, keep_missing=True)
    config_group_cols = ["model", "fp kernel", "count kernel", "mixing method"]

    metric_frames = []
    for metric in metrics:
        metric_df = _expand_master_scores_for_profile(
            filtered_df,
            metric=metric,
            model=model,
            fp_kernels=fp_kernels,
            count_kernels=count_kernels,
            mixing_methods=mixing_methods,
            tree_feature_importance=tree_feature_importance,
        )
        if metric_df.empty:
            raise ValueError(
                f"No {metric} values found for model={model} with "
                f"fp_kernels={fp_kernels}, count_kernels={count_kernels}, "
                f"mixing_methods={mixing_methods}, and "
                f"kernel_triples={kernel_triples}."
            )

        group_cols = [
            column for column in config_group_cols if column in metric_df.columns
        ]
        if dataset_as_experiment_points:
            dataset_group_cols = [
                column
                for column in ["dataset", "target"] + group_cols
                if column in metric_df.columns
            ]
            metric_df = (
                metric_df.groupby(dataset_group_cols, dropna=False, as_index=False)[metric]
                .mean()
                .copy()
            )
        metric_df = (
            metric_df.groupby(group_cols, dropna=False, as_index=False)[metric]
            .mean()
            .copy()
        )
        metric_df["metric"] = metric
        metric_df["metric value"] = metric_df[metric]
        metric_frames.append(metric_df[group_cols + ["metric", "metric value"]])

    combined_df = pd.concat(metric_frames, ignore_index=True)
    index_cols = [
        column
        for column in config_group_cols
        if column in combined_df.columns
    ]
    missing_index_value = "__missing_kernel_config__"
    pivot_df = combined_df.copy()
    for column in index_cols:
        pivot_df[column] = pivot_df[column].where(
            pivot_df[column].notna(),
            missing_index_value,
        )
    plot_df = (
        pivot_df.pivot_table(
            index=index_cols,
            columns="metric",
            values="metric value",
            aggfunc="mean",
        )
        .reset_index()
    )
    plot_df = plot_df.replace(missing_index_value, np.nan)
    missing_metrics = [metric for metric in metrics if metric not in plot_df.columns]
    if missing_metrics:
        raise ValueError(
            "No TOPSIS values could be assembled for metrics: "
            + ", ".join(missing_metrics)
        )

    plot_df = plot_df.dropna(subset=metrics).copy()
    if len(plot_df) < 2:
        raise ValueError(
            "TOPSIS needs at least two model configurations with complete "
            "mean metric values."
        )

    plot_df["TOPSIS alternative"] = plot_df.apply(
        lambda row: _kernel_label(row, include_model=True),
        axis=1,
    )
    plot_df["model configuration"] = plot_df.apply(
        lambda row: _model_config_label(row, include_kernel_config),
        axis=1,
    )

    topsis_input = plot_df.set_index("TOPSIS alternative")[metrics].astype(float)
    topsis_scores = run_topsis(
        topsis_input,
        criteria_weights=criteria_weights,
        criteria_types=criteria_types,
    )
    plot_df["TOPSIS_Score"] = plot_df["TOPSIS alternative"].map(topsis_scores)

    selected_models = _selection_values(model) or plot_df["model"].drop_duplicates().tolist()
    model_order = {str(model_name).lower(): idx for idx, model_name in enumerate(selected_models)}
    config_order_cols = [
        "model configuration",
        "model",
        "fp kernel",
        "count kernel",
        "mixing method",
    ]
    config_order = plot_df[config_order_cols].drop_duplicates(
        subset=["model configuration"]
    )
    config_order["sort_key"] = config_order.apply(
        lambda row: (
            *_model_config_sort_key(row, model_order),
            str(row.get("fp kernel", "")),
            str(row.get("count kernel", "")),
            str(row.get("mixing method", "")),
        ),
        axis=1,
    )
    order = (
        config_order.sort_values("sort_key", kind="mergesort")["model configuration"]
        .tolist()
    )
    order.extend(
        label
        for label in plot_df["model configuration"].drop_duplicates()
        if label not in order
    )

    order_lookup = {label: idx for idx, label in enumerate(order)}
    plot_df["_plot_order"] = plot_df["model configuration"].map(order_lookup)
    plot_df = (
        plot_df.sort_values("_plot_order", kind="stable")
        .drop(columns="_plot_order")
        .copy()
    )
    bar_palette = {
        label: _model_type_color(
            plot_df.loc[plot_df["model configuration"] == label, "model"].iloc[0]
        )
        for label in order
    }

    fig, ax = plt.subplots(figsize=figsize)
    sns.barplot(
        data=plot_df,
        x="model configuration",
        y="TOPSIS_Score",
        hue="model configuration",
        order=order,
        hue_order=order,
        palette=bar_palette,
        width=0.67,
        dodge=False,
        legend=False,
        errorbar=None,
        ax=ax,
    )

    if show_values:
        for patch in ax.patches:
            height = patch.get_height()
            if np.isnan(height):
                continue
            ax.text(
                patch.get_x() + patch.get_width() / 2,
                height + 0.01,
                f"{height:.2f}",
                ha="center",
                va="bottom",
                fontsize=fontsize - 2,
            )

    ax.set_xlabel(x_label, fontsize=fontsize, fontweight="bold")
    ax.set_ylabel(y_label, fontsize=fontsize, fontweight="bold")
    if title is not None:
        ax.set_title(title, fontsize=fontsize + 2)
    ax.tick_params(axis="both", labelsize=fontsize - 2)
    ax.set_xticks(range(len(order)))
    ax.set_xticklabels(
        order,
        rotation=x_tick_rotation,
        ha="right" if x_tick_rotation else "center",
    )
    ax.set_ylim(*y_lim)
    plt.tight_layout()

    if save_dir is not None:
        save_dir = ensure_long_path(Path(save_dir))
        os.makedirs(save_dir, exist_ok=True)
        if file_name is None:
            model_name = "_".join(str(value) for value in (_selection_values(model) or ["all"]))
            metric_name = "_".join(metrics)
            file_name = f"{model_name}_{metric_name}_model_performance_TOPSIS.png"
        fig.savefig(
            ensure_long_path(save_dir / file_name),
            bbox_inches="tight",
            format="png",
            dpi=900 if high_quality else 100,
        )

    if show:
        plt.show()
    else:
        plt.close(fig)

    return plot_df


def plot_hybridization_topsis_comparison(
    df: pd.DataFrame,
    metrics: List[str],
    criteria_weights: Optional[List[float]] = None,
    criteria_types: Optional[List[int]] = None,
    model: Any = "GPytorchMAP",
    fp_kernels: Any = None,
    count_kernels: Any = None,
    mixing_methods: Any = None,
    kernel_triples: Any = None,
    tree_feature_importance: str = "MDI",
    dataset_as_experiment_points: bool = True,
    figsize: tuple = (5, 5),
    fontsize: int = 12,
    title: Optional[str] = None,
    x_label: str = "Hybridization method",
    y_label: str = "TOPSIS score",
    x_tick_rotation: int = 0,
    y_lim: tuple = (0, 1.05),
    show_values: bool = True,
    show: bool = True,
    high_quality: bool = True,
    save_dir: Optional[Path] = HERE / "result_analysis",
    file_name: Optional[str] = None,
) -> pd.DataFrame:
    """
    Compare hybridization methods with TOPSIS using one or more metrics.

    ``df`` should be the master ``result_df`` from
    ``build_master_performance_data``. Each metric is averaged for every exact
    model/FP/count/mixing configuration first. TOPSIS is then run on that
    configuration-mean table, with configurations as alternatives and
    ``metrics`` as the criteria. The bar plot summarizes those configuration
    TOPSIS scores by hybridization method.

    ``criteria_weights`` must match ``metrics``; equal weights are used when it
    is omitted. ``criteria_types`` uses 1 for benefit criteria and -1 for cost
    criteria. When omitted, directions are inferred from
    ``_metric_higher_is_better``.
    """
    if not metrics:
        raise ValueError("metrics must contain at least one metric name.")

    metrics = list(dict.fromkeys(metrics))
    if criteria_weights is None:
        criteria_weights = [1.0] * len(metrics)
    if len(criteria_weights) != len(metrics):
        raise ValueError("criteria_weights must have the same length as metrics.")

    if criteria_types is None:
        criteria_types = [
            1 if _metric_higher_is_better(metric) else -1
            for metric in metrics
        ]
    if len(criteria_types) != len(metrics):
        raise ValueError("criteria_types must have the same length as metrics.")
    if any(criteria_type not in {-1, 1} for criteria_type in criteria_types):
        raise ValueError("criteria_types values must be 1 or -1.")

    filtered_df = _filter_kernel_triples(df, kernel_triples, keep_missing=False)
    metric_frames = []
    config_group_cols = [
        "model",
        "fp kernel",
        "count kernel",
        "mixing method",
    ]

    for metric in metrics:
        metric_df = _expand_master_scores_for_profile(
            filtered_df,
            metric=metric,
            model=model,
            fp_kernels=fp_kernels,
            count_kernels=count_kernels,
            mixing_methods=mixing_methods,
            tree_feature_importance=tree_feature_importance,
        )
        metric_df = metric_df.dropna(
            subset=["fp kernel", "count kernel", "mixing method"]
        ).copy()
        if metric_df.empty:
            raise ValueError(
                f"No {metric} values found for model={model} with "
                f"fp_kernels={fp_kernels}, count_kernels={count_kernels}, "
                f"mixing_methods={mixing_methods}, and "
                f"kernel_triples={kernel_triples}."
            )

        group_cols = [
            column for column in config_group_cols if column in metric_df.columns
        ]
        if dataset_as_experiment_points:
            dataset_group_cols = [
                column
                for column in ["dataset", "target"] + group_cols
                if column in metric_df.columns
            ]
            metric_df = (
                metric_df.groupby(dataset_group_cols, dropna=False, as_index=False)[metric]
                .mean()
                .copy()
            )
        metric_df = (
            metric_df.groupby(group_cols, dropna=False, as_index=False)[metric]
            .mean()
            .copy()
        )
        metric_df["metric"] = metric
        metric_df["metric value"] = metric_df[metric]
        metric_frames.append(metric_df[group_cols + ["metric", "metric value"]])

    combined_df = pd.concat(metric_frames, ignore_index=True)
    index_cols = [
        column
        for column in config_group_cols
        if column in combined_df.columns
    ]
    metric_matrix_df = (
        combined_df.pivot_table(
            index=index_cols,
            columns="metric",
            values="metric value",
            aggfunc="mean",
        )
        .reset_index()
    )
    missing_metrics = [metric for metric in metrics if metric not in metric_matrix_df.columns]
    if missing_metrics:
        raise ValueError(
            "No TOPSIS values could be assembled for metrics: "
            + ", ".join(missing_metrics)
        )

    plot_df = metric_matrix_df.dropna(subset=metrics).copy()
    if plot_df.empty or plot_df["mixing method"].nunique(dropna=True) < 2:
        raise ValueError(
            "TOPSIS needs at least two hybridization methods with complete "
            "configuration-mean metric values."
        )
    plot_df["configuration"] = plot_df.apply(
        lambda row: _kernel_label(row, include_model=True),
        axis=1,
    )
    topsis_input = plot_df.set_index("configuration")[metrics].astype(float)
    topsis_scores = run_topsis(
        topsis_input,
        criteria_weights=criteria_weights,
        criteria_types=criteria_types,
    )
    plot_df["TOPSIS_Score"] = plot_df["configuration"].map(topsis_scores)
    plot_df["hybridization method"] = plot_df["mixing method"].map(
        _mixing_method_label
    )

    selected_models = _selection_values(model)
    include_model = selected_models is None or len(selected_models) > 1
    x_col = "model hybridization method" if include_model else "hybridization method"
    if include_model:
        plot_df[x_col] = (
            plot_df["model"].astype(str)
            + "\n"
            + plot_df["hybridization method"].astype(str)
        )

    selected_mixes = _selection_values(mixing_methods) or globals()["mixing_methods"]
    selected_models = selected_models or plot_df["model"].drop_duplicates().tolist()
    method_order = {str(method).lower(): idx for idx, method in enumerate(selected_mixes)}
    model_order = {str(model_name).lower(): idx for idx, model_name in enumerate(selected_models)}

    config_order = plot_df[[x_col, "model", "mixing method"]].drop_duplicates(subset=[x_col])
    config_order["sort_key"] = config_order.apply(
        lambda row: (
            *_model_order_sort_key(row["model"], model_order),
            method_order.get(str(row["mixing method"]).lower(), len(method_order)),
            str(row["mixing method"]),
        ),
        axis=1,
    )
    config_order = config_order.sort_values("sort_key", kind="mergesort")
    order = config_order[x_col].tolist()
    order.extend(label for label in plot_df[x_col].drop_duplicates() if label not in order)

    summary_group_cols = list(dict.fromkeys([
        x_col,
        "hybridization method",
        "model",
        "mixing method",
    ]))
    summary_df = (
        plot_df.groupby(summary_group_cols, dropna=False)["TOPSIS_Score"]
        .agg(["mean", "std", "count"])
        .reset_index()
        .rename(columns={
            "mean": "TOPSIS_Score_mean",
            "std": "TOPSIS_Score_std",
            "count": "TOPSIS_Score_count",
        })
    )
    summary_df["TOPSIS_Score_std"] = summary_df["TOPSIS_Score_std"].fillna(0.0)

    order_lookup = {label: idx for idx, label in enumerate(order)}
    summary_df["_plot_order"] = summary_df[x_col].map(order_lookup)
    summary_df = (
        summary_df.sort_values("_plot_order", kind="stable")
        .drop(columns="_plot_order")
        .copy()
    )
    plot_df = plot_df.join(
        summary_df.set_index(x_col)[
            ["TOPSIS_Score_mean", "TOPSIS_Score_std", "TOPSIS_Score_count"]
        ],
        on=x_col,
    )

    fallback_colors = sns.color_palette("Set2", n_colors=max(len(order), 1))
    bar_palette = {}
    for idx, (_, row) in enumerate(config_order.iterrows()):
        method = str(row["mixing method"])
        bar_palette[row[x_col]] = HYBRIDIZATION_METHOD_COLORS.get(
            method,
            fallback_colors[idx % len(fallback_colors)],
        )

    fig, ax = plt.subplots(figsize=figsize)
    sns.barplot(
        data=summary_df,
        x=x_col,
        y="TOPSIS_Score_mean",
        hue=x_col,
        order=order,
        hue_order=order,
        palette=bar_palette,
        width=0.67,
        dodge=False,
        legend=False,
        errorbar=None,
        ax=ax,
    )

    summary_by_label = summary_df.set_index(x_col)
    max_label_y = None
    for patch, label in zip(ax.patches, order):
        if label not in summary_by_label.index:
            continue
        mean_value = float(summary_by_label.loc[label, "TOPSIS_Score_mean"])
        std_value = float(summary_by_label.loc[label, "TOPSIS_Score_std"])
        x_position = patch.get_x() + patch.get_width() / 2
        ax.errorbar(
            x_position,
            mean_value,
            yerr=std_value,
            fmt="none",
            ecolor="black",
            elinewidth=1.4,
            capsize=4,
            capthick=1.4,
            zorder=4,
        )
        label_y = mean_value + std_value + 0.015
        max_label_y = label_y if max_label_y is None else max(max_label_y, label_y)

    if show_values:
        for patch, label in zip(ax.patches, order):
            if label not in summary_by_label.index:
                continue
            mean_value = float(summary_by_label.loc[label, "TOPSIS_Score_mean"])
            std_value = float(summary_by_label.loc[label, "TOPSIS_Score_std"])
            ax.text(
                patch.get_x() + patch.get_width() / 2,
                mean_value + std_value + 0.015,
                f"{mean_value:.2f} +/- {std_value:.2f}",
                ha="center",
                va="bottom",
                fontsize=fontsize - 2,
            )

    ax.set_xlabel(x_label, fontsize=fontsize, fontweight="bold")
    ax.set_ylabel(y_label, fontsize=fontsize, fontweight="bold")
    if title is not None:
        ax.set_title(title, fontsize=fontsize + 2)
    ax.tick_params(axis="both", labelsize=fontsize - 2)
    ax.set_xticks(range(len(order)))
    ax.set_xticklabels(
        order,
        rotation=x_tick_rotation,
        ha="right" if x_tick_rotation else "center",
    )
    ax.set_ylim(*y_lim)
    if max_label_y is not None:
        bottom, top = ax.get_ylim()
        ax.set_ylim(bottom, max(top, max_label_y + 0.08))
    plt.tight_layout()

    if save_dir is not None:
        save_dir = ensure_long_path(Path(save_dir))
        os.makedirs(save_dir, exist_ok=True)
        if file_name is None:
            model_name = "_".join(str(value) for value in (_selection_values(model) or ["all"]))
            metric_name = "_".join(metrics)
            file_name = f"{model_name}_{metric_name}_hybridization_topsis_barplot.png"
        fig.savefig(
            ensure_long_path(save_dir / file_name),
            bbox_inches="tight",
            format="png",
            dpi=900 if high_quality else 100,
        )

    if show:
        plt.show()
    else:
        plt.close(fig)


def _rank_percentile_matrix(
    score_matrix: pd.DataFrame,
    higher_is_better: bool,
) -> pd.DataFrame:
    rank_matrix = score_matrix.rank(
        axis=1,
        ascending=not higher_is_better,
        method="min",
    )
    valid_counts = score_matrix.notna().sum(axis=1).replace(0, np.nan)
    return rank_matrix.div(valid_counts, axis=0)


def _profile_task_columns(profile_df: pd.DataFrame) -> List[str]:
    task_cols = [column for column in ["dataset", "target"] if column in profile_df.columns]
    if not task_cols:
        raise ValueError("profile_df must contain at least a dataset or target column.")
    return task_cols


def _profile_index_columns(
    profile_df: pd.DataFrame,
    include_repeat: bool,
) -> List[str]:
    index_cols = _profile_task_columns(profile_df)
    if include_repeat and "repeat" in profile_df.columns:
        index_cols = index_cols + ["repeat"]
    return index_cols


def _hybridization_profile_index_columns(
    profile_df: pd.DataFrame,
    include_repeat: bool,
) -> List[str]:
    index_cols = _profile_index_columns(profile_df, include_repeat=include_repeat)
    for column in ["fp kernel", "count kernel"]:
        if column in profile_df.columns and column not in index_cols:
            index_cols.append(column)
    return index_cols


def _aggregate_profile_repeats(
    profile_df: pd.DataFrame,
    metric: str,
) -> pd.DataFrame:
    group_cols = [
        column
        for column in [
            "dataset",
            "target",
            "kernel",
            "model",
            "fp kernel",
            "count kernel",
            "mixing method",
        ]
        if column in profile_df.columns
    ]
    return (
        profile_df.groupby(group_cols, dropna=False, as_index=False)[metric]
        .mean()
        .copy()
    )


def _cat_gp_percentile_rank_matrix(
    profile_df: pd.DataFrame,
    metric: str,
    higher_is_better: bool,
    index_cols: List[str],
) -> pd.DataFrame:
    ranked_df = profile_df.copy()
    task_cols = _profile_task_columns(ranked_df)
    ranked_df["rank_repeat"] = ranked_df.groupby(
        task_cols,
        dropna=False,
    )[metric].rank(
        ascending=not higher_is_better,
        method="min",
    )
    max_rank = ranked_df.groupby(
        task_cols,
        dropna=False,
    )["rank_repeat"].transform("max")
    ranked_df["percentile_rank_repeat"] = ranked_df["rank_repeat"].div(
        max_rank.replace(0, np.nan)
    )

    return ranked_df.pivot_table(
        index=index_cols,
        columns="kernel",
        values="percentile_rank_repeat",
        aggfunc="first",
    )


def _profile_score_matrix(
    profile_df: pd.DataFrame,
    metric: str,
    index_cols: List[str],
) -> pd.DataFrame:
    return profile_df.pivot_table(
        index=index_cols,
        columns="kernel",
        values=metric,
        aggfunc="first",
    )


def _profile_curve_from_percentiles(
    percentiles: pd.Series,
    tau_grid: np.ndarray,
) -> Optional[np.ndarray]:
    values = percentiles.dropna().to_numpy(dtype=float)
    if len(values) == 0:
        return None

    return np.array([(values <= tau).mean() for tau in tau_grid], dtype=float)


def _profile_curves(
    percentile_matrix: pd.DataFrame,
    tau_grid: np.ndarray,
) -> Dict[str, np.ndarray]:
    curves = {}
    for kernel_name in percentile_matrix.columns:
        curve = _profile_curve_from_percentiles(
            percentile_matrix[kernel_name],
            tau_grid,
        )
        if curve is not None:
            curves[kernel_name] = curve

    return curves


def _repeat_profile_error_bands(
    profile_df: pd.DataFrame,
    metric: str,
    higher_is_better: bool,
    tau_grid: np.ndarray,
    error_band_stat: str,
) -> Dict[str, np.ndarray]:
    error_band_stat = error_band_stat.lower()
    if error_band_stat not in {"std", "sem"}:
        raise ValueError("error_band_stat must be either 'std' or 'sem'.")

    curves_by_kernel = {}
    if "repeat" not in profile_df.columns:
        return curves_by_kernel

    index_cols = _profile_index_columns(profile_df, include_repeat=False)
    for _, repeat_df in profile_df.groupby("repeat", dropna=False):
        repeat_percentile_matrix = _cat_gp_percentile_rank_matrix(
            repeat_df,
            metric,
            higher_is_better,
            index_cols,
        )
        repeat_curves = _profile_curves(repeat_percentile_matrix, tau_grid)
        for kernel_name, curve in repeat_curves.items():
            curves_by_kernel.setdefault(kernel_name, []).append(curve)

    error_bands = {}
    for kernel_name, curves in curves_by_kernel.items():
        curve_matrix = np.vstack(curves)
        if curve_matrix.shape[0] <= 1:
            error_bands[kernel_name] = np.zeros(curve_matrix.shape[1])
            continue

        error = np.nanstd(curve_matrix, axis=0, ddof=1)
        if error_band_stat == "sem":
            error = error / np.sqrt(curve_matrix.shape[0])
        error_bands[kernel_name] = error

    return error_bands


def _hybridization_repeat_profile_error_bands(
    profile_df: pd.DataFrame,
    metric: str,
    higher_is_better: bool,
    tau_grid: np.ndarray,
    error_band_stat: str,
) -> Dict[str, np.ndarray]:
    error_band_stat = error_band_stat.lower()
    if error_band_stat not in {"std", "sem"}:
        raise ValueError("error_band_stat must be either 'std' or 'sem'.")

    curves_by_method = {}
    if "repeat" not in profile_df.columns:
        return curves_by_method

    index_cols = _hybridization_profile_index_columns(
        profile_df,
        include_repeat=False,
    )
    for _, repeat_df in profile_df.groupby("repeat", dropna=False):
        repeat_percentile_matrix = _cat_gp_percentile_rank_matrix(
            repeat_df,
            metric,
            higher_is_better,
            index_cols,
        )
        repeat_curves = _profile_curves(repeat_percentile_matrix, tau_grid)
        for method_name, curve in repeat_curves.items():
            curves_by_method.setdefault(method_name, []).append(curve)

    error_bands = {}
    for method_name, curves in curves_by_method.items():
        curve_matrix = np.vstack(curves)
        if curve_matrix.shape[0] <= 1:
            error_bands[method_name] = np.zeros(curve_matrix.shape[1])
            continue

        error = np.nanstd(curve_matrix, axis=0, ddof=1)
        if error_band_stat == "sem":
            error = error / np.sqrt(curve_matrix.shape[0])
        error_bands[method_name] = error

    return error_bands


def _profile_auc_dataframe(
    profile_df: pd.DataFrame,
    aucs: Dict[str, float],
    sorted_names: List[str],
    metric: str,
) -> pd.DataFrame:
    metadata_cols = ["kernel", "model", "fp kernel", "count kernel", "mixing method"]
    optional_metadata_cols = ["feature stability source", "tree feature importance"]
    metadata_cols.extend(
        column for column in optional_metadata_cols if column in profile_df.columns
    )
    metadata = (
        profile_df[metadata_cols]
        .drop_duplicates(subset=["kernel"])
        .set_index("kernel")
    )
    rows = []

    for rank, kernel_name in enumerate(sorted_names, start=1):
        meta = metadata.loc[kernel_name]
        fp_kernel = meta["fp kernel"]
        count_kernel = meta["count kernel"]
        kernel_combination = (
            None
            if pd.isna(fp_kernel) or pd.isna(count_kernel)
            else f"{fp_kernel}-{count_kernel}"
        )

        row = {
            "rank": rank,
            "metric": metric,
            "model": meta["model"],
            "fp kernel": fp_kernel,
            "count kernel": count_kernel,
            "mixing method": meta["mixing method"],
            "auc": aucs[kernel_name],
        }
        for column in optional_metadata_cols:
            if column in metadata.columns:
                row[column] = meta[column]
        rows.append(row)

    return pd.DataFrame(rows)


def _hybridization_profile_auc_dataframe(
    profile_df: pd.DataFrame,
    aucs: Dict[str, float],
    sorted_names: List[str],
    metric: str,
) -> pd.DataFrame:
    metadata_cols = ["kernel", "model", "mixing method"]
    metadata = (
        profile_df[metadata_cols]
        .drop_duplicates(subset=["kernel"])
        .set_index("kernel")
    )
    fp_kernel_values = sorted(
        profile_df["fp kernel"].dropna().astype(str).unique().tolist()
    )
    count_kernel_values = sorted(
        profile_df["count kernel"].dropna().astype(str).unique().tolist()
    )
    rows = []

    for rank, method_name in enumerate(sorted_names, start=1):
        meta = metadata.loc[method_name]
        rows.append({
            "rank": rank,
            "metric": metric,
            "model": meta["model"],
            "fp kernel": np.nan,
            "count kernel": np.nan,
            "mixing method": meta["mixing method"],
            "hybridization method": method_name,
            "auc": aucs[method_name],
            "profile grouping": "hybridization method",
            "fp kernels": fp_kernel_values,
            "count kernels": count_kernel_values,
            "n_profile_points": profile_df.loc[
                profile_df["kernel"] == method_name
            ].shape[0],
        })

    return pd.DataFrame(rows)


def performance_plot_with_ranks(
    df: pd.DataFrame,
    metric: str = "r2",
    model: Any = "GPytorchMAP",
    fp_kernels: Any = None,
    count_kernels: Any = None,
    mixing_methods: Any = None,
    kernel_triples: Any = None,
    tree_feature_importance: str = "MDI",
    p_threshold: float = 0.05,
    title: Optional[str] = None,
    show: bool = True,
    high_quality: bool = True,
    dataset_as_experiment_points: bool = False,
    show_error_band: bool = False,
    error_band_stat: str = "std",
    error_band_alpha: float = 0.18,
    fontsize: float = 12,
    figsize: tuple = (7, 5),
    legend_ncols: Optional[int] = None,
    legend_nrows: Optional[int] = None,
    save_dir: Optional[Path] = HERE / "result_analysis",
    file_name: Optional[str] = "performance_profile.png",
):
    """
    Plot model/kernel performance profiles from the master result dataset.

    Following the cat_gp profile code, each row is one
    ``dataset * target * seed/fold`` repeat for a model/kernel. Percentile
    ranks are computed within each ``dataset * target`` task using all selected
    model/kernel repeats, then the plotted profile is the CDF of those ranks.
    Returns a DataFrame with one AUC row per model/kernel combination.

    For exact GP kernel choices, pass ``kernel_triples`` as one triple or a
    list of ``(fp_kernel, count_kernel, mixing_method)`` triples. The older
    separate ``fp_kernels``, ``count_kernels``, and ``mixing_methods`` filters
    are still supported.

    If ``dataset_as_experiment_points`` is True, the main profile first averages
    seed/fold scores within each dataset/target so each target contributes one
    profile point. If ``show_error_band`` is True, this dataset-level mode is
    used and the shaded band is the std/sem across seed/fold repeat profiles.
    ``legend_ncols`` and ``legend_nrows`` control the legend layout above the
    plot. If both are omitted, up to three columns are used.

    Use ``metric="feature_stability"`` to rank the mixed feature-stability
    metric used by ``plot_model_feature_importance_stability_comparison``:
    GP-style models use ``lengthscale_kendalls_w`` and tree models use MDI or
    SHAP stability according to ``tree_feature_importance``.
    """
    df = _filter_kernel_triples(df, kernel_triples, keep_missing=True)
    profile_df = _expand_master_scores_for_profile(
        df,
        metric=metric,
        model=model,
        fp_kernels=fp_kernels,
        count_kernels=count_kernels,
        mixing_methods=mixing_methods,
        tree_feature_importance=tree_feature_importance,
    )
    if profile_df.empty:
        raise ValueError(
            f"No {metric} values found for model={model} "
            f"with fp_kernels={fp_kernels}, count_kernels={count_kernels}, "
            f"mixing_methods={mixing_methods}, and kernel_triples={kernel_triples}."
        )

    higher_is_better = _metric_higher_is_better(metric)

    if show_error_band:
        dataset_as_experiment_points = True

    profile_rank_df = profile_df
    if dataset_as_experiment_points:
        profile_rank_df = _aggregate_profile_repeats(profile_df, metric)

    include_repeat = not dataset_as_experiment_points
    profile_index_cols = _profile_index_columns(
        profile_rank_df,
        include_repeat=include_repeat,
    )
    profile_score_matrix = _profile_score_matrix(
        profile_rank_df,
        metric,
        profile_index_cols,
    )
    percentile_matrix = _cat_gp_percentile_rank_matrix(
        profile_rank_df,
        metric,
        higher_is_better,
        profile_index_cols,
    )

    linn = np.linspace(0, 1, max(len(percentile_matrix), 2))
    all_percentages = _profile_curves(percentile_matrix, linn)
    error_bands = (
        _repeat_profile_error_bands(
            profile_df,
            metric,
            higher_is_better,
            linn,
            error_band_stat,
        )
        if show_error_band
        else {}
    )

    aucs = {
        kernel_name: auc(linn, percentages)
        for kernel_name, percentages in all_percentages.items()
    }
    sorted_names = [
        kernel_name
        for kernel_name, _ in sorted(aucs.items(), key=lambda x: x[1], reverse=True)
    ]
    auc_df = _profile_auc_dataframe(profile_df, aucs, sorted_names, metric)


    groups_of_same_ranks = []
    curr_l = [0]
    alternative = "greater" if higher_is_better else "less"
    for i in range(len(sorted_names)):
        if i == len(sorted_names) - 1:
            if len(curr_l) >= 2:
                groups_of_same_ranks.append(curr_l)
            continue

        paired_scores = profile_score_matrix[
            [sorted_names[i], sorted_names[i + 1]]
        ].dropna()
        if paired_scores.empty:
            continue

        try:
            _, p = wilcoxon(
                paired_scores[sorted_names[i]],
                paired_scores[sorted_names[i + 1]],
                alternative=alternative,
            )
        except ValueError:
            p = 1.0

        if p < p_threshold:
            if len(curr_l) >= 2:
                groups_of_same_ranks.append(curr_l)
            curr_l = [i + 1]
        else:
            curr_l.append(i + 1)

    fig, ax = plt.subplots(figsize=figsize)
    colors = sns.color_palette("tab20", n_colors=max(len(sorted_names), 1))

    for color, kernel_name in zip(colors, sorted_names):
        curve = np.asarray(all_percentages[kernel_name], dtype=float)
        ax.plot(
            linn,
            curve,
            color=color,
            label=f"{kernel_name}, {aucs[kernel_name]:.2f}",
        )
        if kernel_name in error_bands:
            error = np.asarray(error_bands[kernel_name], dtype=float)
            ax.fill_between(
                linn,
                np.clip(curve - error, 0, 1),
                np.clip(curve + error, 0, 1),
                color=color,
                alpha=error_band_alpha,
                linewidth=0,
            )

    if legend_ncols is not None and legend_ncols < 1:
        raise ValueError("legend_ncols must be at least 1.")
    if legend_nrows is not None and legend_nrows < 1:
        raise ValueError("legend_nrows must be at least 1.")

    if legend_ncols is None:
        if legend_nrows is None:
            legend_ncols = max(1, min(3, len(sorted_names)))
        else:
            legend_ncols = int(np.ceil(len(sorted_names) / legend_nrows))

    legend = ax.legend(
        loc="lower center",
        bbox_to_anchor=(0.5, 1.02),
        ncol=legend_ncols,
        borderaxespad=0,
        frameon=False,
        fontsize=fontsize,
    )

    fig.canvas.draw()

    ax.set_xlabel(r"$\tau$", fontsize=fontsize, fontweight="bold")
    ax.set_ylabel(r"$p_i(\tau)$", fontsize=fontsize, fontweight="bold")
    if title is not None:
        ax.set_title(
            title or f"{model} {metric} performance profile",
            fontsize=fontsize+2,
            fontweight="bold"
        )
    ax.yaxis.set_tick_params(labelsize=fontsize-2)
    ax.xaxis.set_tick_params(labelsize=fontsize-2)

    fig.canvas.draw()
    for group_rank in groups_of_same_ranks:
        add_legend_box(group_rank, legend, linewidth=1.2, alpha=1.0, pad=3)

    if save_dir is not None:
        save_dir = ensure_long_path(Path(save_dir))
        os.makedirs(save_dir, exist_ok=True)
        if file_name is None:
            file_name = f"{model}_{metric}_performance_profile.png"
        fig.savefig(
            ensure_long_path(save_dir / file_name),
            bbox_inches="tight",
            format="png",
            dpi=900 if high_quality else 100,
        )

    if show:
        plt.show()
    else:
        plt.close(fig)

    return auc_df


def performance_plot_hybridization_with_ranks(
    df: pd.DataFrame,
    metric: str = "r2",
    model: Any = "GPytorchMAP",
    fp_kernels: Any = None,
    count_kernels: Any = None,
    mixing_methods: Any = None,
    kernel_triples: Any = None,
    p_threshold: float = 0.05,
    title: Optional[str] = None,
    show: bool = True,
    high_quality: bool = True,
    dataset_as_experiment_points: bool = False,
    show_error_band: bool = False,
    error_band_stat: str = "std",
    error_band_alpha: float = 0.18,
    fontsize: float = 12,
    figsize: tuple = (7, 5),
    legend_ncols: Optional[int] = None,
    legend_nrows: Optional[int] = None,
    save_dir: Optional[Path] = HERE / "result_analysis",
    file_name: Optional[str] = "hybridization_performance_profile.png",
):
    """
    Plot performance profiles comparing hybridization methods.

    Unlike ``performance_plot_with_ranks``, this treats selected FP/count
    kernel combinations as experiment points together with datasets, targets,
    and seed/fold repeats. The returned DataFrame therefore contains one AUC
    row per hybridization method instead of one AUC per exact kernel
    combination.
    """
    df = _filter_kernel_triples(df, kernel_triples, keep_missing=False)
    profile_df = _expand_master_scores_for_profile(
        df,
        metric=metric,
        model=model,
        fp_kernels=fp_kernels,
        count_kernels=count_kernels,
        mixing_methods=mixing_methods,
    )
    profile_df = profile_df.dropna(
        subset=["fp kernel", "count kernel", "mixing method"]
    ).copy()
    if profile_df.empty:
        raise ValueError(
            f"No {metric} seed/fold scores found for model={model} "
            f"with fp_kernels={fp_kernels}, count_kernels={count_kernels}, "
            f"mixing_methods={mixing_methods}, and kernel_triples={kernel_triples}."
        )

    selected_models = _selection_values(model)
    include_model = selected_models is None or len(selected_models) > 1
    profile_df["hybridization method"] = profile_df["mixing method"].map(
        _mixing_method_label
    )
    if include_model:
        profile_df["kernel"] = (
            profile_df["model"].astype(str)
            + ": "
            + profile_df["hybridization method"].astype(str)
        )
    else:
        profile_df["kernel"] = profile_df["hybridization method"]

    higher_is_better = _metric_higher_is_better(metric)

    if show_error_band:
        dataset_as_experiment_points = True

    profile_rank_df = profile_df
    if dataset_as_experiment_points:
        profile_rank_df = _aggregate_profile_repeats(profile_df, metric)

    include_repeat = not dataset_as_experiment_points
    profile_index_cols = _hybridization_profile_index_columns(
        profile_rank_df,
        include_repeat=include_repeat,
    )
    profile_score_matrix = _profile_score_matrix(
        profile_rank_df,
        metric,
        profile_index_cols,
    )
    percentile_matrix = _cat_gp_percentile_rank_matrix(
        profile_rank_df,
        metric,
        higher_is_better,
        profile_index_cols,
    )

    linn = np.linspace(0, 1, max(len(percentile_matrix), 2))
    all_percentages = _profile_curves(percentile_matrix, linn)
    error_bands = (
        _hybridization_repeat_profile_error_bands(
            profile_df,
            metric,
            higher_is_better,
            linn,
            error_band_stat,
        )
        if show_error_band
        else {}
    )

    aucs = {
        method_name: auc(linn, percentages)
        for method_name, percentages in all_percentages.items()
    }
    sorted_names = [
        method_name
        for method_name, _ in sorted(aucs.items(), key=lambda x: x[1], reverse=True)
    ]
    auc_df = _hybridization_profile_auc_dataframe(
        profile_df,
        aucs,
        sorted_names,
        metric,
    )

    groups_of_same_ranks = []
    curr_l = [0]
    alternative = "greater" if higher_is_better else "less"
    for i in range(len(sorted_names)):
        if i == len(sorted_names) - 1:
            if len(curr_l) >= 2:
                groups_of_same_ranks.append(curr_l)
            continue

        paired_scores = profile_score_matrix[
            [sorted_names[i], sorted_names[i + 1]]
        ].dropna()
        if paired_scores.empty:
            continue

        try:
            _, p = wilcoxon(
                paired_scores[sorted_names[i]],
                paired_scores[sorted_names[i + 1]],
                alternative=alternative,
            )
        except ValueError:
            p = 1.0

        if p < p_threshold:
            if len(curr_l) >= 2:
                groups_of_same_ranks.append(curr_l)
            curr_l = [i + 1]
        else:
            curr_l.append(i + 1)

    method_lookup = (
        profile_df[["kernel", "mixing method"]]
        .drop_duplicates(subset=["kernel"])
        .set_index("kernel")["mixing method"]
        .to_dict()
    )

    fig, ax = plt.subplots(figsize=figsize)
    fallback_colors = sns.color_palette("Set2", n_colors=max(len(sorted_names), 1))
    for idx, method_name in enumerate(sorted_names):
        mix_method = str(method_lookup.get(method_name, method_name))
        color = HYBRIDIZATION_METHOD_COLORS.get(
            mix_method,
            fallback_colors[idx % len(fallback_colors)],
        )
        curve = np.asarray(all_percentages[method_name], dtype=float)
        ax.plot(
            linn,
            curve,
            color=color,
            label=f"{method_name}, {aucs[method_name]:.2f}",
        )
        if method_name in error_bands:
            error = np.asarray(error_bands[method_name], dtype=float)
            ax.fill_between(
                linn,
                np.clip(curve - error, 0, 1),
                np.clip(curve + error, 0, 1),
                color=color,
                alpha=error_band_alpha,
                linewidth=0,
            )

    if legend_ncols is not None and legend_ncols < 1:
        raise ValueError("legend_ncols must be at least 1.")
    if legend_nrows is not None and legend_nrows < 1:
        raise ValueError("legend_nrows must be at least 1.")

    if legend_ncols is None:
        if legend_nrows is None:
            legend_ncols = max(1, min(3, len(sorted_names)))
        else:
            legend_ncols = int(np.ceil(len(sorted_names) / legend_nrows))

    legend = ax.legend(
        loc="lower center",
        bbox_to_anchor=(0.5, 1.02),
        ncol=legend_ncols,
        borderaxespad=0,
        frameon=False,
        fontsize=fontsize,
    )

    fig.canvas.draw()

    ax.set_xlabel(r"$\tau$", fontsize=fontsize, fontweight="bold")
    ax.set_ylabel(r"$p_i(\tau)$", fontsize=fontsize, fontweight="bold")
    if title is not None:
        ax.set_title(title, fontsize=fontsize + 2, fontweight="bold")
    ax.yaxis.set_tick_params(labelsize=fontsize - 2)
    ax.xaxis.set_tick_params(labelsize=fontsize - 2)

    fig.canvas.draw()
    for group_rank in groups_of_same_ranks:
        add_legend_box(group_rank, legend, linewidth=1.2, alpha=1.0, pad=3)

    if save_dir is not None:
        save_dir = ensure_long_path(Path(save_dir))
        os.makedirs(save_dir, exist_ok=True)
        if file_name is None:
            model_name = "_".join(str(value) for value in (_selection_values(model) or ["all"]))
            file_name = f"{model_name}_{metric}_hybridization_performance_profile.png"
        fig.savefig(
            ensure_long_path(save_dir / file_name),
            bbox_inches="tight",
            format="png",
            dpi=900 if high_quality else 100,
        )

    if show:
        plt.show()
    else:
        plt.close(fig)

    return auc_df


def plot_profile_auc_heatmap(
    auc_df: pd.DataFrame,
    model: Any = None,
    fp_kernels: Any = None,
    count_kernels: Any = None,
    mixing_methods: Any = None,
    auc_column: str = "auc",
    figsize: tuple = (12, 4),
    title: Optional[str] = None,
    vmin: float = 0,
    vmax: float = 1,
    fontsize: int = 12,
    show: bool = True,
    high_quality: bool = True,
    plot_heatmap: bool = True,
    save_dir: Optional[Path] = HERE / "result_analysis",
    file_name: Optional[str] = "profile_auc_heatmap.png",
) -> pd.DataFrame:
    """
    Plot profile-performance AUC values as a heatmap.

    The input should be the DataFrame returned by ``performance_plot_with_ranks``.
    Rows are hybridization methods and columns are FP/count kernel combinations.
    """
    required_columns = {
        "model",
        "fp kernel",
        "count kernel",
        "mixing method",
        auc_column,
    }
    missing_columns = required_columns.difference(auc_df.columns)
    if missing_columns:
        raise ValueError(
            "auc_df is missing required columns: "
            + ", ".join(sorted(missing_columns))
        )

    plot_df = auc_df.copy()
    plot_df = _filter_selection(plot_df, "model", model)
    plot_df = _filter_selection(plot_df, "fp kernel", fp_kernels)
    plot_df = _filter_selection(plot_df, "count kernel", count_kernels)
    plot_df = _filter_selection(plot_df, "mixing method", mixing_methods)
    plot_df = plot_df.dropna(
        subset=["fp kernel", "count kernel", "mixing method", auc_column]
    ).copy()

    if plot_df.empty:
        raise ValueError(
            "No profile-AUC rows available for the requested model/kernel filters."
        )

    plot_df[auc_column] = pd.to_numeric(plot_df[auc_column], errors="coerce")
    plot_df = plot_df.dropna(subset=[auc_column])
    if plot_df.empty:
        raise ValueError("No numeric profile-AUC values available to plot.")

    plot_df["hybridization method"] = plot_df["mixing method"].map(
        lambda value: mixing_labels.get(str(value), str(value))
    )
    plot_df["kernel combination"] = (
        plot_df["fp kernel"].astype(str)
        + "-"
        + plot_df["count kernel"].astype(str)
    )

    heatmap_matrix = plot_df.pivot_table(
        index="hybridization method",
        columns="kernel combination",
        values=auc_column,
        aggfunc="mean",
    )

    selected_mixes = _selection_values(mixing_methods) or globals()["mixing_methods"]
    row_order = [
        mixing_labels.get(str(method), str(method))
        for method in selected_mixes
        if mixing_labels.get(str(method), str(method)) in heatmap_matrix.index
    ]
    row_order.extend(row for row in heatmap_matrix.index if row not in row_order)
    heatmap_matrix = heatmap_matrix.reindex(row_order)

    selected_fp = _selection_values(fp_kernels) or (fp_sk_kernels + fp_bit_kernels)
    selected_count = _selection_values(count_kernels) or globals()["count_kernels"]
    column_order = [
        f"{fp_kernel}-{count_kernel}"
        for fp_kernel in selected_fp
        for count_kernel in selected_count
        if f"{fp_kernel}-{count_kernel}" in heatmap_matrix.columns
    ]
    column_order.extend(
        column for column in heatmap_matrix.columns if column not in column_order
    )
    heatmap_matrix = heatmap_matrix.reindex(columns=column_order)

    annotations = heatmap_matrix.copy().astype(object)
    for row_label in heatmap_matrix.index:
        for column_label in heatmap_matrix.columns:
            value = heatmap_matrix.loc[row_label, column_label]
            annotations.loc[row_label, column_label] = (
                "" if pd.isna(value) else f"{value:.2f}"
            )

    if not plot_heatmap:
        return heatmap_matrix

    x_tick_labels = [
        str(column).replace("-", ":\n", 1)
        for column in heatmap_matrix.columns
    ]

    fig, ax = plt.subplots(figsize=figsize)
    custom_cmap = sns.color_palette("viridis", as_cmap=True)
    custom_cmap.set_bad(color="lightgray")
    hmap = sns.heatmap(
        heatmap_matrix,
        annot=annotations,
        fmt="",
        cmap=custom_cmap,
        cbar=True,
        vmin=vmin,
        vmax=vmax,
        ax=ax,
        mask=heatmap_matrix.isnull(),
        linewidths=0.5,
        linecolor="white",
        xticklabels=x_tick_labels,
        yticklabels=heatmap_matrix.index.tolist(),
        annot_kws={"fontsize": fontsize},
    )

    cbar = hmap.collections[0].colorbar
    if cbar is not None:
        cbar.set_label(
            "Performance profile AUC",
            fontsize=fontsize,
            fontweight="bold",
        )
        cbar.ax.tick_params(labelsize=fontsize)

    ax.set_xlabel("FP Kernel : Count Kernel", fontsize=fontsize + 2, fontweight="bold")
    ax.set_ylabel("Hybridization method", fontsize=fontsize + 2, fontweight="bold")
    if title is not None:
        ax.set_title(title, fontsize=fontsize + 3, fontweight="bold")
    ax.set_xticklabels(
        x_tick_labels,
        rotation=45,
        ha="right",
        fontsize=fontsize,
    )
    ax.set_yticklabels(ax.get_yticklabels(), rotation=0, fontsize=fontsize)

    plt.tight_layout()

    if save_dir is not None:
        save_dir = ensure_long_path(Path(save_dir))
        os.makedirs(save_dir, exist_ok=True)
        if file_name is None:
            file_name = "profile_auc_heatmap.png"
        fig.savefig(
            ensure_long_path(save_dir / file_name),
            bbox_inches="tight",
            format="png",
            dpi=900 if high_quality else 100,
        )

    if show:
        plt.show()
    else:
        plt.close(fig)

    return heatmap_matrix


def _datapoint_lookup() -> Dict[tuple[str, str], Any]:
    datapoint_lookup = {}
    for paper_name, paper_info in PAPER.items():
        targets = paper_info.get("target", [])
        n_datapoints = paper_info.get("n_datapoints", [])
        if len(targets) != len(n_datapoints):
            raise ValueError(
                f"PAPER entry for {paper_name!r} has {len(targets)} targets "
                f"but {len(n_datapoints)} datapoint counts."
            )
        for target, n_data in zip(targets, n_datapoints):
            datapoint_lookup[(paper_name, target)] = n_data

    return datapoint_lookup


def _add_datapoint_counts(plot_df: pd.DataFrame) -> pd.DataFrame:
    datapoint_lookup = _datapoint_lookup()
    plot_df = plot_df.copy()
    plot_df["n datapoints"] = plot_df.apply(
        lambda row: datapoint_lookup.get((row["dataset"], row["target"])),
        axis=1,
    )
    missing_size = plot_df["n datapoints"].isna()
    if missing_size.any():
        missing_pairs = (
            plot_df.loc[missing_size, ["dataset", "target"]]
            .drop_duplicates()
            .apply(lambda row: f"{row['dataset']} / {row['target']}", axis=1)
            .tolist()
        )
        raise ValueError(
            "Missing datapoint counts in PAPER for: "
            + "; ".join(missing_pairs)
        )
    plot_df["n datapoints"] = pd.to_numeric(plot_df["n datapoints"])
    return plot_df


def plot_model_performance_vs_data_number(
    df: pd.DataFrame,
    metric: str = "r2",
    model: Any = "GPytorchMAP",
    kernel_triples: Any = None,
    include_kernel_config: bool = True,
    dataset_as_experiment_points: bool = False,
    figsize: tuple = (10, 5),
    fontsize: int = 12,
    title: Optional[str] = None,
    x_label: str = "Number of datapoints",
    y_label: Optional[str] = None,
    x_tick_rotation: int = 0,
    y_lim: Optional[tuple] = None,
    log_y: bool = False,
    show: bool = True,
    high_quality: bool = True,
    save_dir: Optional[Path] = HERE / "result_analysis",
    file_name: Optional[str] = None,
) -> pd.DataFrame:
    """
    Make a box plot of model performance grouped by dataset-target size.

    Dataset sizes are read from ``PAPER[paper]["n_datapoints"]`` and matched
    to each target in ``PAPER[paper]["target"]``. Score metrics use
    ``<metric>_seed_fold_scores`` when available; scalar columns are plotted
    directly. For GP models, pass ``kernel_triples`` as one triple or a list of
    ``(fp_kernel, count_kernel, mixing_method)`` triples.
    """
    df = _filter_kernel_triples(df, kernel_triples, keep_missing=True)
    is_scalar_metric = metric in df.columns and f"{metric}_seed_fold_scores" not in df.columns
    plot_df = _expand_master_scores_for_profile(
        df,
        metric=metric,
        model=model,
    )
    if plot_df.empty:
        raise ValueError(
            f"No {metric} values found for model={model} "
            f"with kernel_triples={kernel_triples}."
        )

    plot_df = _add_datapoint_counts(plot_df)

    hue_col = "model configuration"
    plot_df[hue_col] = plot_df.apply(
        lambda row: _model_config_label(row, include_kernel_config),
        axis=1,
    )

    if dataset_as_experiment_points:
        group_cols = [
            "dataset",
            "target",
            "n datapoints",
            hue_col,
            "model",
            "fp kernel",
            "count kernel",
            "mixing method",
        ]
        plot_df = (
            plot_df.groupby(group_cols, dropna=False, as_index=False)[metric]
            .mean()
            .copy()
        )

    selected_models = _selection_values(model) or MODELS
    model_order = {str(model_name).lower(): idx for idx, model_name in enumerate(selected_models)}
    config_order = plot_df[
        [hue_col, "model", "fp kernel", "count kernel", "mixing method"]
    ].drop_duplicates(subset=[hue_col])
    config_order["sort_key"] = config_order.apply(
        lambda row: _model_config_sort_key(row, model_order),
        axis=1,
    )
    hue_order = config_order.sort_values("sort_key", kind="mergesort")[hue_col].tolist()
    data_number_order = sorted(plot_df["n datapoints"].dropna().unique())
    palette = {
        row[hue_col]: INDIVIDUAL_MODEL_COLORS.get(
            row[hue_col],
            INDIVIDUAL_MODEL_COLORS.get(
                str(row["model"]),
                _model_type_color(row["model"]),
            ),
        )
        for _, row in config_order.iterrows()
        if row[hue_col] in hue_order
    }

    fig, ax = plt.subplots(figsize=figsize)
    sns.boxplot(
        data=plot_df,
        x="n datapoints",
        y=metric,
        hue=hue_col,
        order=data_number_order,
        hue_order=hue_order,
        ax=ax,
        width=0.72,
        palette=palette,
        linewidth=1,
        fliersize=2.5,
    )

    ax.set_xlabel(x_label, fontsize=fontsize, fontweight="bold")
    ax.set_ylabel(y_label or metric.capitalize(), fontsize=fontsize, fontweight="bold")
    if title is not None:
        ax.set_title(title, fontsize=fontsize + 2)
    ax.tick_params(axis="both", labelsize=fontsize - 2)
    ax.set_xticks(range(len(data_number_order)))
    ax.set_xticklabels(
        [
            f"{int(value)}" if float(value).is_integer() else f"{value:g}"
            for value in data_number_order
        ],
        rotation=x_tick_rotation,
        ha="right" if x_tick_rotation else "center",
    )
    if log_y:
        positive_values = pd.to_numeric(plot_df[metric], errors="coerce")
        if (positive_values.dropna() <= 0).any():
            raise ValueError("log_y=True requires all plotted values to be positive.")
        ax.set_yscale("log")

    if y_lim is not None:
        ax.set_ylim(*y_lim)
    elif log_y:
        ax.set_ylim(bottom=pd.to_numeric(plot_df[metric], errors="coerce").min() * 0.8)
    elif metric == "r2":
        ax.set_ylim(0, 1.05)
    elif is_scalar_metric or not _metric_higher_is_better(metric):
        ax.set_ylim(bottom=0)

    legend = ax.get_legend()
    if legend is not None:
        legend.set_title(None)
        for text in legend.get_texts():
            text.set_fontsize(fontsize - 3)

    plt.tight_layout()

    if save_dir is not None:
        save_dir = ensure_long_path(Path(save_dir))
        os.makedirs(save_dir, exist_ok=True)
        if file_name is None:
            model_name = "_".join(str(value) for value in (_selection_values(model) or ["all"]))
            file_name = f"{model_name}_{metric}_performance_vs_data_number.png"
        fig.savefig(
            ensure_long_path(save_dir / file_name),
            bbox_inches="tight",
            format="png",
            dpi=900 if high_quality else 100,
        )

    if show:
        plt.show()
    else:
        plt.close(fig)

    return plot_df


def plot_model_average_performance_vs_data_number(
    df: pd.DataFrame,
    metric: str = "r2",
    model: Any = "GPytorchMAP",
    kernel_triples: Any = None,
    include_kernel_config: bool = True,
    dataset_as_experiment_points: bool = True,
    figsize: tuple = (10, 5),
    fontsize: int = 12,
    title: Optional[str] = None,
    x_label: str = "Number of datapoints",
    y_label: Optional[str] = None,
    x_tick_rotation: int = 0,
    y_lim: Optional[tuple] = None,
    log_y: bool = False,
    show_values: bool = True,
    show: bool = True,
    high_quality: bool = True,
    save_dir: Optional[Path] = HERE / "result_analysis",
    file_name: Optional[str] = None,
) -> pd.DataFrame:
    """
    Make a bar plot of average model performance by dataset-target size.

    This follows ``plot_model_performance_vs_data_number`` for metric expansion,
    model/kernel filtering, and datapoint lookup, then averages the selected
    metric for each ``n datapoints`` and model configuration before plotting.
    """
    df = _filter_kernel_triples(df, kernel_triples, keep_missing=True)
    is_scalar_metric = metric in df.columns and f"{metric}_seed_fold_scores" not in df.columns
    plot_df = _expand_master_scores_for_profile(
        df,
        metric=metric,
        model=model,
    )
    if plot_df.empty:
        raise ValueError(
            f"No {metric} values found for model={model} "
            f"with kernel_triples={kernel_triples}."
        )

    plot_df = _add_datapoint_counts(plot_df)
    hue_col = "model configuration"
    plot_df[hue_col] = plot_df.apply(
        lambda row: _model_config_label(row, include_kernel_config),
        axis=1,
    )

    if dataset_as_experiment_points:
        dataset_group_cols = [
            "dataset",
            "target",
            "n datapoints",
            hue_col,
            "model",
            "fp kernel",
            "count kernel",
            "mixing method",
        ]
        plot_df = (
            plot_df.groupby(dataset_group_cols, dropna=False, as_index=False)[metric]
            .mean()
            .copy()
        )

    summary_group_cols = [
        "n datapoints",
        hue_col,
        "model",
        "fp kernel",
        "count kernel",
        "mixing method",
    ]
    summary_df = (
        plot_df.groupby(summary_group_cols, dropna=False)[metric]
        .agg(["mean", "std", "count"])
        .reset_index()
        .rename(columns={
            "mean": f"{metric}_mean",
            "std": f"{metric}_std",
            "count": f"{metric}_count",
        })
    )
    summary_df[f"{metric}_std"] = summary_df[f"{metric}_std"].fillna(0.0)

    selected_models = _selection_values(model) or MODELS
    model_order = {str(model_name).lower(): idx for idx, model_name in enumerate(selected_models)}
    config_order = summary_df[
        [hue_col, "model", "fp kernel", "count kernel", "mixing method"]
    ].drop_duplicates(subset=[hue_col])
    config_order["sort_key"] = config_order.apply(
        lambda row: _model_config_sort_key(row, model_order),
        axis=1,
    )
    hue_order = config_order.sort_values("sort_key", kind="mergesort")[hue_col].tolist()
    data_number_order = sorted(summary_df["n datapoints"].dropna().unique())
    order_lookup = {label: idx for idx, label in enumerate(hue_order)}
    summary_df["_plot_order"] = summary_df[hue_col].map(order_lookup)
    summary_df = (
        summary_df.sort_values(["n datapoints", "_plot_order"], kind="stable")
        .drop(columns="_plot_order")
        .copy()
    )

    palette = {
        row[hue_col]: INDIVIDUAL_MODEL_COLORS.get(
            row[hue_col],
            INDIVIDUAL_MODEL_COLORS.get(
                str(row["model"]),
                _model_type_color(row["model"]),
            ),
        )
        for _, row in config_order.iterrows()
        if row[hue_col] in hue_order
    }

    fig, ax = plt.subplots(figsize=figsize)
    sns.barplot(
        data=summary_df,
        x="n datapoints",
        y=f"{metric}_mean",
        hue=hue_col,
        order=data_number_order,
        hue_order=hue_order,
        palette=palette,
        width=0.72,
        errorbar=None,
        ax=ax,
    )

    summary_lookup = (
        summary_df.groupby(["n datapoints", hue_col], dropna=False)
        .agg({
            f"{metric}_mean": "mean",
            f"{metric}_std": "mean",
        })
    )
    max_label_y = None
    for hue_idx, container in enumerate(ax.containers[:len(hue_order)]):
        if hue_idx >= len(hue_order):
            continue
        hue_value = hue_order[hue_idx]
        for data_idx, patch in enumerate(container.patches):
            if data_idx >= len(data_number_order):
                continue
            data_value = data_number_order[data_idx]
            if (data_value, hue_value) not in summary_lookup.index:
                continue

            mean_value = float(summary_lookup.loc[(data_value, hue_value), f"{metric}_mean"])
            std_value = float(summary_lookup.loc[(data_value, hue_value), f"{metric}_std"])
            if np.isnan(mean_value):
                continue

            x_position = patch.get_x() + patch.get_width() / 2
            if std_value > 0:
                ax.errorbar(
                    x_position,
                    mean_value,
                    yerr=std_value,
                    fmt="none",
                    ecolor="black",
                    elinewidth=1.1,
                    capsize=3,
                    capthick=1.1,
                    zorder=4,
                )

            label_y = mean_value + (std_value if std_value > 0 else 0) + 0.01
            max_label_y = label_y if max_label_y is None else max(max_label_y, label_y)
            if show_values:
                bottom, top = ax.get_ylim()
                value_y = bottom + 0.02 * (top - bottom)
                ax.text(
                    x_position,
                    value_y,
                    f"{mean_value:.2f}",
                    ha="center",
                    va="bottom",
                    fontsize=max(fontsize - 10, 4),
                    fontweight="bold",
                    color="white",
                    rotation=90,
                )

    ax.set_xlabel(x_label, fontsize=fontsize, fontweight="bold")
    ax.set_ylabel(y_label or f"Mean {metric}", fontsize=fontsize, fontweight="bold")
    if title is not None:
        ax.set_title(title, fontsize=fontsize + 2)
    ax.tick_params(axis="both", labelsize=fontsize - 2)
    ax.set_xticks(range(len(data_number_order)))
    ax.set_xticklabels(
        [
            f"{int(value)}" if float(value).is_integer() else f"{value:g}"
            for value in data_number_order
        ],
        rotation=x_tick_rotation,
        ha="right" if x_tick_rotation else "center",
    )
    if log_y:
        positive_values = pd.to_numeric(summary_df[f"{metric}_mean"], errors="coerce")
        if (positive_values.dropna() <= 0).any():
            raise ValueError("log_y=True requires all plotted values to be positive.")
        ax.set_yscale("log")

    if y_lim is not None:
        ax.set_ylim(*y_lim)
    elif log_y:
        ax.set_ylim(
            bottom=pd.to_numeric(summary_df[f"{metric}_mean"], errors="coerce").min() * 0.8
        )
    elif str(metric).strip().lower() in {"r2", "oof_r2"} or _is_feature_stability_metric(metric):
        ax.set_ylim(0, 1.05)
    elif is_scalar_metric or not _metric_higher_is_better(metric):
        ax.set_ylim(bottom=0)
    if max_label_y is not None:
        bottom, top = ax.get_ylim()
        ax.set_ylim(bottom, max(top, max_label_y + 0.05))

    legend = ax.get_legend()
    if legend is not None:
        handles, labels = ax.get_legend_handles_labels()
        legend.remove()
        legend = ax.legend(
            handles,
            labels,
            loc="lower center",
            bbox_to_anchor=(0.5, 1.02),
            ncol=max(1, min(3, len(labels))),
            frameon=False,
            title=None,
        )
        for text in legend.get_texts():
            text.set_fontsize(fontsize - 3)

    plt.tight_layout(rect=(0, 0, 1, 0.92))

    if save_dir is not None:
        save_dir = ensure_long_path(Path(save_dir))
        os.makedirs(save_dir, exist_ok=True)
        if file_name is None:
            model_name = "_".join(str(value) for value in (_selection_values(model) or ["all"]))
            file_name = f"{model_name}_{metric}_average_performance_vs_data_number.png"
        fig.savefig(
            ensure_long_path(save_dir / file_name),
            bbox_inches="tight",
            format="png",
            dpi=900 if high_quality else 100,
        )

    if show:
        plt.show()
    else:
        plt.close(fig)

    return summary_df





if __name__ == "__main__":

    build_master_performance_data(save_path=RESULTS/"master_performance_data"/"Tree_and_GP",
                                   score_metrics=DEFAULT_SCORE_METRICS)

    # COMBINED_RESULTS: pd.DataFrame = RESULTS / "master_performance_data"/ "Tree_and_GP.pkl"
    # result_df = pd.read_pickle(ensure_long_path(COMBINED_RESULTS))

    # prof_results = performance_plot_with_ranks(
    #     df=result_df,
    #     metric="r2",
    #     model=["GPytorchMAP"],
    #     fp_kernels=["TanimotoRBF", "TanimotoMatern32", "TanimotoMatern52", "Tanimoto"],
    #     count_kernels=["RBF", "Matern32", "Matern52"],
    #     # title="GPyTorch MAP R² Performance Profile",
    #     dataset_as_experiment_points=False,
    #     # show_error_band=True,
    #     show=False,
    #     save_dir= HERE / "result_analysis",
    #     file_name="r2_GpytorchMAP_SK_performance_profile.png",
    # )
    #"r2", "nll", "cvpp_ama", "ece"
    # for metric in ["feature_stability"]:
    #     tree_fi = "MDI"  # or "SHAP"
    #     prof_results = performance_plot_with_ranks(
    #         df=result_df,
    #         metric=metric,
    #         model=["RF", "XGBR", "NGB", "GPytorchMAP", "GpyroHMC"],
    #         kernel_triples=[
    #             ("Matern32", "Matern32", "product"),
    #             ("TanimotoMatern32", "Matern32", "product"),
    #             # ("Graph", "Matern32", "product"),
    #         ],
    #         tree_feature_importance=tree_fi,
    #         legend_ncols=2,
    #         fontsize=17,
    #         figsize=(7, 5),
    #         show=True,
    #         show_error_band=False,
    #         save_dir= HERE / "result_analysis",
    #         file_name=f"{metric}_{tree_fi}_lengthscale_model_performance_profile_curve_comparison.png",
    #     )
    #     plot_model_profile_comparison(
    #         prof_results=prof_results,
    #         model=["RF", "XGBR", "NGB", "GPytorchMAP", "GpyroHMC"],
    #         kernel_triples=[
    #             ("Matern32", "Matern32", "product"),
    #             ("TanimotoMatern32", "Matern32", "product"),
    #             # ("Graph", "Matern32", "product"),
    #         ],
    #         metric=metric,
    #         tree_feature_importance=tree_fi,
    #         y_label=f"Profile AUC of feature stability",
    #         fontsize=17,
    #         figsize=(5, 5),
    #         save_dir=HERE / "result_analysis",
    #         file_name=f"{metric}_{tree_fi}_lengthscale_model_profile_comparison.png",
    #     )

    # plot_profile_auc_heatmap(
    #     auc_df=prof_results,
    #     fontsize=13.4,
    #     figsize=(15, 4.1),
    #     model=["GPytorchMAP"],
    #     fp_kernels=["RBF", "Matern32", "Matern52"],
    #     count_kernels=["RBF", "Matern32", "Matern52"],
    #     # title="GPyTorch MAP R² Profile AUC Heatmap",
    #     save_dir=HERE / "result_analysis",
    #     file_name="GPytorchMAP_bitwise_r2_profile_auc_heatmap.png",
    # )


    # plot_model_comparison(
    #     df=result_df,
    #     metric="OOF_R2",
    #     model=["RF", "XGBR","NGB","GPytorchMAP", "GpyroHMC","MGK"],
    #     kernel_triples=[
    #         ("Matern32", "Matern32", "product"),
    #         ("TanimotoMatern32", "Matern32", "product"),
    #         ("Graph", "Matern32", "product"),
    #         ],
    #     y_label="R²",
    #     fontsize=17,
    #     show=True,
    #     y_lim=(0,1.05),
    #     figsize=(6, 5),
    #     # log_y=True,
    #     save_dir=HERE / "result_analysis",
    #     file_name="OOF_R2_distributional_model_comparison.png",
    # )

    # plot_hybridization_method_comparison(
    #     df=result_df,
    #     metric="r2",
    #     model="GPytorchMAP",
    #     fp_kernels=["TanimotoRBF", "TanimotoMatern32", "TanimotoMatern52", "Tanimoto"],
    #     count_kernels=["RBF", "Matern32", "Matern52"],
    #     figsize=(5, 5),
    #     fontsize=17,
    #     y_label="R²",
    #     show=True,
    #     save_dir=HERE / "result_analysis",
    #     file_name="r2_GPytorchMAP_SK_hybridization_comparison.png",
    # )
    
    # plot_hybridization_profile_comparison(
    #     prof_results=prof_results,
    #     model="GPytorchMAP",
    #     metric="r2",
    #     fp_kernels=["TanimotoRBF", "TanimotoMatern32", "TanimotoMatern52", "Tanimoto"],
    #     count_kernels=["RBF", "Matern32", "Matern52"],
    #     y_label="Profile AUC of R²",
    #     fontsize=17,
    #     figsize=(5, 5),
    #     save_dir=HERE / "result_analysis",
    #     file_name="r2_GPytorchMAP_SK_hybridization_profile_comparison_avg_over_config.png",
    # )

    # plot_hybridization_profile_comparison(
    # prof_results=hybrid_prof_results,
    # model="GPytorchMAP",
    # metric="r2",
    # y_label="Profile AUC of R²",
    # fontsize=17,
    # figsize=(5, 5),
    # save_dir=HERE / "result_analysis",
    # file_name="r2_GPytorchMAP_bitwise_hybridization_profile_comparison_pure.png",
    # )

    # plot_hybridization_topsis_comparison(
    #     df=result_df,
    #     metrics=["OOF_R2"],
    #     criteria_weights=[1],
    #     model="GPytorchMAP",
    #     fp_kernels=["TanimotoRBF", "TanimotoMatern32", "TanimotoMatern52", "Tanimoto"],
    #     count_kernels=["RBF", "Matern32", "Matern52"],
    #     show=True,
    #     fontsize=17,
    #     figsize=(5, 6),
    #     save_dir=HERE / "result_analysis",
    #     file_name="OOF_R2_TOPSIS_GPytorchMAP_SK_hybridization_profile_comparison_avg_over_config.png",
    # )
    


    # plot_model_performance_vs_data_number(
    #     df=result_df,
    #     metric="OOF_R2",
    #     model=["RF", "XGBR", "NGB", "GPytorchMAP", "GpyroHMC", "MGK"],
    #     kernel_triples=[
    #         ("Matern32", "Matern32", "product"),
    #         ("TanimotoMatern32", "Matern32", "product"),
    #         ("Graph", "Matern32", "product"),
    #     ],
    #     y_label=label_conversion_source["OOF_R2"],
    #     x_label="# Datapoints",
    #     fontsize=18,
    #     y_lim=(0, 1.05),
    #     # log_y=True,
    #     figsize=(9, 5),
    #     show=True,
    #     save_dir=HERE / "result_analysis",
    #     file_name="OOF_R2_model_performance_vs_data_number.png",
    # )


    # plot_model_average_performance_vs_data_number(
    #     df=result_df,
    #     metric="OOF_R2",
    #     model=["RF", "XGBR", "NGB", "GPytorchMAP", "GpyroHMC", "MGK"],
    #     kernel_triples=[
    #         ("Matern32", "Matern32", "product"),
    #         ("TanimotoMatern32", "Matern32", "product"),
    #         ("Graph", "Matern32", "product"),
    #     ],
    #     y_label=label_conversion_source["OOF_R2"],
    #     x_label="# Datapoints",
    #     fontsize=18,
    #     y_lim=(0, 1.05),
    #     # log_y=True,
    #     figsize=(9, 7),
    #     show=True,
    #     save_dir=HERE / "result_analysis",
    #     file_name="OOF_R2_model_performance_vs_data_number.png",
    # )


    # plot_model_performance_TOPSIS(
    #     df=result_df,
    #     metrics=["OOF_R2"],
    #     criteria_weights=[1],
    #     model=["RF", "XGBR", "NGB", "GPytorchMAP", "GpyroHMC", "MGK"],
    #     kernel_triples=[
    #         ("TanimotoMatern32", "Matern32", "product"),
    #         ("Graph", "Matern32", "product"),
    #     ],
    #     show=True,
    #     fontsize=17,
    #     figsize=(5, 5),
    #     save_dir=HERE / "result_analysis",
    #     file_name=f"OOF_R2_TOPSIS_model_comparison.png",
    # )

    # plot_model_feature_importance_stability_comparison(
    #     result_df,
    #     model=["RF", "XGBR", "NGB", "GPytorchMAP", "GpyroHMC"],
    #     tree_feature_importance="MDI",
    #     kernel_triples=[
    #         ("Matern32", "Matern32", "product"),
    #         ("TanimotoMatern32", "Matern32", "product"),
    #         ],
    #     y_label="Stability (Kendall's W)",
    #     fontsize=19,
    #     show=True,
    #     y_lim=(0,1.05),
    #     figsize=(6, 5),
    #     # log_y=True,
    #     save_dir=HERE / "result_analysis",
    #     file_name="feature_importance_stability_lengthscale_MDI_comparison.png",
    # )
